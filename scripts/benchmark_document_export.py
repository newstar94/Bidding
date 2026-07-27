"""Benchmark detailed-evaluation DOCX rendering through the production pipeline."""

from __future__ import annotations

import argparse
from concurrent.futures import ThreadPoolExecutor
from copy import deepcopy
from io import BytesIO
import json
import math
from pathlib import Path
from statistics import median
import sys
import tempfile
import time

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from backend.documents.custom_exporter import generate_report_from_custom_template
from backend.documents.document_worker import run_document_job
from backend.documents.docx_context_policy import seal_docx_context


DEFAULT_ROW_COUNTS = (10, 100, 500)
DEFAULT_ITERATIONS = 5
MAX_BENCHMARK_ROWS = 5_000
_MAPPINGS = [("ds_dgct", "detailed_evaluation_rows", "")]


def nearest_rank_percentile(values: list[float], percentile: float) -> float:
    if not values:
        raise ValueError("Không có số đo để tính percentile.")
    if not 0 < percentile <= 1:
        raise ValueError("Percentile phải nằm trong khoảng (0, 1].")
    ordered = sorted(float(value) for value in values)
    return ordered[max(0, math.ceil(len(ordered) * percentile) - 1)]


def build_benchmark_template(template_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=4, cols=5)
    for cell, value in zip(
        table.rows[0].cells,
        ("STT", "Tiêu chí", "Yêu cầu", "Nội dung HSDT", "Nhận xét"),
        strict=True,
    ):
        cell.text = value
    table.rows[1].cells[0].text = "{%tr for tc in ds_dgct %}"
    for cell, value in zip(
        table.rows[2].cells,
        (
            "{{ tc.stt }}",
            "{{ tc.ten_tieu_chi }}",
            "{{ tc.yeu_cau }}",
            "{{ tc.noi_dung_hsdt }}",
            "{{ tc.nhan_xet }}",
        ),
        strict=True,
    ):
        cell.text = value
    table.rows[3].cells[0].text = "{%tr endfor %}"
    document.save(template_path)


def build_benchmark_context(row_count: int) -> tuple[dict, dict, float]:
    if row_count < 1 or row_count > MAX_BENCHMARK_ROWS:
        raise ValueError(
            f"Số dòng benchmark phải từ 1 đến {MAX_BENCHMARK_ROWS}."
        )
    rows = [
        {
            "stt": str(index + 1),
            "ten_tieu_chi": f"Tiêu chí đánh giá chi tiết số {index + 1}",
            "yeu_cau": "Đáp ứng đầy đủ yêu cầu của hồ sơ mời thầu",
            "noi_dung_hsdt": "Nhà thầu cung cấp tài liệu chứng minh tương ứng",
            "nhan_xet": "Đạt yêu cầu",
        }
        for index in range(row_count)
    ]
    started_at = time.perf_counter()
    context, manifest = seal_docx_context(
        "evaluation",
        {"ds_dgct": rows},
        _MAPPINGS,
    )
    seal_ms = (time.perf_counter() - started_at) * 1_000
    return context, manifest, seal_ms


def _summary(values: list[float]) -> dict[str, float]:
    return {
        "medianMs": round(median(values), 3),
        "p95Ms": round(nearest_rank_percentile(values, 0.95), 3),
    }


def benchmark_case(
    template_path: Path,
    row_count: int,
    iterations: int,
    mode: str = "both",
    parallel_jobs: int = 0,
) -> dict:
    if iterations < 1 or iterations > 50:
        raise ValueError("Số lượt benchmark phải từ 1 đến 50.")
    if mode not in {"direct", "worker", "both"}:
        raise ValueError("Chế độ benchmark không hợp lệ.")
    if parallel_jobs < 0 or parallel_jobs > 32:
        raise ValueError("Số tác vụ song song phải từ 0 đến 32.")

    seal_durations: list[float] = []
    context = manifest = None
    for _ in range(iterations):
        context, manifest, seal_ms = build_benchmark_context(row_count)
        seal_durations.append(seal_ms)
    assert context is not None and manifest is not None

    result: dict = {
        "rows": row_count,
        "iterations": iterations,
        "contextSeal": _summary(seal_durations),
    }
    output = b""
    if mode in {"direct", "both"}:
        durations = []
        for _ in range(iterations):
            render_context = deepcopy(context)
            started_at = time.perf_counter()
            stream = generate_report_from_custom_template(
                template_path,
                render_context,
                manifest,
            )
            durations.append((time.perf_counter() - started_at) * 1_000)
            output = stream.getvalue()
        result["directRender"] = _summary(durations)

    if mode in {"worker", "both"}:
        durations = []
        for _ in range(iterations):
            started_at = time.perf_counter()
            output = run_document_job(
                "render_docx",
                {
                    "template_path": str(template_path),
                    "context": context,
                    "context_manifest": manifest,
                },
                timeout_seconds=60,
            )
            durations.append((time.perf_counter() - started_at) * 1_000)
        result["isolatedWorker"] = _summary(durations)

        if parallel_jobs:
            payload = {
                "template_path": str(template_path),
                "context": context,
                "context_manifest": manifest,
            }

            def render_one() -> tuple[float, bytes]:
                started_at = time.perf_counter()
                rendered_bytes = run_document_job(
                    "render_docx",
                    payload,
                    timeout_seconds=60,
                )
                return (time.perf_counter() - started_at) * 1_000, rendered_bytes

            batch_started_at = time.perf_counter()
            with ThreadPoolExecutor(max_workers=parallel_jobs) as executor:
                parallel_results = list(
                    executor.map(lambda _index: render_one(), range(parallel_jobs))
                )
            batch_ms = (time.perf_counter() - batch_started_at) * 1_000
            parallel_durations = [item[0] for item in parallel_results]
            output = parallel_results[-1][1]
            result["parallelWorker"] = {
                "jobs": parallel_jobs,
                "completed": len(parallel_results),
                "batchMs": round(batch_ms, 3),
                **_summary(parallel_durations),
            }

    rendered = Document(BytesIO(output))
    rendered_rows = len(rendered.tables[0].rows) - 1
    if rendered_rows != row_count:
        raise RuntimeError(
            f"DOCX kết quả có {rendered_rows} dòng, dự kiến {row_count}."
        )
    result["outputBytes"] = len(output)
    result["renderedRows"] = rendered_rows
    return result


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--rows",
        nargs="+",
        type=int,
        default=list(DEFAULT_ROW_COUNTS),
    )
    parser.add_argument("--iterations", type=int, default=DEFAULT_ITERATIONS)
    parser.add_argument(
        "--mode",
        choices=("direct", "worker", "both"),
        default="both",
    )
    parser.add_argument("--parallel-jobs", type=int, default=0)
    args = parser.parse_args()

    with tempfile.TemporaryDirectory(
        prefix="biddingflow-document-export-benchmark-"
    ) as raw_directory:
        template_path = Path(raw_directory) / "evaluation-template.docx"
        build_benchmark_template(template_path)
        results = [
            benchmark_case(
                template_path,
                row_count,
                args.iterations,
                args.mode,
                args.parallel_jobs,
            )
            for row_count in args.rows
        ]
    print(
        json.dumps(
            {
                "mode": args.mode,
                "iterations": args.iterations,
                "results": results,
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
