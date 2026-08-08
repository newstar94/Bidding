"""Seed, verify, and remove an isolated joint-venture browser E2E package."""

from __future__ import annotations

import json
import os
from pathlib import Path
import shutil
import sys
import time

import psycopg
from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from backend.auth.auth_helper import hash_password
from backend.shared.paths import WORD_TEMPLATE_DIR


def _database_url() -> str:
    value = str(os.environ.get("DATABASE_URL") or "").strip()
    if not value:
        raise RuntimeError("DATABASE_URL is required")
    return value


def _payload() -> dict:
    return json.loads(sys.stdin.buffer.read().decode("utf-8"))


def _account_id(cursor, username: str) -> str:
    row = cursor.execute(
        """SELECT id FROM tai_khoan
            WHERE lower(COALESCE(username_norm, '')) = lower(%s)
               OR lower(COALESCE(ten_dang_nhap, '')) = lower(%s)
            LIMIT 1""",
        (username, username),
    ).fetchone()
    if not row:
        raise RuntimeError(f"E2E account {username!r} does not exist")
    return str(row[0])


def _setup(data: dict) -> dict:
    run_id = str(data["runId"])
    organization_id = str(data["organizationId"])
    package = data["package"]
    owner_id = f"{run_id}-owner"
    plan_id = f"{run_id}-plan"
    with psycopg.connect(_database_url()) as connection:
        with connection.cursor() as cursor:
            account = data.get("account")
            if account:
                user_id = str(account["id"])
                cursor.execute(
                    """INSERT INTO tai_khoan (
                           id, ten_dang_nhap, username_norm, mat_khau, ho_ten,
                           vai_tro, email, email_norm, da_xac_minh, username_da_dat
                       ) VALUES (%s, %s, %s, %s, %s, 'user', %s, %s, 1, 1)""",
                    (
                        user_id, account["username"], str(account["username"]).lower(),
                        hash_password(str(data["password"])), account["name"],
                        account["email"], str(account["email"]).lower(),
                    ),
                )
            else:
                user_id = _account_id(cursor, str(data["username"]))
            cursor.execute(
                "INSERT INTO to_chuc (id, ten_to_chuc) VALUES (%s, %s)",
                (organization_id, f"Joint venture E2E {run_id}"),
            )
            now = int(time.time())
            cursor.execute(
                """INSERT INTO organization_subscriptions (
                       organization_id, package_id, status, starts_at,
                       expires_at, member_quota
                   ) VALUES (%s, 'diamond', 'active', %s, %s, 999)""",
                (organization_id, now, now + 86400),
            )
            cursor.execute(
                """INSERT INTO danh_muc_trang_thai_hop_dong
                       (id, organization_id, owner_type, name, color)
                   VALUES (%s, %s, 'organization', 'Đang thực hiện', '#2563eb')""",
                (f"{run_id}-contract-status-active", organization_id),
            )
            cursor.execute(
                """INSERT INTO thanh_vien_to_chuc (
                       user_id, organization_id, vai_tro_trong_to_chuc,
                       ten_nhan_su, trang_thai_thanh_vien
                   ) VALUES (%s, %s, 'manager', %s, 'active')""",
                (user_id, organization_id, "Quản lý JV E2E"),
            )
            cursor.execute(
                """INSERT INTO chu_dau_tu (
                       id, organization_id, owner_type, id_goc,
                       ma_chu_dau_tu, ma_so_thue, ten_chu_dau_tu
                   ) VALUES (%s, %s, 'organization', %s, %s, %s, %s)""",
                (owner_id, organization_id, owner_id, f"{run_id}-CDT", "0100000001", f"Chủ đầu tư {run_id}"),
            )
            cursor.execute(
                """INSERT INTO ke_hoach_lcnt (
                       id, organization_id, owner_type, id_goc, ma_ke_hoach,
                       ten_ke_hoach, ten_du_an_du_toan, loai_hinh_mua_sam,
                       chu_dau_tu_id, ngay_phe_duyet, quyet_dinh_phe_duyet
                   ) VALUES (%s, %s, 'organization', %s, %s, %s, %s, %s, %s, %s, %s)""",
                (
                    plan_id,
                    organization_id,
                    plan_id,
                    f"{run_id}-KH",
                    f"Kế hoạch {run_id}",
                    f"Dự toán {run_id}",
                    "Mua sắm hàng hóa",
                    owner_id,
                    "2026-07-01",
                    f"{run_id}/QD-KH",
                ),
            )
            for index, contractor in enumerate(data["contractors"], start=1):
                cursor.execute(
                    """INSERT INTO nha_thau (
                           id, organization_id, owner_type, id_goc,
                           ma_nha_thau, ten_nha_thau, ma_so_thue, loai_nha_thau
                       ) VALUES (%s, %s, 'organization', %s, %s, %s, %s, 'Độc lập')""",
                    (
                        contractor["id"],
                        organization_id,
                        contractor["id"],
                        contractor["code"],
                        contractor["name"],
                        f"01000000{index:02d}",
                    ),
                )
            cursor.execute(
                """INSERT INTO goi_thau (
                       id, organization_id, owner_type, id_goc, ma_goi_thau,
                       ke_hoach_id, ten_goi_thau, gia_goi_thau, linh_vuc,
                       hinh_thuc_lua_chon, phuong_thuc_lua_chon,
                       phuong_phap_danh_gia, thoi_gian_thuc_hien, nguon_von,
                       thoi_gian_to_chuc, thoi_gian_bat_dau_to_chuc,
                       thoi_gian_dang_tai, thoi_gian_dong_thau, thoi_gian_mo_thau,
                       gia_tri_dam_bao_du_thau, hieu_luc_hsdt,
                       hieu_luc_dam_bao_du_thau, phan_lo, trang_thai
                   ) VALUES (
                       %s, %s, 'organization', %s, %s, %s, %s, %s,
                       'Xây lắp', 'Đấu thầu rộng rãi',
                       'Một giai đoạn một túi hồ sơ', 'Giá thấp nhất',
                       '90 ngày', 'Nguồn vốn E2E', '30 ngày', '2026-07-01',
                       '2026-07-01 08:00:00', '2026-07-05 08:00:00',
                       '2026-07-05 08:05:00', 10000, 90, 120,
                       'Không', 'OPENED'
                   )""",
                (
                    package["id"],
                    organization_id,
                    package["id"],
                    package["code"],
                    plan_id,
                    package["name"],
                    int(package["price"]),
                ),
            )
            cursor.execute(
                """INSERT INTO phan_cong_nhan_su (
                       id, organization_id, owner_type, id_nhan_vien,
                       id_muc_tieu, loai_doi_tuong
                   ) VALUES (%s, %s, 'organization', %s, %s, 'goithau')""",
                (f"{run_id}-package-assignment", organization_id, user_id, package["id"]),
            )
            lot_package = data.get("lotPackage")
            if lot_package:
                cursor.execute(
                    """INSERT INTO goi_thau (
                           id, organization_id, owner_type, id_goc, ma_goi_thau,
                           ke_hoach_id, ten_goi_thau, gia_goi_thau, linh_vuc,
                           hinh_thuc_lua_chon, phuong_thuc_lua_chon,
                           phuong_phap_danh_gia, thoi_gian_thuc_hien, nguon_von,
                           thoi_gian_to_chuc, thoi_gian_bat_dau_to_chuc,
                           thoi_gian_dang_tai, thoi_gian_dong_thau, thoi_gian_mo_thau,
                           hieu_luc_hsdt, hieu_luc_dam_bao_du_thau, phan_lo, trang_thai
                       ) VALUES (
                           %s, %s, 'organization', %s, %s, %s, %s, %s,
                           'Xây lắp', 'Đấu thầu rộng rãi',
                           'Một giai đoạn một túi hồ sơ', 'Giá thấp nhất',
                           '90 ngày', 'Nguồn vốn E2E', '30 ngày', '2026-07-01',
                           '2026-07-01 08:00:00', '2026-07-05 08:00:00',
                           '2026-07-05 08:05:00', 90, 120, 'Có', 'OPENED'
                       )""",
                    (
                        lot_package["id"], organization_id, lot_package["id"],
                        lot_package["code"], plan_id, lot_package["name"], int(lot_package["price"]),
                    ),
                )
                for sort_order, lot in enumerate(lot_package["lots"]):
                    cursor.execute(
                        """INSERT INTO goi_thau_phan_lo (
                               id, organization_id, owner_type, goi_thau_id,
                               ma_phan_lo, ten_phan_lo, gia_tri_phan_lo,
                               bao_dam_du_thau, thoi_gian_thuc_hien, sort_order
                           ) VALUES (%s, %s, 'organization', %s, %s, %s, %s, 10000, '90 ngày', %s)""",
                        (
                            lot["id"], organization_id, lot_package["id"], lot["code"],
                            lot["name"], int(lot["price"]), sort_order,
                        ),
                    )
                cursor.execute(
                    """INSERT INTO phan_cong_nhan_su (
                           id, organization_id, owner_type, id_nhan_vien,
                           id_muc_tieu, loai_doi_tuong
                       ) VALUES (%s, %s, 'organization', %s, %s, 'goithau')""",
                    (f"{run_id}-lot-package-assignment", organization_id, user_id, lot_package["id"]),
                )
            two_envelope_package = data.get("twoEnvelopePackage")
            if two_envelope_package:
                cursor.execute(
                    """INSERT INTO goi_thau (
                           id, organization_id, owner_type, id_goc, ma_goi_thau,
                           ke_hoach_id, ten_goi_thau, gia_goi_thau, linh_vuc,
                           hinh_thuc_lua_chon, phuong_thuc_lua_chon,
                           phuong_phap_danh_gia, thoi_gian_thuc_hien, nguon_von,
                           thoi_gian_to_chuc, thoi_gian_bat_dau_to_chuc,
                           thoi_gian_dang_tai, thoi_gian_dong_thau, thoi_gian_mo_thau,
                           gia_tri_dam_bao_du_thau, hieu_luc_hsdt,
                           hieu_luc_dam_bao_du_thau, phan_lo, trang_thai
                       ) VALUES (
                           %s, %s, 'organization', %s, %s, %s, %s, %s,
                           'Xây lắp', 'Đấu thầu rộng rãi',
                           'Một giai đoạn hai túi hồ sơ', 'Giá thấp nhất',
                           '90 ngày', 'Nguồn vốn E2E', '30 ngày', '2026-07-01',
                           '2026-07-01 08:00:00', '2026-07-05 08:00:00',
                           '2026-07-05 08:05:00', 10000, 90, 120,
                           'Không', 'OPENED'
                       )""",
                    (
                        two_envelope_package["id"], organization_id,
                        two_envelope_package["id"], two_envelope_package["code"],
                        plan_id, two_envelope_package["name"],
                        int(two_envelope_package["price"]),
                    ),
                )
                cursor.execute(
                    """INSERT INTO phan_cong_nhan_su (
                           id, organization_id, owner_type, id_nhan_vien,
                           id_muc_tieu, loai_doi_tuong
                       ) VALUES (%s, %s, 'organization', %s, %s, 'goithau')""",
                    (
                        f"{run_id}-two-envelope-assignment", organization_id,
                        user_id, two_envelope_package["id"],
                    ),
                )
    return {
        "organizationId": organization_id,
        "packageId": package["id"],
        "lotPackageId": data.get("lotPackage", {}).get("id"),
        "twoEnvelopePackageId": data.get("twoEnvelopePackage", {}).get("id"),
    }


def _create_word_template(data: dict) -> dict:
    output_path = Path(str(data["wordTemplatePath"])).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    title = document.add_heading("BÁO CÁO KẾT QUẢ ĐÁNH GIÁ HỒ SƠ DỰ THẦU", level=1)
    for run in title.runs:
        run.font.name = "Arial"
    document.add_paragraph("Gói thầu: {{ goi_thau.ten_goi_thau }}")
    document.add_paragraph("Giá trúng thầu: {{ goi_thau.gia_trung_thau }}")
    document.add_paragraph("{% for bid in nha_thau %}")
    document.add_paragraph("Nhà thầu: {{ bid.ten_nha_thau }}")
    document.add_paragraph("Loại nhà thầu: {{ bid.loai_nha_thau }}")
    document.add_paragraph("Giá đề nghị trúng thầu: {{ bid.gia_de_nghi_trung_thau }}")
    document.add_paragraph("{% if bid.chap_thuan_gia_de_nghi_trung_thau_duoi_50 %}")
    document.add_paragraph("Quyết định giá dưới 50%: Chấp thuận")
    document.add_paragraph("{% else %}")
    document.add_paragraph("Quyết định giá dưới 50%: Không chấp thuận")
    document.add_paragraph("{% endif %}")
    document.add_paragraph("{% for member in bid.thanh_vien_lien_danh %}")
    document.add_paragraph("Thành viên: {{ member.ten_nha_thau }} — {{ member.vai_tro }}")
    document.add_paragraph("{% endfor %}")
    document.add_paragraph("{% endfor %}")
    document.save(output_path)
    return {"path": str(output_path), "size": output_path.stat().st_size}


def _inspect_docx(data: dict) -> dict:
    export_path = Path(str(data["exportPath"])).resolve()
    document = Document(export_path)
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    text = "\n".join(values)
    for expected in data.get("expectedNames", []):
        if str(expected) not in text:
            raise AssertionError(f"DOCX is missing expected value: {expected}")
    if "Quyết định giá dưới 50%: Chấp thuận" not in text:
        raise AssertionError("DOCX is missing the accepted low-price decision")
    if "{{" in text or "{%" in text:
        raise AssertionError("DOCX still contains unresolved template syntax")
    fonts = sorted({
        run.font.name
        for paragraph in document.paragraphs
        for run in paragraph.runs
        if run.font.name
    })
    return {
        "path": str(export_path),
        "size": export_path.stat().st_size,
        "paragraphs": len(document.paragraphs),
        "fonts": fonts,
    }


def _verify(data: dict) -> dict:
    organization_id = str(data["organizationId"])
    package_id = str(data["package"]["id"])
    with psycopg.connect(_database_url()) as connection:
        with connection.cursor() as cursor:
            rows = cursor.execute(
                """SELECT opening.id, opening.loai_nha_thau,
                          count(member.id),
                          sum(CASE WHEN member.vai_tro = 'Đứng đầu liên danh' THEN 1 ELSE 0 END)
                     FROM thong_tin_mo_thau AS opening
                     LEFT JOIN thong_tin_mo_thau_lien_danh_thanh_vien AS member
                       ON member.organization_id = opening.organization_id
                      AND member.thong_tin_mo_thau_id = opening.id
                    WHERE opening.organization_id = %s AND opening.goi_thau_id = %s
                    GROUP BY opening.id, opening.loai_nha_thau
                    ORDER BY opening.id""",
                (organization_id, package_id),
            ).fetchall()
            if len(rows) != 2:
                raise AssertionError(f"Expected two opening bids, got {len(rows)}")
            joint = next((row for row in rows if str(row[1]) == "Liên danh"), None)
            if not joint or int(joint[2]) != 3 or int(joint[3] or 0) != 1:
                raise AssertionError(f"Invalid joint-venture persistence: {joint}")
            return {"openingRows": len(rows), "jointMembers": int(joint[2]), "leaders": int(joint[3])}


def _verify_evaluation(data: dict) -> dict:
    organization_id = str(data["organizationId"])
    package_id = str(data["package"]["id"])
    raw_expected = data["expectedLowPriceDecision"]
    expected = None if raw_expected is None else bool(raw_expected)
    with psycopg.connect(_database_url()) as connection:
        with connection.cursor() as cursor:
            row = cursor.execute(
                """SELECT result.chap_thuan_gia_de_nghi_trung_thau_duoi_50,
                          COALESCE(result.ly_do_loai, ''),
                          result.gia_de_nghi_trung_thau,
                          COALESCE(result.danh_gia_tai_chinh, '')
                     FROM thong_tin_mo_thau AS opening
                     JOIN ket_qua_danh_gia_nha_thau AS result
                       ON result.organization_id = opening.organization_id
                      AND result.thong_tin_mo_thau_id = opening.id
                    WHERE opening.organization_id = %s
                      AND opening.goi_thau_id = %s
                      AND opening.loai_nha_thau = 'Liên danh'
                    LIMIT 1""",
                (organization_id, package_id),
            ).fetchone()
            if not row:
                raise AssertionError("Joint-venture evaluation was not persisted")
            actual = bool(row[0]) if row[0] is not None else None
            reason = str(row[1] or "")
            if actual is not expected:
                raise AssertionError(f"Expected low-price decision {expected}, got {actual}")
            if expected is None and reason:
                raise AssertionError(f"Non-warning price kept a rejection reason: {reason}")
            if expected is True and reason:
                raise AssertionError(f"Accepted low-price decision kept rejection reason: {reason}")
            if expected is False and "nhỏ hơn 50%" not in reason:
                raise AssertionError(f"Rejected low-price decision has wrong reason: {reason}")
            financial_result = str(row[3] or "")
            if expected in {None, True} and "Xếp hạng 1" not in financial_result:
                raise AssertionError(f"Accepted joint venture has wrong ranking: {financial_result}")
            if expected is False and financial_result not in {"--", "Không xếp hạng"}:
                raise AssertionError(f"Rejected joint venture was not removed from ranking: {financial_result}")
            independent_row = cursor.execute(
                """SELECT COALESCE(result.danh_gia_tai_chinh, '')
                     FROM thong_tin_mo_thau AS opening
                     JOIN ket_qua_danh_gia_nha_thau AS result
                       ON result.organization_id = opening.organization_id
                      AND result.thong_tin_mo_thau_id = opening.id
                    WHERE opening.organization_id = %s
                      AND opening.goi_thau_id = %s
                      AND opening.loai_nha_thau = 'Độc lập'
                    LIMIT 1""",
                (organization_id, package_id),
            ).fetchone()
            independent_ranking = str(independent_row[0] or "") if independent_row else ""
            expected_independent = "Xếp hạng 2" if expected in {None, True} else "Xếp hạng 1"
            if expected_independent not in independent_ranking:
                raise AssertionError(
                    f"Independent bidder was not reranked to {expected_independent}: {independent_ranking}"
                )
            return {
                "decision": actual,
                "reasonCleared": not reason,
                "proposedPrice": int(row[2] or 0),
                "financialResult": financial_result,
                "independentRanking": independent_ranking,
            }


def _verify_lot_outcomes(data: dict) -> dict:
    organization_id = str(data["organizationId"])
    package = data["lotPackage"]
    contractors = data["contractors"]
    with psycopg.connect(_database_url()) as connection:
        with connection.cursor() as cursor:
            rows = cursor.execute(
                """SELECT opening.ma_phan_lo, opening.loai_nha_thau,
                          opening.nha_thau_id, count(member.id),
                          result.chap_thuan_gia_de_nghi_trung_thau_duoi_50,
                          COALESCE(result.danh_gia_tai_chinh, ''),
                          COALESCE(result.ly_do_loai, '')
                     FROM thong_tin_mo_thau AS opening
                     LEFT JOIN thong_tin_mo_thau_lien_danh_thanh_vien AS member
                       ON member.organization_id = opening.organization_id
                      AND member.thong_tin_mo_thau_id = opening.id
                     LEFT JOIN ket_qua_danh_gia_nha_thau AS result
                       ON result.organization_id = opening.organization_id
                      AND result.thong_tin_mo_thau_id = opening.id
                    WHERE opening.organization_id = %s
                      AND opening.goi_thau_id = %s
                    GROUP BY opening.id, opening.ma_phan_lo, opening.loai_nha_thau,
                             opening.nha_thau_id,
                             result.chap_thuan_gia_de_nghi_trung_thau_duoi_50,
                             result.danh_gia_tai_chinh, result.ly_do_loai
                    ORDER BY opening.ma_phan_lo, opening.loai_nha_thau""",
                (organization_id, package["id"]),
            ).fetchall()
            if len(rows) != 3:
                raise AssertionError(f"Expected three multi-lot bids, got {rows}")
            by_scope = {(str(row[0]), str(row[1])): row for row in rows}
            lot1_joint = by_scope.get((package["lots"][0]["code"], "Liên danh"))
            lot2_joint = by_scope.get((package["lots"][1]["code"], "Liên danh"))
            lot2_independent = by_scope.get((package["lots"][1]["code"], "Độc lập"))
            if not lot1_joint or not lot2_joint or not lot2_independent:
                raise AssertionError(f"Missing expected lot participation: {rows}")
            if int(lot1_joint[3]) != 3 or int(lot2_joint[3]) != 3:
                raise AssertionError(f"Joint-venture members were lost by lot: {rows}")
            if "Xếp hạng 1" not in str(lot1_joint[5]):
                raise AssertionError(f"Joint venture did not win lot 1: {lot1_joint}")
            if lot2_joint[4] not in (False, 0) or str(lot2_joint[5]) not in {"--", "Không xếp hạng"}:
                raise AssertionError(f"Joint venture was not rejected only in lot 2: {lot2_joint}")
            if "nhỏ hơn 50%" not in str(lot2_joint[6]):
                raise AssertionError(f"Lot 2 rejection reason is wrong: {lot2_joint}")
            if "Xếp hạng 1" not in str(lot2_independent[5]):
                raise AssertionError(f"Independent bidder did not rerank in lot 2: {lot2_independent}")

            lot_rows = cursor.execute(
                """SELECT ma_phan_lo, nha_thau_trung_thau_id, gia_trung_thau
                     FROM goi_thau_phan_lo
                    WHERE organization_id = %s AND goi_thau_id = %s
                    ORDER BY sort_order""",
                (organization_id, package["id"]),
            ).fetchall()
            winners = {str(row[0]): str(row[1] or "") for row in lot_rows}
            if winners.get(package["lots"][0]["code"]) != contractors[0]["id"]:
                raise AssertionError(f"Lot 1 winner is wrong: {lot_rows}")
            if winners.get(package["lots"][1]["code"]) != contractors[3]["id"]:
                raise AssertionError(f"Lot 2 winner is wrong: {lot_rows}")
            return {
                "openingRows": len(rows),
                "jointLotCount": 2,
                "membersPerJointLot": [int(lot1_joint[3]), int(lot2_joint[3])],
                "lot1Result": str(lot1_joint[5]),
                "lot2JointResult": str(lot2_joint[5]),
                "lot2IndependentResult": str(lot2_independent[5]),
                "winners": winners,
            }


def _verify_two_envelope(data: dict) -> dict:
    organization_id = str(data["organizationId"])
    package = data["twoEnvelopePackage"]
    contractors = data["contractors"]
    with psycopg.connect(_database_url()) as connection:
        with connection.cursor() as cursor:
            rows = cursor.execute(
                """SELECT opening.loai_nha_thau, opening.nha_thau_id,
                          count(member.id),
                          sum(CASE WHEN member.vai_tro = 'Đứng đầu liên danh' THEN 1 ELSE 0 END),
                          COALESCE(result.danh_gia_ky_thuat, ''),
                          COALESCE(result.danh_gia_tai_chinh, ''),
                          result.chap_thuan_gia_de_nghi_trung_thau_duoi_50,
                          result.gia_de_nghi_trung_thau
                     FROM thong_tin_mo_thau AS opening
                     LEFT JOIN thong_tin_mo_thau_lien_danh_thanh_vien AS member
                       ON member.organization_id = opening.organization_id
                      AND member.thong_tin_mo_thau_id = opening.id
                     LEFT JOIN ket_qua_danh_gia_nha_thau AS result
                       ON result.organization_id = opening.organization_id
                      AND result.thong_tin_mo_thau_id = opening.id
                    WHERE opening.organization_id = %s
                      AND opening.goi_thau_id = %s
                    GROUP BY opening.id, opening.loai_nha_thau, opening.nha_thau_id,
                             result.danh_gia_ky_thuat, result.danh_gia_tai_chinh,
                             result.chap_thuan_gia_de_nghi_trung_thau_duoi_50,
                             result.gia_de_nghi_trung_thau
                    ORDER BY opening.loai_nha_thau""",
                (organization_id, package["id"]),
            ).fetchall()
            if len(rows) != 2:
                raise AssertionError(f"Expected two 1G2T bids, got {rows}")
            joint = next((row for row in rows if str(row[0]) == "Liên danh"), None)
            independent = next((row for row in rows if str(row[0]) == "Độc lập"), None)
            if not joint or not independent:
                raise AssertionError(f"Missing 1G2T bidder type: {rows}")
            if str(joint[1]) != contractors[0]["id"] or int(joint[2]) != 3 or int(joint[3] or 0) != 1:
                raise AssertionError(f"Invalid 1G2T joint venture: {joint}")
            if "Đạt" not in str(joint[4]) or "Xếp hạng 1" not in str(joint[5]):
                raise AssertionError(f"Qualified joint venture did not reach financial ranking: {joint}")
            if joint[6] not in (True, 1) or int(joint[7] or 0) != 400_000:
                raise AssertionError(f"Wrong 1G2T low-price outcome: {joint}")
            if "Không đạt" not in str(independent[4]) or str(independent[5]) not in {"", "--", "Không xếp hạng"}:
                raise AssertionError(f"Technically rejected bidder leaked into financial evaluation: {independent}")
            winner = cursor.execute(
                """SELECT nha_thau_trung_thau_id, gia_trung_thau, trang_thai
                     FROM goi_thau
                    WHERE organization_id = %s AND id = %s""",
                (organization_id, package["id"]),
            ).fetchone()
            if not winner or str(winner[0] or "") != contractors[0]["id"] or int(winner[1] or 0) != 400_000:
                raise AssertionError(f"Wrong 1G2T package winner: {winner}")
            return {
                "openingRows": len(rows),
                "jointMembers": int(joint[2]),
                "leaders": int(joint[3] or 0),
                "lowPriceDecision": bool(joint[6]),
                "proposedPrice": int(joint[7] or 0),
                "rejectedBidFinancialResult": str(independent[5]),
                "winnerId": str(winner[0]),
                "awardPrice": int(winner[1]),
                "packageStatus": str(winner[2]),
            }


def _verify_contract(data: dict) -> dict:
    organization_id = str(data["organizationId"])
    package = data["package"]
    contractors = data["contractors"]
    with psycopg.connect(_database_url()) as connection:
        with connection.cursor() as cursor:
            contract_rows = cursor.execute(
                """SELECT id, so_hop_dong, nha_thau_id, gia_tri, trang_thai_hop_dong
                     FROM hop_dong
                    WHERE organization_id = %s
                    ORDER BY so_hop_dong""",
                (organization_id,),
            ).fetchall()
            expected_contract_number = f"{data['runId']}/HD-JV"
            contract = next(
                (
                    item for item in contract_rows
                    if str(item[1]).strip().casefold() == expected_contract_number.casefold()
                ),
                None,
            )
            if not contract:
                raise AssertionError(
                    f"Joint-venture contract was not persisted; organization rows={contract_rows}"
                )
            links = cursor.execute(
                """SELECT goi_thau_id FROM hop_dong_goi_thau
                    WHERE organization_id = %s AND hop_dong_id = %s
                    ORDER BY goi_thau_id""",
                (organization_id, contract[0]),
            ).fetchall()
            opening = cursor.execute(
                """SELECT id, nha_thau_id
                     FROM thong_tin_mo_thau
                    WHERE organization_id = %s AND goi_thau_id = %s
                      AND loai_nha_thau = 'Liên danh'""",
                (organization_id, package["id"]),
            ).fetchone()
            if not opening:
                raise AssertionError("Winning joint-venture opening is missing")
            member_counts = cursor.execute(
                """SELECT count(*),
                          sum(CASE WHEN vai_tro = 'Đứng đầu liên danh' THEN 1 ELSE 0 END)
                     FROM thong_tin_mo_thau_lien_danh_thanh_vien
                    WHERE organization_id = %s AND thong_tin_mo_thau_id = %s""",
                (organization_id, opening[0]),
            ).fetchone()
    linked_package_ids = [str(item[0]) for item in links]
    if str(contract[2]) != contractors[0]["id"] or str(contract[2]) != str(opening[1]) or int(contract[3]) != 400_000:
        raise AssertionError(
            f"Joint-venture contract binding is wrong: contract={contract}, opening={opening}, links={linked_package_ids}"
        )
    if package["id"] not in linked_package_ids:
        raise AssertionError(f"Contract lost its winning package link: {linked_package_ids}")
    if int(member_counts[0]) != 3 or int(member_counts[1] or 0) != 1:
        raise AssertionError(f"Contract cannot resolve the full joint venture: {member_counts}")
    return {
        "contractNumber": str(contract[1]),
        "leadContractorId": str(contract[2]),
        "value": int(contract[3]),
        "status": str(contract[4]),
        "packageId": package["id"],
        "jointMembers": int(member_counts[0]),
        "leaders": int(member_counts[1] or 0),
    }


def _inspect_export(data: dict) -> dict:
    export_path = str(data["exportPath"])
    expected_names = [str(value) for value in data.get("expectedNames", [])]
    source = open(export_path, "rb")
    workbook = load_workbook(source, read_only=True, data_only=False)
    try:
        values = []
        fonts = set()
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        values.append(str(cell.value))
                    if cell.font and cell.font.name:
                        fonts.add(str(cell.font.name))
        text = "\n".join(values)
        for expected in expected_names:
            if expected not in text:
                raise AssertionError(f"Export is missing expected value: {expected}")
        if "Chấp thuận" not in text:
            raise AssertionError("Export is missing the accepted low-price decision")
        if "\ufffd" in text or "Ã" in text:
            raise AssertionError("Export contains broken Vietnamese text")
        return {
            "sheets": workbook.sheetnames,
            "fonts": sorted(fonts),
            "values": len(values),
        }
    finally:
        workbook.close()
        source.close()


def _cleanup(data: dict) -> dict:
    organization_id = str(data["organizationId"])
    with psycopg.connect(_database_url()) as connection:
        with connection.cursor() as cursor:
            for table_name in (
                "ho_so_nghiep_vu_lcnt_phan_lo",
                "ho_so_nghiep_vu_lcnt",
                "nhom_phu_thuoc_phan_lo_thanh_vien",
                "nhom_phu_thuoc_phan_lo",
                "dot_xu_ly_phan_lo_chi_tiet",
                "dot_xu_ly_phan_lo",
                "chi_tiet_danh_gia_nha_thau",
                "bao_cao_danh_gia_nha_thau",
                "tieu_chi_danh_gia",
                "ket_qua_danh_gia_nha_thau",
                "vong_danh_gia",
                "nha_thau_tham_du_mo_thau",
                "thong_tin_mo_thau_lien_danh_thanh_vien",
                "thong_tin_mo_thau",
                "goi_thau_hang_hoa",
                "goi_thau_phan_lo",
                "hop_dong_goi_thau",
                "phan_cong_nhan_su",
                "hop_dong",
                "danh_muc_trang_thai_hop_dong",
                "goi_thau",
                "nha_thau_lien_danh_thanh_vien",
                "nha_thau",
                "ke_hoach_lcnt",
                "chu_dau_tu",
                "record_edit_ownership",
                "sync_mutations",
                "deleted_records",
            ):
                cursor.execute(f"DELETE FROM {table_name} WHERE organization_id = %s", (organization_id,))
            cursor.execute("DELETE FROM thanh_vien_to_chuc WHERE organization_id = %s", (organization_id,))
            cursor.execute("DELETE FROM sync_metadata WHERE organization_id = %s", (organization_id,))
            cursor.execute("DELETE FROM organization_subscriptions WHERE organization_id = %s", (organization_id,))
            cursor.execute("DELETE FROM to_chuc WHERE id = %s", (organization_id,))
            if data.get("account"):
                cursor.execute("DELETE FROM tai_khoan WHERE id = %s", (str(data["account"]["id"]),))
    if data.get("account"):
        shutil.rmtree(WORD_TEMPLATE_DIR / str(data["account"]["id"]), ignore_errors=True)
    safe_organization_id = organization_id.replace("..", "").replace("/", "").replace("\\", "").strip()
    shutil.rmtree(
        WORD_TEMPLATE_DIR / "organizations" / safe_organization_id,
        ignore_errors=True,
    )
    word_template_path = str(data.get("wordTemplatePath") or "").strip()
    if word_template_path:
        Path(word_template_path).unlink(missing_ok=True)
    return {"removed": organization_id, "accountRemoved": bool(data.get("account"))}


def main() -> None:
    if len(sys.argv) != 2 or sys.argv[1] not in {"setup", "create_word_template", "verify", "verify_evaluation", "verify_lot_outcomes", "verify_two_envelope", "verify_contract", "inspect_export", "inspect_docx", "cleanup"}:
        raise SystemExit("Usage: joint_venture_e2e_fixture.py setup|create_word_template|verify|verify_evaluation|verify_lot_outcomes|verify_two_envelope|verify_contract|inspect_export|inspect_docx|cleanup")
    data = _payload()
    action = sys.argv[1]
    result = {
        "setup": _setup,
        "create_word_template": _create_word_template,
        "verify": _verify,
        "verify_evaluation": _verify_evaluation,
        "verify_lot_outcomes": _verify_lot_outcomes,
        "verify_two_envelope": _verify_two_envelope,
        "verify_contract": _verify_contract,
        "inspect_export": _inspect_export,
        "inspect_docx": _inspect_docx,
        "cleanup": _cleanup,
    }[action](data)
    sys.stdout.buffer.write(json.dumps(result, ensure_ascii=False).encode("utf-8"))


if __name__ == "__main__":
    main()
