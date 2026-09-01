from io import BytesIO
import hashlib
from types import SimpleNamespace
from zipfile import ZipFile

from docx import Document
from PIL import Image
import pytest

from backend.documents import custom_exporter
from backend.documents.document_ipc import (
    DocumentIpcError,
    read_job_manifest,
    write_job_manifest,
)
from backend.documents.docx_context_policy import (
    MANIFEST_VERSION,
    project_docx_context,
    validate_docx_context_manifest,
)
from backend.shared.media_helper import managed_image_tenant_segment


def _tenant_image(image_root, organization_id, *, filename="signature.png"):
    tenant_segment = managed_image_tenant_segment(organization_id)
    image_path = image_root / "chuyen_gia" / tenant_segment / filename
    image_path.parent.mkdir(parents=True, exist_ok=True)
    Image.new("RGB", (8, 8), color=(20, 80, 140)).save(image_path, format="PNG")
    return f"images/chuyen_gia/{tenant_segment}/{filename}", image_path


def _signature_capabilities():
    return SimpleNamespace(financial=False, identity=False, signature=True)


def test_docx_projection_preserves_only_current_tenant_media(tmp_path):
    own_path, _ = _tenant_image(tmp_path / "images", "org-a")
    foreign_path, _ = _tenant_image(
        tmp_path / "images",
        "org-b",
        filename="foreign.png",
    )
    context = {
        "to_chuyen_gia": [
            {"id": "expert-a", "anh_chu_ky": own_path},
            {"id": "expert-b", "anh_chu_ky": foreign_path},
        ]
    }

    projected = project_docx_context(
        "evaluation",
        context,
        _signature_capabilities(),
        organization_id="org-a",
    )

    assert projected["to_chuyen_gia"][0]["anh_chu_ky"] == own_path
    assert projected["to_chuyen_gia"][1]["anh_chu_ky"] == ""


def test_document_ipc_copies_current_tenant_media_with_namespace(tmp_path):
    image_root = tmp_path / "source-images"
    managed_path, source_path = _tenant_image(image_root, "org-a")
    job_dir = tmp_path / "job"
    job_dir.mkdir()

    write_job_manifest(
        job_dir / "input.json",
        "render_docx",
        {
            "context": {"signature_image": managed_path},
            "context_manifest": {"media_organization_id": "org-a"},
        },
        image_root=image_root,
    )

    copied = job_dir / "assets" / "images" / source_path.relative_to(image_root)
    assert copied.read_bytes() == source_path.read_bytes()


def test_replayed_worker_manifest_reuses_staged_images_without_second_copy(tmp_path):
    image_root = tmp_path / "source-images"
    managed_path, source_path = _tenant_image(image_root, "org-a")
    template_path = tmp_path / "template.docx"
    Document().save(template_path)
    job_dir = tmp_path / "job"
    job_dir.mkdir()
    payload = {
        "template_path": str(template_path),
        "context": {"anh_chu_ky": managed_path},
        "context_manifest": {"media_organization_id": "org-a"},
    }
    write_job_manifest(
        job_dir / "input.json", "render_docx", payload, image_root=image_root,
    )
    copied = job_dir / "assets" / "images" / source_path.relative_to(image_root)
    before = copied.stat().st_mtime_ns
    _operation, materialized = read_job_manifest(job_dir / "input.json", job_dir)
    write_job_manifest(
        job_dir / "retry-input.json",
        "render_docx",
        materialized,
        image_root=job_dir / "assets" / "images",
        copy_images=False,
        sidecar_prefix="input-retry",
    )

    assert copied.read_bytes() == source_path.read_bytes()
    assert copied.stat().st_mtime_ns == before
    assert list((job_dir / "assets" / "images").rglob(source_path.name)) == [copied]


def test_document_ipc_rejects_foreign_tenant_media(tmp_path):
    image_root = tmp_path / "source-images"
    foreign_path, _ = _tenant_image(image_root, "org-b")
    job_dir = tmp_path / "job"
    job_dir.mkdir()

    with pytest.raises(DocumentIpcError, match="tổ chức"):
        write_job_manifest(
            job_dir / "input.json",
            "render_docx",
            {
                "context": {"signature_image": foreign_path},
                "context_manifest": {"media_organization_id": "org-a"},
            },
            image_root=image_root,
        )


def test_docx_renderer_embeds_tenant_scoped_media(tmp_path, monkeypatch):
    image_root = tmp_path / "worker-images"
    managed_path, _ = _tenant_image(image_root, "org-a")
    template_path = tmp_path / "template.docx"
    template = Document()
    template.add_paragraph("{{ signature_image }}")
    template.save(template_path)
    monkeypatch.setattr(custom_exporter, "IMAGE_DIR", str(image_root))
    context = {"signature_image": managed_path}
    manifest = {
        "version": MANIFEST_VERSION,
        "document_type": "evaluation",
        "root_keys": ["signature_image"],
        "custom_root_keys": ["signature_image"],
        "image_fields": {"signature_image": "chuyen_gia"},
        "media_organization_id": "org-a",
    }

    output = custom_exporter.generate_report_from_custom_template(
        str(template_path),
        context,
        manifest,
    )

    with ZipFile(BytesIO(output.getvalue())) as archive:
        media_files = [name for name in archive.namelist() if name.startswith("word/media/")]
        assert media_files


@pytest.mark.parametrize(
    ("image_format", "extension"),
    (("PNG", "png"), ("JPEG", "jpg"), ("WEBP", "webp")),
)
def test_legal_word_media_is_embedded_byte_for_byte(
    tmp_path, monkeypatch, image_format, extension,
):
    image_root = tmp_path / "worker-images"
    tenant_segment = managed_image_tenant_segment("org-a")
    source_path = (
        image_root / "chuyen_gia" / tenant_segment / f"signature.{extension}"
    )
    source_path.parent.mkdir(parents=True)
    Image.new("RGB", (19, 11), color=(45, 90, 135)).save(
        source_path, format=image_format,
    )
    source = source_path.read_bytes()
    managed_path = f"images/chuyen_gia/{tenant_segment}/{source_path.name}"
    template_path = tmp_path / "template.docx"
    template = Document()
    template.add_paragraph("{{ anh_chu_ky }}")
    template.save(template_path)
    monkeypatch.setattr(custom_exporter, "IMAGE_DIR", str(image_root))
    manifest = {
        "version": MANIFEST_VERSION,
        "document_type": "evaluation",
        "root_keys": ["anh_chu_ky"],
        "custom_root_keys": ["anh_chu_ky"],
        "image_fields": {"anh_chu_ky": "chuyen_gia"},
        "media_organization_id": "org-a",
    }

    output = custom_exporter.generate_report_from_custom_template(
        str(template_path), {"anh_chu_ky": managed_path}, manifest,
    ).getvalue()

    with ZipFile(BytesIO(output)) as archive:
        embedded = [
            archive.read(name)
            for name in archive.namelist()
            if name.startswith("word/media/")
        ]
    assert len(embedded) == 1
    assert embedded[0] == source
    assert hashlib.sha256(embedded[0]).digest() == hashlib.sha256(source).digest()


def test_document_worker_rejects_malformed_managed_media_path():
    context = {"signature_image": "images/chuyen_gia/../escape.png"}
    manifest = {
        "version": MANIFEST_VERSION,
        "document_type": "evaluation",
        "root_keys": ["signature_image"],
        "custom_root_keys": ["signature_image"],
        "image_fields": {"signature_image": "chuyen_gia"},
        "media_organization_id": "org-a",
    }

    with pytest.raises(ValueError, match="Ảnh Word"):
        validate_docx_context_manifest(context, manifest)


def test_docx_renderer_fails_clearly_when_legal_media_is_missing(
    tmp_path, monkeypatch,
):
    image_root = tmp_path / "worker-images"
    template_path = tmp_path / "template.docx"
    template = Document()
    template.add_paragraph("{{ anh_chu_ky }}")
    template.save(template_path)
    monkeypatch.setattr(custom_exporter, "IMAGE_DIR", str(image_root))
    tenant_segment = managed_image_tenant_segment("org-a")
    context = {
        "anh_chu_ky": f"images/chuyen_gia/{tenant_segment}/missing.png"
    }
    manifest = {
        "version": MANIFEST_VERSION,
        "document_type": "evaluation",
        "root_keys": ["anh_chu_ky"],
        "custom_root_keys": ["anh_chu_ky"],
        "image_fields": {"anh_chu_ky": "chuyen_gia"},
        "media_organization_id": "org-a",
    }

    with pytest.raises(
        custom_exporter.TemplateRenderError,
        match="ảnh pháp lý",
    ):
        custom_exporter.generate_report_from_custom_template(
            str(template_path), context, manifest,
        )
