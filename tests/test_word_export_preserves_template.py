from io import BytesIO
from zipfile import ZIP_STORED, ZipFile

from docx import Document
from docx.shared import Pt, RGBColor

import backend.documents.document_worker as document_worker
from backend.documents.document_worker import run_document_job
from backend.documents.docx_context_policy import MANIFEST_VERSION


def _empty_manifest() -> dict:
    return {
        "version": MANIFEST_VERSION,
        "document_type": "plan",
        "root_keys": [],
        "custom_root_keys": [],
        "datetime_root_keys": [],
        "date_root_keys": [],
        "money_root_keys": [],
        "image_fields": {},
        "media_organization_id": "org-a",
    }


def _template() -> bytes:
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("GIỮ NGUYÊN ĐỊNH DẠNG TEMPLATE")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x22, 0x22)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _first_run_format(content: bytes) -> tuple[str | None, int | None, str | None]:
    run = Document(BytesIO(content)).paragraphs[0].runs[0]
    color = str(run.font.color.rgb) if run.font.color.rgb is not None else None
    size = int(run.font.size.pt) if run.font.size is not None else None
    return run.font.name, size, color


def test_word_export_uses_original_template_format_even_if_legacy_mode_is_set(
    monkeypatch,
):
    monkeypatch.setenv("WORD_EXPORT_STANDARDIZATION_MODE", "apply_safe")
    source = _template()

    output = run_document_job(
        "render_docx",
        {
            "template_content": source,
            "context": {},
            "context_manifest": _empty_manifest(),
        },
        timeout_seconds=30,
    )

    assert _first_run_format(source) == ("Arial", 11, "662222")
    assert _first_run_format(output) == _first_run_format(source)


def test_batch_word_export_preserves_each_template_and_uses_one_worker(
    monkeypatch,
):
    monkeypatch.setenv("WORD_EXPORT_STANDARDIZATION_MODE", "apply_safe")
    original_start_worker = document_worker._start_worker_process
    process_count = 0

    def counting_start_worker(*args, **kwargs):
        nonlocal process_count
        process_count += 1
        return original_start_worker(*args, **kwargs)

    monkeypatch.setattr(
        document_worker,
        "_start_worker_process",
        counting_start_worker,
    )
    first = _template()
    second_document = Document(BytesIO(_template()))
    second_run = second_document.paragraphs[0].runs[0]
    second_run.font.name = "Calibri"
    second_run.font.size = Pt(13)
    second_run.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
    second_stream = BytesIO()
    second_document.save(second_stream)
    second = second_stream.getvalue()

    result = run_document_job(
        "render_docx_batch",
        {
            "templates": [
                {"template_content": first, "filename": "A.docx"},
                {"template_content": second, "filename": "B.docx"},
            ],
            "context": {},
            "context_manifest": _empty_manifest(),
        },
        timeout_seconds=30,
    )

    assert process_count == 1
    with ZipFile(BytesIO(result)) as archive:
        assert [item.filename for item in archive.infolist()] == ["A.docx", "B.docx"]
        assert {item.compress_type for item in archive.infolist()} == {ZIP_STORED}
        assert _first_run_format(archive.read("A.docx")) == _first_run_format(first)
        assert _first_run_format(archive.read("B.docx")) == _first_run_format(second)
