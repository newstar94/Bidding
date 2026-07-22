from __future__ import annotations

import json
from datetime import timedelta

from docx import Document

from backend.documents import custom_exporter
from backend.documents import timeline_context_service as context_service
from backend.documents import timeline_document_service as document_service
from backend.shared.date_utils import vietnam_today


class _Cursor:
    def __init__(self, responses):
        self.responses = responses
        self.rows = []

    def execute(self, sql, params=()):
        normalized = " ".join(sql.split())
        for marker, rows in self.responses.items():
            if marker in normalized:
                self.rows = rows
                break
        else:
            self.rows = []
        return self

    def fetchone(self):
        return self.rows[0] if self.rows else None

    def fetchall(self):
        return list(self.rows)


class _Connection:
    def __init__(self, cursor):
        self._cursor = cursor
        self.closed = False

    def cursor(self):
        return self._cursor

    def close(self):
        self.closed = True


class _Database:
    def __init__(self, connection):
        self.connection = connection

    def get_connection(self):
        return self.connection


def test_timeline_template_renders_sections_and_replaces_every_token(tmp_path):
    template = tmp_path / "timeline.docx"
    assert document_service.create_timeline_template(template) == template

    context = {
        "to_chuc": {"ten_to_chuc": "Công ty kiểm thử"},
        "goi_thau": {"ma_goi_thau": "GT-01", "ten_goi_thau": "Gói mẫu"},
        "ke_hoach": {"ma_ke_hoach": "KH-01"},
        "planned_date_note": "Ngày đỏ là dự kiến",
        "generated_date": "19/07/2026",
        "timeline_sections": [
            {
                "code": "I",
                "title": "Chuẩn bị",
                "items": [
                    {
                        "display_code": "1.1",
                        "cong_viec": "Việc một",
                        "don_vi_ban_hanh": "Chủ đầu tư",
                        "so_van_ban": "01/QĐ",
                        "display_date": "18/07/2026",
                        "ghi_chu": "Đã xong",
                        "is_planned_date": False,
                    },
                    {
                        "ma_moc": "1.2",
                        "cong_viec": "Việc hai",
                        "don_vi_ban_hanh": "Tư vấn",
                        "so_van_ban": "",
                        "display_date": "20/07/2026",
                        "ghi_chu": "Dự kiến",
                        "is_planned_date": True,
                    },
                ],
            }
        ],
    }
    stream = document_service.render_timeline_document(template, context)
    result = Document(stream)
    all_text = "\n".join(
        [paragraph.text for paragraph in result.paragraphs]
        + [cell.text for table in result.tables for row in table.rows for cell in row.cells]
        + [paragraph.text for section in result.sections for paragraph in section.header.paragraphs]
    )
    assert "Công ty kiểm thử" in all_text
    assert "GT-01" in all_text
    assert "KH-01" in all_text
    assert "Chuẩn bị" in all_text
    assert "Việc một" in all_text
    assert "20/07/2026" in all_text
    assert len(result.tables[0].rows) == 4


def test_timeline_context_filters_sources_dates_and_contracts(monkeypatch):
    yesterday = (vietnam_today() - timedelta(days=1)).isoformat()
    responses = {
        "SELECT * FROM goi_thau WHERE organization_id": [{
            "id": "pkg-1",
            "ke_hoach_id": "plan-1",
            "hinh_thuc_lua_chon": "Đấu thầu rộng rãi",
            "phuong_thuc_lua_chon": "Một giai đoạn hai túi hồ sơ",
            "yeu_cau_tham_dinh_hsmt": "Có",
            "so_quyet_dinh": "12/QĐ",
            "ngay_quyet_dinh": "2026-07-18",
        }],
        "SELECT * FROM ke_hoach_lcnt": [{
            "id": "plan-1",
            "chu_dau_tu_id": "investor-1",
            "phe_duyet": "Kế hoạch",
            "so_to_trinh_ke_hoach": "11/TTr-KH",
            "ngay_trinh_ke_hoach": "2026-07-16",
            "quyet_dinh_phe_duyet": "22/QĐ",
            "ngay_phe_duyet": "2026-07-17",
        }],
        "SELECT * FROM chu_dau_tu": [{"id": "investor-1", "ten_chu_dau_tu": "CĐT"}],
        "SELECT * FROM to_chuc": [{"id": "org-1", "ten_to_chuc": "Tổ chức"}],
        "SELECT * FROM goi_thau_moc_tien_do": [{
            "ma_moc": "2.1",
            "id": "milestone-1",
            "source_mode": "MANUAL",
            "ngay_du_kien": yesterday,
            "trang_thai": "PENDING",
            "is_optional": False,
        }],
        "SELECT loai_vong": [
            {"loai_vong": "technical", "so_bao_cao": "BC-KT", "ngay_bao_cao": "2026-07-15", "extension_json": json.dumps({"soBctdKt": "TĐ-KT", "ngayBctdKt": "2026-07-16"})},
            {"loai_vong": "financial", "so_bao_cao": "BC-TC", "ngay_bao_cao": "2026-07-17", "extension_json": {}},
        ],
        "SELECT hd.*": [
            {"phan_loai": "Tư vấn lập", "so_qd_chi_dinh": "QĐ-TVL", "ngay_qd_chi_dinh": "2026-06-01", "so_hop_dong": "HĐ-TVL", "ngay_ky": "2026-06-02", "ngay_thanh_ly": "2026-06-30"},
            {"phan_loai": "Tư vấn thẩm", "so_qd_chi_dinh": "QĐ-TVT", "ngay_qd_chi_dinh": "2026-06-03", "so_hop_dong": "HĐ-TVT", "ngay_ky": "2026-06-04", "ngay_thanh_ly": ""},
            {"phan_loai": "Khác", "so_hop_dong": "ignored"},
        ],
    }
    connection = _Connection(_Cursor(responses))
    monkeypatch.setattr(context_service, "database", _Database(connection))

    result = context_service.build_timeline_context("pkg-1", "user-1", "org-1")
    assert connection.closed
    assert result["goi_thau"]["id"] == "pkg-1"
    assert result["to_chuc"]["ten_to_chuc"] == "Tổ chức"
    flat = [item for section in result["timeline_sections"] for item in section["items"]]
    by_code = {item["ma_moc"]: item for item in flat}
    assert "1.7" not in by_code and "1.8" not in by_code
    assert by_code["1.5"]["so_van_ban"] == "11/TTr-KH"
    assert by_code["1.5"]["display_date"] == "16/07/2026"
    assert by_code["1.6"]["so_van_ban"] == "22/QĐ"
    assert by_code["2.1"]["is_overdue"] is True
    assert by_code["2.5"]["so_van_ban"] == "QĐ-TVL"
    assert by_code["5.2"]["so_van_ban"] == "BC-KT"
    assert by_code["5.3"]["so_van_ban"] == "TĐ-KT"
    assert by_code["5.6"]["so_van_ban"] == "BC-TC"
    assert by_code["4.3"]["display_date"] == "18/07/2026"
    assert result["planned_date_note"]


def test_timeline_export_sources_include_all_plan_submission_numbers():
    items = [
        {
            "ma_moc": code,
            "source_mode": "AUTO",
            "source_key": "",
            "so_van_ban": "",
            "ngay_thuc_te": "",
            "trang_thai": "PENDING",
        }
        for code in ("1.3", "1.5", "1.7")
    ]
    plan = {
        "so_to_trinh_du_toan": "01/TTr-DT",
        "ngay_trinh_du_toan": "2026-07-14",
        "so_to_trinh_ke_hoach": "02/TTr-KH",
        "so_to_trinh_du_toan_ke_hoach": "03/TTr-DTKH",
        "ngay_trinh_ke_hoach": "2026-07-15",
    }

    by_code = {
        item["ma_moc"]: item
        for item in context_service._apply_sources(items, {}, plan)
    }

    assert by_code["1.3"]["so_van_ban"] == "01/TTr-DT"
    assert by_code["1.3"]["ngay_thuc_te"] == "2026-07-14"
    assert by_code["1.5"]["so_van_ban"] == "02/TTr-KH"
    assert by_code["1.5"]["ngay_thuc_te"] == "2026-07-15"
    assert by_code["1.7"]["so_van_ban"] == "03/TTr-DTKH"
    assert by_code["1.7"]["ngay_thuc_te"] == "2026-07-15"


def test_timeline_context_rejects_missing_package_and_covers_applicability(monkeypatch):
    connection = _Connection(_Cursor({}))
    monkeypatch.setattr(context_service, "database", _Database(connection))
    try:
        context_service.build_timeline_context("missing", "user", "org")
    except ValueError as exc:
        assert "Không tìm thấy" in str(exc)
    else:
        raise AssertionError("Missing package must be rejected")
    assert connection.closed

    competitive = {
        "hinh_thuc_lua_chon": "Chào hàng cạnh tranh",
        "phuong_thuc_lua_chon": "Một giai đoạn một túi hồ sơ",
        "yeu_cau_tham_dinh_hsmt": "Không",
    }
    plan = {"phe_duyet": "Dự toán và kế hoạch"}
    assert not context_service._timeline_item_is_applicable({"ma_moc": "1.3"}, competitive, plan)
    assert not context_service._timeline_item_is_applicable({"ma_moc": "3.1"}, competitive, plan)
    assert not context_service._timeline_item_is_applicable({"ma_moc": "4.2"}, competitive, plan)
    assert not context_service._timeline_item_is_applicable({"ma_moc": "5.5"}, competitive, plan)
    assert context_service._timeline_item_is_applicable({"ma_moc": "2.1"}, competitive, plan)
    assert context_service._date_only("19/07/2026") == "2026-07-19"
    assert context_service._date_only("2026/07/19") == "2026-07-19"
    assert context_service._date_only("invalid") == ""
    assert context_service._parse_metadata("[]") == {}
    assert context_service._parse_metadata("invalid") == {}


def test_custom_exporter_context_helpers_cover_dates_words_and_template_state(tmp_path, monkeypatch):
    assert custom_exporter.number_to_vietnamese_words(None) == ""
    assert custom_exporter.number_to_vietnamese_words("invalid") == ""
    assert custom_exporter.number_to_vietnamese_words(0) == "Không"
    assert "tỷ" in custom_exporter.number_to_vietnamese_words(1_005_004_021)
    assert "lăm" in custom_exporter.number_to_vietnamese_words(25)
    assert "mốt" in custom_exporter.number_to_vietnamese_words(21)
    assert "tư" in custom_exporter.number_to_vietnamese_words(24)

    data = {"{TongGiaTri}": 125, "Items": [{"NgayKy": "2026-07-19"}]}
    custom_exporter.enrich_context_with_lowercase_keys(data)
    custom_exporter.enrich_context_with_words(data)
    custom_exporter.format_context_dates(data)
    assert data["tonggiatri"] == 125
    assert data["bangchu_TongGiaTri"].endswith(" đồng")
    assert str(data["Items"][0]["NgayKy"]) == "ngày 19 tháng 7 năm 2026"
    assert data["Items"][0]["s_ngay_ky"] == "19/7/2026"

    assert custom_exporter.normalize_date_str("2026-01-02T03:04:05Z") == "02/01/2026 03:04"
    assert custom_exporter.normalize_date_str("02-03-2026 04:05:06") == "02/03/2026 04:05"
    assert custom_exporter.normalize_date_str(123) == 123
    assert custom_exporter.format_vietnamese_datetime("2026-01-02 03:04", "thoiGianMoThau") == "03:04 ngày 02/01/2026"
    assert custom_exporter.format_vietnamese_datetime("not-a-date") == "not-a-date"
    assert custom_exporter.SmartDate("19/07/2026") - "18/07/2026" == 1
    assert "20/07/2026" - custom_exporter.SmartDate("19/07/2026") == 1
    assert custom_exporter.SmartDate("invalid") - "also-invalid" == ""

    monkeypatch.setattr(custom_exporter, "TEMPLATE_DIR", str(tmp_path))
    (tmp_path / "mau_bao_cao_dau_thau.docx").write_bytes(b"docx")
    user_dir = custom_exporter.get_user_template_dir("../user/path")
    assert ".." not in user_dir
    (tmp_path / "userpath" / "custom.docx").write_bytes(b"docx")
    custom_exporter.set_active_template("custom.docx", "../user/path")
    assert custom_exporter.get_active_template("../user/path") == "custom.docx"
    templates = custom_exporter.list_templates("../user/path")
    assert any(item["filename"] == "custom.docx" and item["is_active"] for item in templates)
    assert custom_exporter.validate_template_syntax(b"anything")[0]

    (tmp_path / "userpath" / "config.json").write_text("not-json", encoding="utf-8")
    assert custom_exporter.get_active_template("../user/path") == "mau_bao_cao_dau_thau.docx"
