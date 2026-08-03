from io import BytesIO
from pathlib import Path

from docx import Document

from backend.documents import custom_exporter
from backend.documents.routes_docx import _resolve_template_path
from backend.documents.timeline_document_service import render_timeline_document
from backend.shared.paths import provision_system_word_templates


TIMELINE_TEMPLATE = "mau_timeline_goi_thau.docx"
LEGACY_DEFAULT_TEMPLATES = {
    "mau_bao_cao_dau_thau.docx",
    "mau_hop_dong_lcnt.docx",
}


def test_timeline_template_is_generated_and_resolvable_without_bundled_data(tmp_path, monkeypatch):
    source_dir = tmp_path / "empty-source"
    target_dir = tmp_path / "runtime-templates"
    result = provision_system_word_templates(
        source_dir=source_dir,
        target_dir=target_dir,
    )

    provisioned = target_dir / TIMELINE_TEMPLATE
    assert provisioned in result["generated"]
    assert provisioned.is_file()

    monkeypatch.setattr(custom_exporter, "TEMPLATE_DIR", str(target_dir))
    resolved, filename = _resolve_template_path(
        "organization",
        "test-organization",
        TIMELINE_TEMPLATE,
    )

    assert Path(resolved) == provisioned
    assert filename == TIMELINE_TEMPLATE
    document = Document(resolved)
    assert len(document.tables) == 1
    assert len(document.tables[0].columns) == 6


def test_generated_timeline_template_renders_effective_rows(tmp_path):
    template = tmp_path / TIMELINE_TEMPLATE
    provision_system_word_templates(
        source_dir=tmp_path / "empty-source",
        target_dir=tmp_path,
    )

    rendered = render_timeline_document(
        template,
        {
            "goi_thau": {"ma_goi_thau": "GT-01", "ten_goi_thau": "Gói thầu thử nghiệm"},
            "ke_hoach": {"ma_ke_hoach": "KH-01"},
            "to_chuc": {"ten_to_chuc": "Đơn vị thử nghiệm"},
            "planned_date_note": "Ngày dự kiến",
            "generated_date": "03/08/2026",
            "timeline_sections": [{
                "code": "I",
                "title": "Chuẩn bị",
                "items": [{
                    "display_code": "1.1",
                    "cong_viec": "Phê duyệt kế hoạch",
                    "don_vi_ban_hanh": "Chủ đầu tư",
                    "so_van_ban": "01/QĐ",
                    "display_date": "03/08/2026",
                    "ghi_chu": "",
                    "is_planned_date": False,
                }],
            }],
        },
    )

    document = Document(BytesIO(rendered.getvalue()))
    assert len(document.tables[0].rows) == 3
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "GT-01 - Gói thầu thử nghiệm" in text
    assert "[[PACKAGE_CODE]]" not in text


def test_fresh_install_does_not_provision_legacy_default_templates(tmp_path):
    source_dir = tmp_path / "bundled"
    source_dir.mkdir()
    for filename in LEGACY_DEFAULT_TEMPLATES:
        (source_dir / filename).write_bytes(b"legacy")

    target_dir = tmp_path / "runtime"
    provision_system_word_templates(source_dir=source_dir, target_dir=target_dir)

    assert all(not (target_dir / filename).exists() for filename in LEGACY_DEFAULT_TEMPLATES)
