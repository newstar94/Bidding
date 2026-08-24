from __future__ import annotations

from io import BytesIO
import hashlib
import json
from datetime import datetime, timezone
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
import pytest
from lxml import etree

from backend.documents.document_worker_entry import _run_operation
from backend.documents.document_worker import run_document_job
from backend.documents.template_catalog.service import (
    CatalogConflictError,
    WordTemplateCatalog,
)
from backend.documents.word_standardizer import (
    WordStandardizationError,
    process_docx,
)
from backend.documents.word_standardizer.engine import (
    MAX_REPORT_BYTES,
    _PPR_ORDER,
    _RPR_ORDER,
    _rules,
    _validate_rule_bundle,
)


PLACEHOLDERS = (
    "{{ document.number }}",
    "{{ document.subject }}",
    "{% for item in items %}",
    "{{ item.name }}",
    "{% endfor %}",
)


def _fixture_docx() -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x22, 0x33, 0x44)

    document.add_paragraph("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    document.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    document.add_paragraph("Số: {{ document.number }}/TB-BF")
    document.add_paragraph("THÔNG BÁO")
    document.add_paragraph("Về việc {{ document.subject }}")
    document.add_paragraph("Kính gửi: Đơn vị liên quan")
    body = document.add_paragraph("Nội dung nghiệp vụ phải được giữ nguyên.")
    body.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = (
        "{% for item in items %}{{ item.name }}{% endfor %}"
    )
    signing = document.add_paragraph("KT. GIÁM ĐỐC")
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), "7")
    bookmark_start.set(qn("w:name"), "signing_block")
    signing._p.insert(0, bookmark_start)
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "7")
    signing._p.append(bookmark_end)
    document.add_paragraph("Nơi nhận:")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _document_xml(content: bytes) -> str:
    with ZipFile(BytesIO(content)) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def _with_zip_entries(content: bytes, additions: dict[str, bytes]) -> bytes:
    source = BytesIO(content)
    output = BytesIO()
    with ZipFile(source) as current, ZipFile(output, "w", ZIP_DEFLATED) as updated:
        for info in current.infolist():
            updated.writestr(info, current.read(info.filename))
        for name, value in additions.items():
            updated.writestr(name, value)
    return output.getvalue()


def test_audit_recognizes_document_and_scopes_sector_fixes():
    content = _fixture_docx()

    sector = process_docx(
        content,
        profile="sector_template",
        mode="preview_fix",
    ).report
    reference = process_docx(
        content,
        profile="reference_only",
        mode="audit",
    ).report

    assert sector["documentType"]["value"] == "thong_bao"
    assert sector["documentType"]["confidence"] >= 0.9
    assert sector["placeholders"]["count"] == len(PLACEHOLDERS)
    assert sector["summary"]["safeFixes"] > 0
    assert sector["summary"]["previewOnly"] > 0
    assert sector["invariants"]["status"] == "PASS"
    assert sector["plannedChanges"]
    assert all(
        change["ruleId"] != "N30-BODY-ALIGN"
        for change in sector["plannedChanges"]
    )
    assert all(
        not change["ruleId"].startswith("N30-PAGE-")
        for change in sector["plannedChanges"]
    )
    assert reference["summary"]["safeFixes"] == 0
    assert reference["plannedChanges"] == []


def test_apply_fix_preserves_text_placeholders_and_structure():
    content = _fixture_docx()
    preview = process_docx(
        content,
        profile="n30_strict",
        mode="preview_fix",
    ).report

    applied = process_docx(
        content,
        profile="n30_strict",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    )

    assert applied.content is not None
    assert applied.content != content
    assert applied.report["changed"] is True
    assert applied.report["invariants"]["status"] == "PASS"
    assert applied.report["outputSha256"]
    before_xml = _document_xml(content)
    after_xml = _document_xml(applied.content)
    for token in PLACEHOLDERS:
        assert before_xml.count(token) == after_xml.count(token) == 1
    assert before_xml.count("<w:tbl>") == after_xml.count("<w:tbl>") == 1
    assert before_xml.count("w:bookmarkStart") == after_xml.count("w:bookmarkStart")
    assert "Nội dung nghiệp vụ phải được giữ nguyên." in after_xml

    post_audit = process_docx(
        applied.content,
        profile="n30_strict",
        mode="audit",
    ).report
    assert post_audit["summary"]["safeFixes"] == 0


def test_apply_records_a_missing_true_boolean_property_as_a_change():
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(20)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    run.font.name = "Times New Roman"
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        run._r.get_or_add_rPr().get_or_add_rFonts().set(
            qn(f"w:{attribute}"), "Times New Roman"
        )
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.italic = False
    for value in ("Độc lập - Tự do - Hạnh phúc", "THÔNG BÁO"):
        shell = document.add_paragraph()
        shell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shell_run = shell.add_run(value)
        shell_run.font.name = "Times New Roman"
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            shell_run._r.get_or_add_rPr().get_or_add_rFonts().set(
                qn(f"w:{attribute}"), "Times New Roman"
            )
        shell_run.font.size = Pt(13)
        shell_run.font.color.rgb = RGBColor(0, 0, 0)
        shell_run.bold = True
        shell_run.italic = False
    output = BytesIO()
    document.save(output)
    content = output.getvalue()

    preview = process_docx(
        content,
        profile="sector_template",
        mode="preview_fix",
    ).report

    assert preview["summary"]["safeFixes"] == 1
    assert [item["ruleId"] for item in preview["plannedChanges"]] == [
        "N30-NATIONAL-HEADER-BOLD"
    ]
    applied = process_docx(
        content,
        profile="sector_template",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    )
    assert applied.content != content
    assert applied.report["changed"] is True
    assert applied.report["postFixSummary"]["safeFixes"] == 0


def test_apply_black_removes_theme_color_attributes():
    document = Document()
    shell_runs = []
    for value in (
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "Độc lập - Tự do - Hạnh phúc",
        "THÔNG BÁO",
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True
        run.italic = False
        shell_runs.append(run)
    color = shell_runs[0]._r.get_or_add_rPr().find(qn("w:color"))
    color.set(qn("w:themeColor"), "accent1")
    color.set(qn("w:themeTint"), "80")
    output = BytesIO()
    document.save(output)
    content = output.getvalue()
    preview = process_docx(
        content,
        profile="sector_template",
        mode="preview_fix",
    ).report

    assert any(
        item["ruleId"] == "N30-SHELL-COLOR"
        for item in preview["plannedChanges"]
    )
    applied = process_docx(
        content,
        profile="sector_template",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    ).content
    root = etree.fromstring(_document_xml(applied).encode("utf-8"))
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    first_color = root.find(".//w:color", namespaces=namespace)
    assert first_color.get(qn("w:val")) == "000000"
    assert first_color.get(qn("w:themeColor")) is None
    assert first_color.get(qn("w:themeTint")) is None


def test_apply_font_removes_valid_complex_script_theme_attribute():
    document = Document()
    runs = []
    for value in (
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "Độc lập - Tự do - Hạnh phúc",
        "THÔNG BÁO",
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True
        run.italic = False
        fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
        fonts.set(qn("w:ascii"), "Arial")
        fonts.set(qn("w:hAnsi"), "Arial")
        fonts.set(qn("w:cs"), "Arial")
        fonts.set(qn("w:asciiTheme"), "minorHAnsi")
        fonts.set(qn("w:cstheme"), "majorBidi")
        runs.append(run)
    output = BytesIO()
    document.save(output)
    content = output.getvalue()
    preview = process_docx(
        content,
        profile="sector_template",
        mode="preview_fix",
    ).report
    applied = process_docx(
        content,
        profile="sector_template",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    ).content
    root = etree.fromstring(_document_xml(applied).encode("utf-8"))
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    first_fonts = root.find(".//w:rFonts", namespaces=namespace)

    assert first_fonts.get(qn("w:cs")) == "Times New Roman"
    assert first_fonts.get(qn("w:asciiTheme")) is None
    assert first_fonts.get(qn("w:cstheme")) is None


def test_mixed_script_fonts_are_detected_and_all_scripts_are_normalized():
    document = Document()
    for value in (
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "Độc lập - Tự do - Hạnh phúc",
        "THÔNG BÁO",
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True
        run.italic = False
        fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
        fonts.set(qn("w:ascii"), "Times New Roman")
        fonts.set(qn("w:hAnsi"), "Arial")
        fonts.set(qn("w:eastAsia"), "Times New Roman")
        fonts.set(qn("w:cs"), "Times New Roman")
    output = BytesIO()
    document.save(output)
    content = output.getvalue()

    preview = process_docx(
        content,
        profile="sector_template",
        mode="preview_fix",
    ).report

    assert any(
        item["ruleId"] == "N30-SHELL-FONT"
        for item in preview["plannedChanges"]
    )
    applied = process_docx(
        content,
        profile="sector_template",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    ).content
    root = etree.fromstring(_document_xml(applied).encode("utf-8"))
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    first_fonts = root.find(".//w:rFonts", namespaces=namespace)
    assert {
        first_fonts.get(qn(f"w:{attribute}"))
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs")
    } == {"Times New Roman"}


def test_partial_direct_font_does_not_hide_inherited_hansi_font():
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    for value in (
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "Độc lập - Tự do - Hạnh phúc",
        "THÔNG BÁO",
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True
        run.italic = False
        fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
        for attribute in tuple(fonts.attrib):
            del fonts.attrib[attribute]
        fonts.set(qn("w:ascii"), "Times New Roman")
    output = BytesIO()
    document.save(output)
    content = output.getvalue()

    preview = process_docx(
        content,
        profile="sector_template",
        mode="preview_fix",
    ).report

    assert any(
        item["ruleId"] == "N30-SHELL-FONT"
        for item in preview["plannedChanges"]
    )
    applied = process_docx(
        content,
        profile="sector_template",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    ).content
    root = etree.fromstring(_document_xml(applied).encode("utf-8"))
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    first_fonts = root.find(".//w:rFonts", namespaces=namespace)
    assert {
        first_fonts.get(qn(f"w:{attribute}"))
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs")
    } == {"Times New Roman"}


def test_apply_rejects_stale_analysis_and_reference_only_profile():
    content = _fixture_docx()

    with pytest.raises(WordStandardizationError, match="stale"):
        process_docx(
            content,
            profile="n30_strict",
            mode="apply_fix",
            expected_analysis_hash="0" * 64,
        )
    with pytest.raises(WordStandardizationError, match="reference_only"):
        process_docx(content, profile="reference_only", mode="apply_fix")


def test_apply_requires_an_accepted_analysis_hash_at_engine_and_worker_seams():
    content = _fixture_docx()

    with pytest.raises(WordStandardizationError, match="requires"):
        process_docx(content, profile="sector_template", mode="apply_fix")
    with pytest.raises(WordStandardizationError, match="requires"):
        _run_operation("standardize_docx", {
            "content": content,
            "profile": "sector_template",
            "mode": "apply_fix",
        })


def test_signed_opc_package_is_manual_review_only_and_cannot_apply():
    content = _with_zip_entries(
        _fixture_docx(),
        {"_xmlsignatures/sig1.xml": b"<Signature xmlns='urn:test'/>"},
    )

    preview = process_docx(
        content,
        profile="sector_template",
        mode="preview_fix",
    ).report

    assert preview["packageSignature"]["detected"] is True
    assert preview["summary"]["safeFixes"] == 0
    assert preview["plannedChanges"] == []
    assert any(
        issue["ruleId"] == "OOXML-PACKAGE-SIGNATURE"
        for issue in preview["issues"]
    )
    with pytest.raises(WordStandardizationError, match="digitally signed"):
        process_docx(
            content,
            profile="sector_template",
            mode="apply_fix",
            expected_analysis_hash=preview["analysisHash"],
        )


def test_quoted_shell_marker_in_body_is_never_a_sector_auto_fix():
    document = Document()
    for value in (
        "Nội dung thứ nhất",
        "Nội dung thứ hai",
        "Trích dẫn sau đây:",
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "chỉ là nội dung được trích dẫn",
        "Kết thúc trích dẫn",
    ):
        document.add_paragraph(value)
    output = BytesIO()
    document.save(output)
    content = output.getvalue()

    preview = process_docx(
        content,
        profile="sector_template",
        mode="preview_fix",
    ).report
    marker = next(
        field
        for field in preview["fields"]
        if field["semantic"] == "document.national_header"
    )

    assert marker["mutationSafe"] is False
    assert preview["summary"]["safeFixes"] == 0
    assert preview["plannedChanges"] == []
    applied = process_docx(
        content,
        profile="sector_template",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    )
    assert applied.content == content


def test_uncalibrated_title_candidates_are_not_strict_body_paragraphs():
    document = Document()
    document.add_paragraph("THÔNG BÁO")
    document.add_paragraph("Về việc kiểm tra thể thức")
    document.add_paragraph("Nội dung nghiệp vụ cần trình bày.")
    output = BytesIO()
    document.save(output)
    content = output.getvalue()

    preview = process_docx(
        content,
        profile="n30_strict",
        mode="preview_fix",
    ).report
    body_targets = {
        target["paragraph"]
        for change in preview["plannedChanges"]
        if change["ruleId"].startswith("N30-BODY-")
        for target in change["targets"]
    }

    assert 0 not in body_targets
    assert 1 not in body_targets
    assert 2 in body_targets


def test_owned_ooxml_properties_are_inserted_in_schema_order():
    content = _fixture_docx()
    preview = process_docx(
        content,
        profile="n30_strict",
        mode="preview_fix",
    ).report
    applied = process_docx(
        content,
        profile="n30_strict",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    ).content
    root = etree.fromstring(_document_xml(applied).encode("utf-8"))
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    for query, order in ((".//w:rPr", _RPR_ORDER), (".//w:pPr", _PPR_ORDER)):
        ranks = {name: index for index, name in enumerate(order)}
        for properties in root.findall(query, namespaces=namespace):
            actual = [
                ranks[etree.QName(child).localname]
                for child in properties
                if etree.QName(child).localname in ranks
            ]
            assert actual == sorted(actual)


def test_exact_line_spacing_is_reported_but_not_auto_converted():
    document = Document(BytesIO(_fixture_docx()))
    body = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Nội dung nghiệp vụ")
    )
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    body.paragraph_format.line_spacing = Pt(12)
    output = BytesIO()
    document.save(output)
    content = output.getvalue()
    preview = process_docx(
        content,
        profile="n30_strict",
        mode="preview_fix",
    ).report

    issue = next(
        item for item in preview["issues"]
        if item["ruleId"] == "N30-BODY-LINE-RULE"
    )
    assert issue["fixPolicy"] == "PREVIEW_ONLY"
    assert all(
        item["ruleId"] != "N30-BODY-LINE-SPACING"
        for item in preview["plannedChanges"]
    )
    applied = process_docx(
        content,
        profile="n30_strict",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    ).content
    assert 'w:lineRule="exact"' in _document_xml(applied)


def test_strict_body_layout_excludes_headings_and_table_cells():
    document = Document()
    for value in (
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "Độc lập - Tự do - Hạnh phúc",
        "THÔNG BÁO",
        "Về việc kiểm tra thể thức",
        "CHƯƠNG I",
        "QUY ĐỊNH CHUNG",
        "Nội dung nghiệp vụ cần căn chỉnh.",
    ):
        document.add_paragraph(value)
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Dữ liệu trong ô bảng"
    output = BytesIO()
    document.save(output)

    preview = process_docx(
        output.getvalue(),
        profile="n30_strict",
        mode="preview_fix",
    ).report
    body_targets = {
        target["paragraph"]
        for change in preview["plannedChanges"]
        if change["ruleId"].startswith("N30-BODY-")
        for target in change["targets"]
    }
    rule_ids = {item["ruleId"] for item in preview["issues"]}

    assert 4 not in body_targets
    assert 5 not in body_targets
    assert 6 in body_targets
    assert 7 not in body_targets
    assert "N30-HEADING-LAYOUT" in rule_ids
    assert "N30-TABLE-CELL-LAYOUT" in rule_ids


def test_strict_body_layout_excludes_direct_and_inherited_numbering():
    document = Document()
    for value in (
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "Độc lập - Tự do - Hạnh phúc",
        "THÔNG BÁO",
        "Về việc kiểm tra thể thức",
    ):
        document.add_paragraph(value)

    direct_numbered = document.add_paragraph("Danh sách đánh số trực tiếp")
    direct_num_pr = OxmlElement("w:numPr")
    direct_num_id = OxmlElement("w:numId")
    direct_num_id.set(qn("w:val"), "1")
    direct_num_pr.append(direct_num_id)
    direct_numbered._p.get_or_add_pPr().append(direct_num_pr)

    numbered_parent = document.styles.add_style(
        "Numbered Layout Parent", WD_STYLE_TYPE.PARAGRAPH
    )
    inherited_num_pr = OxmlElement("w:numPr")
    inherited_num_id = OxmlElement("w:numId")
    inherited_num_id.set(qn("w:val"), "2")
    inherited_num_pr.append(inherited_num_id)
    numbered_parent.element.get_or_add_pPr().append(inherited_num_pr)
    numbered_child = document.styles.add_style(
        "Numbered Layout Child", WD_STYLE_TYPE.PARAGRAPH
    )
    numbered_child.base_style = numbered_parent
    document.add_paragraph(
        "Danh sách đánh số kế thừa", style=numbered_child
    )
    document.add_paragraph("Nội dung thân bài thông thường")

    output = BytesIO()
    document.save(output)
    preview = process_docx(
        output.getvalue(),
        profile="n30_strict",
        mode="preview_fix",
    ).report
    body_targets = {
        target["paragraph"]
        for change in preview["plannedChanges"]
        if change["ruleId"].startswith("N30-BODY-")
        for target in change["targets"]
    }
    numbered_issue = next(
        issue
        for issue in preview["issues"]
        if issue["ruleId"] == "N30-NUMBERED-LAYOUT"
    )

    assert {4, 5}.isdisjoint(body_targets)
    assert 6 in body_targets
    assert numbered_issue["target"] == "numbered_layout"
    assert numbered_issue["fixPolicy"] == "MANUAL_REVIEW"
    assert numbered_issue["affectedCount"] == 2


def test_landscape_section_is_reported_even_when_document_has_a_table():
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Bảng hợp lệ"
    output = BytesIO()
    document.save(output)

    report = process_docx(
        output.getvalue(),
        profile="sector_template",
        mode="audit",
    ).report

    orientation = next(
        item for item in report["issues"]
        if item["ruleId"] == "N30-PAGE-ORIENTATION"
    )
    assert orientation["fixPolicy"] == "PREVIEW_ONLY"
    assert orientation["affectedCount"] == 1


def test_auto_text_color_is_normalized_to_explicit_black():
    document = Document()
    for value in (
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "Độc lập - Tự do - Hạnh phúc",
        "THÔNG BÁO",
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        run.font.size = Pt(13)
        run.bold = True
        run.italic = False
        fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attribute}"), "Times New Roman")
        color = run._r.get_or_add_rPr().get_or_add_color()
        color.set(qn("w:val"), "auto")
    output = BytesIO()
    document.save(output)
    content = output.getvalue()
    preview = process_docx(
        content,
        profile="sector_template",
        mode="preview_fix",
    ).report

    assert any(
        item["ruleId"] == "N30-SHELL-COLOR"
        for item in preview["plannedChanges"]
    )
    applied = process_docx(
        content,
        profile="sector_template",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    ).content
    root = etree.fromstring(_document_xml(applied).encode("utf-8"))
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    assert {
        color.get(qn("w:val"))
        for color in root.findall(".//w:color", namespaces=namespace)
    } == {"000000"}


def test_character_unit_indent_override_is_removed_by_strict_fix():
    document = Document(BytesIO(_fixture_docx()))
    body = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Nội dung nghiệp vụ")
    )
    ppr = body._p.get_or_add_pPr()
    indent = ppr.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        ppr.append(indent)
    indent.set(qn("w:firstLine"), "567")
    indent.set(qn("w:firstLineChars"), "200")
    indent.set(qn("w:hangingChars"), "100")
    output = BytesIO()
    document.save(output)
    content = output.getvalue()
    preview = process_docx(
        content,
        profile="n30_strict",
        mode="preview_fix",
    ).report

    assert any(
        item["ruleId"] == "N30-BODY-INDENT"
        for item in preview["plannedChanges"]
    )
    applied = process_docx(
        content,
        profile="n30_strict",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    ).content
    root = etree.fromstring(_document_xml(applied).encode("utf-8"))
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    body_paragraph = next(
        paragraph
        for paragraph in root.findall(".//w:p", namespaces=namespace)
        if "".join(
            node.text or ""
            for node in paragraph.findall(".//w:t", namespaces=namespace)
        ).startswith("Nội dung nghiệp vụ")
    )
    applied_indent = body_paragraph.find("w:pPr/w:ind", namespaces=namespace)
    assert applied_indent.get(qn("w:firstLine")) == "567"
    assert applied_indent.get(qn("w:firstLineChars")) is None
    assert applied_indent.get(qn("w:hanging")) is None
    assert applied_indent.get(qn("w:hangingChars")) is None


def test_automatic_paragraph_spacing_is_preview_only():
    document = Document(BytesIO(_fixture_docx()))
    body = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Nội dung nghiệp vụ")
    )
    ppr = body._p.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:afterAutospacing"), "1")
    output = BytesIO()
    document.save(output)
    content = output.getvalue()
    preview = process_docx(
        content,
        profile="n30_strict",
        mode="preview_fix",
    ).report

    issue = next(
        item for item in preview["issues"]
        if item["ruleId"] == "N30-BODY-SPACING-AUTOMATIC"
    )
    assert issue["fixPolicy"] == "PREVIEW_ONLY"
    assert all(
        not (
            item["ruleId"] == "N30-BODY-SPACING"
            and any(target["paragraph"] == 6 for target in item["targets"])
        )
        for item in preview["plannedChanges"]
    )
    applied = process_docx(
        content,
        profile="n30_strict",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    ).content
    assert 'w:afterAutospacing="1"' in _document_xml(applied)


def test_contextual_spacing_is_preview_only_for_adjacent_same_style():
    document = Document(BytesIO(_fixture_docx()))
    body = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Nội dung nghiệp vụ")
    )
    ppr = body._p.get_or_add_pPr()
    contextual = OxmlElement("w:contextualSpacing")
    contextual.set(qn("w:val"), "1")
    ppr.append(contextual)
    output = BytesIO()
    document.save(output)
    content = output.getvalue()

    preview = process_docx(
        content,
        profile="n30_strict",
        mode="preview_fix",
    ).report

    automatic = next(
        item for item in preview["issues"]
        if item["ruleId"] == "N30-BODY-SPACING-AUTOMATIC"
    )
    assert automatic["fixPolicy"] == "PREVIEW_ONLY"
    applied = process_docx(
        content,
        profile="n30_strict",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    ).content
    assert "w:contextualSpacing" in _document_xml(applied)


def test_default_paragraph_style_applies_to_implicit_and_explicit_normal():
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    style_fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        style_fonts.set(qn(f"w:{slot}"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Mm(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    contextual = OxmlElement("w:contextualSpacing")
    contextual.set(qn("w:val"), "1")
    normal.element.get_or_add_pPr().append(contextual)

    for value in (
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "Độc lập - Tự do - Hạnh phúc",
        "THÔNG BÁO",
        "Về việc kiểm tra thể thức",
    ):
        document.add_paragraph(value)
    implicit_normal = document.add_paragraph("Nội dung dùng ngầm định Normal.")
    explicit_normal = document.add_paragraph("Nội dung khai báo rõ Normal.")
    explicit_style = OxmlElement("w:pStyle")
    explicit_style.set(qn("w:val"), "Normal")
    explicit_normal._p.get_or_add_pPr().insert(0, explicit_style)
    assert implicit_normal._p.find(qn("w:pPr")) is None

    output = BytesIO()
    document.save(output)
    preview = process_docx(
        output.getvalue(),
        profile="n30_strict",
        mode="preview_fix",
    ).report

    body_indexes = {4, 5}
    assert all(
        not (
            change["ruleId"].startswith("N30-BODY-")
            and any(
                target["paragraph"] in body_indexes
                for target in change["targets"]
            )
        )
        for change in preview["plannedChanges"]
    )
    automatic = next(
        issue
        for issue in preview["issues"]
        if issue["ruleId"] == "N30-BODY-SPACING-AUTOMATIC"
    )
    assert automatic["current"] == [
        {"paragraph": 4, "contextualSpacing": True}
    ]


def test_large_semantic_report_is_bounded_and_inventory_is_hashed():
    document = Document()
    for index in range(3000):
        document.add_paragraph(
            f"CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM - mục kiểm tra {index:04d}"
        )
    output = BytesIO()
    document.save(output)

    report = process_docx(
        output.getvalue(),
        profile="sector_template",
        mode="preview_fix",
    ).report
    encoded = json.dumps(
        report,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")

    assert len(encoded) < MAX_REPORT_BYTES
    assert report["fieldInventory"]["totalCount"] == 3000
    assert report["fieldInventory"]["truncated"] is True
    assert report["issueInventory"]["totalCount"] > len(report["issues"])
    assert report["issueInventory"]["truncated"] is True


def test_versioned_rule_bundle_is_validated_and_policy_changes_are_hashed():
    rules, semantic_fields, digest = _rules()
    assert rules["rule_set"]["official_source"].startswith(
        "https://vanban.chinhphu.vn/"
    )
    assert len(digest) == 64
    invalid = json.loads(json.dumps(rules))
    invalid["profiles"]["sector_template"]["auto_fix"] = True

    with pytest.raises(RuntimeError, match="profile policy"):
        _validate_rule_bundle(invalid, semantic_fields)


def test_document_worker_operation_returns_report_or_docx_bytes():
    content = _fixture_docx()
    report = _run_operation("standardize_docx", {
        "content": content,
        "profile": "sector_template",
        "mode": "preview_fix",
    })

    assert report["mode"] == "preview_fix"
    output = _run_operation("standardize_docx", {
        "content": content,
        "profile": "sector_template",
        "mode": "apply_fix",
        "expectedAnalysisHash": report["analysisHash"],
    })
    assert isinstance(output, bytes)
    assert output.startswith(b"PK")


def test_standardizer_crosses_the_isolated_document_worker_boundary():
    content = _fixture_docx()
    report = run_document_job(
        "standardize_docx",
        {
            "content": content,
            "profile": "sector_template",
            "mode": "preview_fix",
        },
        timeout_seconds=30,
    )

    output = run_document_job(
        "standardize_docx",
        {
            "content": content,
            "profile": "sector_template",
            "mode": "apply_fix",
            "expectedAnalysisHash": report["analysisHash"],
        },
        timeout_seconds=30,
    )

    assert report["invariants"]["status"] == "PASS"
    assert isinstance(output, bytes)
    assert output.startswith(b"PK")


def test_catalog_creates_versioned_standardized_draft_with_provenance():
    source_content = _fixture_docx()
    preview = process_docx(
        source_content,
        profile="sector_template",
        mode="preview_fix",
    ).report
    corrected = process_docx(
        source_content,
        profile="sector_template",
        mode="apply_fix",
        expected_analysis_hash=preview["analysisHash"],
    ).content
    source_sha = hashlib.sha256(source_content).hexdigest()
    run = {
        "id": "preflight-a",
        "templateVersionId": "version-a",
        "templateSha256": source_sha,
        "reportHash": "f" * 64,
        "report": {"standardization": preview},
        "runAt": datetime.now(timezone.utc).isoformat(),
    }
    captured = {}
    repository_state = {
        "rowVersion": 4,
        "draftVersionId": "version-a",
        "existing": None,
        "creates": 0,
    }

    class Repository:
        cursor = object()

        @staticmethod
        def get_version(_organization_id, version_id):
            assert version_id == "version-a"
            return {
                "id": "version-a",
                "templateId": "template-a",
                "sha256": source_sha,
                "originalFilename": "source.docx",
            }

        @staticmethod
        def get_preflight(_organization_id, run_id):
            assert run_id == "preflight-a"
            return run

        @staticmethod
        def create_draft_version(**values):
            captured.update(values)
            repository_state["creates"] += 1
            version_id = (
                "version-b" if repository_state["creates"] == 1 else "version-c"
            )
            repository_state["rowVersion"] = values["expected_row_version"] + 1
            repository_state["draftVersionId"] = version_id
            repository_state["existing"] = {"id": version_id}
            return {
                "id": "template-a",
                "rowVersion": repository_state["rowVersion"],
                "draftVersionId": version_id,
            }, None

        @staticmethod
        def get_template(_organization_id, _template_id):
            return {
                "id": "template-a",
                "rowVersion": repository_state["rowVersion"],
                "publishedVersionId": None,
                "draftVersionId": repository_state["draftVersionId"],
            }

        @staticmethod
        def validate_template_cas(
            _organization_id, _template_id, expected_row_version,
        ):
            template = Repository.get_template("org-a", "template-a")
            if expected_row_version != repository_state["rowVersion"]:
                return template, "STALE"
            return template, None

        @staticmethod
        def find_standardized_version(**_values):
            return repository_state["existing"]

    class Storage:
        @staticmethod
        def put(_organization_id, content):
            digest = hashlib.sha256(content).hexdigest()
            return f"v1/test/{digest}.docx", digest, len(content)

        @staticmethod
        def read(_organization_id, _storage_key, _sha256):
            return source_content

    audits = []
    catalog = WordTemplateCatalog(
        Repository(),
        Storage(),
        audit=lambda action, **values: audits.append((action, values)),
    )

    result = catalog.create_standardized_draft(
        organization_id="org-a",
        template_id="template-a",
        source_version_id="version-a",
        accepted_preflight_run_id="preflight-a",
        expected_row_version=4,
        profile="sector_template",
        standardized_content=corrected,
        actor_user_id="manager-a",
        reason="Chuẩn hóa trước khi phát hành",
    )

    assert result == {
        "template": {
            "id": "template-a",
            "rowVersion": 5,
            "draftVersionId": "version-b",
        },
        "created": True,
        "sourceVersionId": "version-a",
        "draftVersionId": "version-b",
    }
    assert captured["expected_row_version"] == 4
    assert captured["version"]["source_version_id"] == "version-a"
    manifest = json.loads(captured["version"]["creation_manifest_json"])
    assert manifest["action"] == "STANDARDIZE"
    assert manifest["sourceVersionId"] == "version-a"
    assert manifest["metadata"]["standardization"]["analysisHash"] == (
        preview["analysisHash"]
    )
    assert audits[0][0] == "document.word_template_standardized"

    replay = catalog.create_standardized_draft(
        organization_id="org-a",
        template_id="template-a",
        source_version_id="version-a",
        accepted_preflight_run_id="preflight-a",
        expected_row_version=5,
        profile="sector_template",
        standardized_content=corrected,
        actor_user_id="manager-a",
        reason="Thử lại cùng một lần chuẩn hóa",
    )
    assert replay["created"] is False
    assert replay["replayed"] is True
    assert replay["draftVersionId"] == "version-b"
    assert repository_state["creates"] == 1
    assert audits[-1][0] == "document.word_template_standardization_replayed"

    repository_state["rowVersion"] = 6
    repository_state["draftVersionId"] = "version-unrelated"
    retired_duplicate = catalog.create_standardized_draft(
        organization_id="org-a",
        template_id="template-a",
        source_version_id="version-a",
        accepted_preflight_run_id="preflight-a",
        expected_row_version=6,
        profile="sector_template",
        standardized_content=corrected,
        actor_user_id="manager-a",
        reason="Tạo lại vì kết quả cũ không còn là draft hiện tại",
    )
    assert retired_duplicate["created"] is True
    assert repository_state["creates"] == 2

    with pytest.raises(CatalogConflictError):
        catalog.create_standardized_draft(
            organization_id="org-a",
            template_id="template-a",
            source_version_id="version-a",
            accepted_preflight_run_id="preflight-a",
            expected_row_version=4,
            profile="sector_template",
            standardized_content=source_content,
            actor_user_id="manager-a",
            reason="No-op nhưng row version đã cũ",
        )

    with pytest.raises(CatalogConflictError):
        catalog.get_standardization_candidate(
            organization_id="org-a",
            version_id="version-a",
            accepted_preflight_run_id="preflight-a",
            profile="n30_strict",
        )
