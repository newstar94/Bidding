from io import BytesIO
import json

from docx import Document

from backend.documents.custom_exporter import generate_report_from_custom_template
from backend.documents.detailed_evaluation_context import build_detailed_evaluation_context
from backend.documents.docx_context_policy import project_docx_context, seal_docx_context
from backend.documents.docx_mapping_service import apply_custom_mappings
from backend.documents.word_defaults import build_default_word_mappings


def _source_context():
    package = {
        "danh_gia_hsdt_metadata": json.dumps({
            "is1G2T": True,
            "technical": {
                "criteria": [
                    {
                        "id": "validity-1",
                        "code": "VALIDITY",
                        "name": "Bảo đảm dự thầu",
                        "group": "validity",
                        "resultType": "pass_fail",
                        "required": True,
                        "stt": "1",
                        "order": 0,
                    },
                    {
                        "id": "capacity-1",
                        "code": "CAPACITY",
                        "name": "Doanh thu bình quân",
                        "group": "capacity",
                        "resultType": "pass_fail",
                        "required": True,
                        "requirement": "Tối thiểu 10 tỷ đồng",
                        "stt": "2.1",
                        "order": 1,
                    },
                    {
                        "id": "technical-1",
                        "code": "TECHNICAL",
                        "name": "Giải pháp kỹ thuật",
                        "group": "technical",
                        "resultType": "score",
                        "required": True,
                        "maxScore": 100,
                        "minScore": 70,
                        "stt": "3",
                        "order": 2,
                    },
                ],
            },
            "financial": {
                "criteria": [{
                    "id": "financial-1",
                    "code": "FINANCIAL",
                    "name": "Giá dự thầu sau sửa lỗi",
                    "group": "financial",
                    "resultType": "number",
                    "required": True,
                    "stt": "1",
                    "order": 0,
                }],
            },
        }, ensure_ascii=False),
    }
    bid = {
        "id": "opening-1",
        "nha_thau_id": "contractor-1",
        "ten_nha_thau": "Công ty ABC",
        "loai_nha_thau": "Độc lập",
        "ma_phan_lo": "LOT-01",
        "ten_phan_lo": "Lô 01",
        "bao_cao_danh_gia_chi_tiet_list": [
            {
                "id": "report-technical",
                "vong_danh_gia_id": "round-technical",
                "loai_vong": "technical",
                "trang_thai": "completed",
                "ket_luan": "Đạt",
                "chi_tiet_list": [
                    {
                        "tieu_chi_danh_gia_id": "validity-1",
                        "ket_qua": "pass",
                        "noi_dung_hsdt": "Thư bảo lãnh số 01",
                        "nhan_xet": "Hợp lệ",
                        "extension": {"ketQuaTuDong": "pass"},
                        "ly_do_khong_dat": "Không được xuất",
                    },
                    {
                        "tieu_chi_danh_gia_id": "capacity-1",
                        "ket_qua": "fail",
                        "noi_dung_hsdt": "9 tỷ đồng",
                        "nhan_xet": "Không đáp ứng",
                        "extension": {"ketQuaTuDong": "fail"},
                    },
                    {
                        "tieu_chi_danh_gia_id": "technical-1",
                        "ket_qua": "pass",
                        "diem": 85,
                        "nhan_xet": "Đáp ứng",
                    },
                ],
            },
            {
                "id": "report-financial",
                "vong_danh_gia_id": "round-financial",
                "loai_vong": "financial",
                "trang_thai": "draft",
                "ket_luan": "",
                "chi_tiet_list": [{
                    "tieu_chi_danh_gia_id": "financial-1",
                    "ket_qua": "pass",
                    "noi_dung_hsdt": "9.500.000.000 đồng",
                }],
            },
        ],
    }
    return package, [bid]


def test_detailed_evaluation_context_flattens_reports_and_group_tables():
    package, bids = _source_context()
    context = build_detailed_evaluation_context(package, bids)

    assert len(context["detailed_evaluation_reports"]) == 2
    assert len(context["detailed_evaluation_rows"]) == 4
    assert len(context["detailed_evaluation_validity_rows"]) == 1
    assert len(context["detailed_evaluation_capacity_rows"]) == 1
    assert len(context["detailed_evaluation_technical_rows"]) == 1
    assert len(context["detailed_evaluation_financial_rows"]) == 1
    validity = context["detailed_evaluation_validity_rows"][0]
    assert validity["stt"] == "1"
    assert validity["ten_tieu_chi"] == "Bảo đảm dự thầu"
    assert validity["ten_nha_thau"] == "Công ty ABC"
    assert validity["ket_qua_tu_dong_dat"] == "x"
    assert validity["ket_qua_tu_dong_khong_dat"] == ""
    assert validity["ket_qua_chuyen_gia_dat"] == "x"
    assert validity["ket_qua_chuyen_gia_khong_dat"] == ""
    assert validity["noi_dung_hsdt"] == "Thư bảo lãnh số 01"
    assert validity["nhan_xet"] == "Hợp lệ"
    capacity = context["detailed_evaluation_capacity_rows"][0]
    assert capacity["yeu_cau"] == "Tối thiểu 10 tỷ đồng"
    assert capacity["ket_qua_chuyen_gia_khong_dat"] == "x"
    technical = context["detailed_evaluation_technical_rows"][0]
    assert technical["diem"] == 85
    assert technical["diem_toi_da"] == 100
    assert technical["diem_toi_thieu"] == 70
    assert "ly_do_khong_dat" not in validity
    assert "yeu_cau_lam_ro" not in validity
    report = context["detailed_evaluation_reports"][0]
    assert report["ten_vong"] == "Đánh giá kỹ thuật"
    assert len(report["ds_hop_le"]) == 1
    assert len(report["ds_nang_luc"]) == 1
    assert len(report["ds_ky_thuat"]) == 1
    assert report["ds_tai_chinh"] == []


def test_detailed_evaluation_lists_are_allowlisted_and_receive_default_aliases():
    package, bids = _source_context()
    raw = {
        "goi_thau": package,
        "nha_thau": bids,
        **build_detailed_evaluation_context(package, bids),
    }
    projected = project_docx_context("evaluation", raw)
    mappings = [
        (item["ten_bien"], item["source_table"], item["source_column"])
        for item in build_default_word_mappings()
        if item["ten_bien"].startswith("ds_dgct")
        or item["ten_bien"] == "ds_bao_cao_dgct"
    ]
    apply_custom_mappings(projected, mappings)
    sealed, manifest = seal_docx_context("evaluation", projected, mappings)

    assert len(sealed["ds_bao_cao_dgct"]) == 2
    assert len(sealed["ds_dgct_hop_le"]) == 1
    assert sealed["ds_bao_cao_dgct"][0]["ds_hop_le"][0]["stt"] == "1"
    assert "unknown" not in sealed["ds_dgct_hop_le"][0]
    assert "ds_bao_cao_dgct" in manifest["custom_root_keys"]


def test_word_row_loop_renders_detailed_evaluation_placeholders(tmp_path):
    template_path = tmp_path / "detailed-template.docx"
    document = Document()
    table = document.add_table(rows=4, cols=3)
    table.rows[0].cells[0].text = "STT"
    table.rows[0].cells[1].text = "Tiêu chí"
    table.rows[0].cells[2].text = "Đạt"
    table.rows[1].cells[0].text = "{#ds_dgct_hop_le}"
    table.rows[2].cells[0].text = "{stt}"
    table.rows[2].cells[1].text = "{ten_tieu_chi}"
    table.rows[2].cells[2].text = "{ket_qua_chuyen_gia_dat}"
    table.rows[3].cells[0].text = "{/ds_dgct_hop_le}"
    document.save(template_path)

    package, bids = _source_context()
    context = build_detailed_evaluation_context(package, bids)
    context["ds_dgct_hop_le"] = context["detailed_evaluation_validity_rows"]
    sealed, manifest = seal_docx_context("evaluation", context, [
        ("ds_dgct_hop_le", "detailed_evaluation_validity_rows", ""),
    ])
    rendered = generate_report_from_custom_template(template_path, sealed, manifest)
    (tmp_path / "nested-detailed-rendered.docx").write_bytes(rendered.getvalue())
    output = Document(BytesIO(rendered.getvalue()))
    cell_text = [
        cell.text
        for rendered_table in output.tables
        for row in rendered_table.rows
        for cell in row.cells
    ]
    assert "1" in cell_text
    assert "Bảo đảm dự thầu" in cell_text
    assert "x" in cell_text


def test_word_nested_report_loop_renders_one_table_per_bidder_round(tmp_path):
    template_path = tmp_path / "nested-detailed-template.docx"
    document = Document()
    document.add_paragraph("{%p for bc in ds_bao_cao_dgct %}")
    document.add_paragraph("Nhà thầu: {{ bc.ten_nha_thau }} - {{ bc.ten_vong }}")
    table = document.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "STT"
    table.rows[0].cells[1].text = "Tiêu chí"
    table.rows[1].cells[0].text = "{%tr for tc in bc.ds_hop_le %}"
    table.rows[2].cells[0].text = "{{ tc.stt }}"
    table.rows[2].cells[1].text = "{{ tc.ten_tieu_chi }}"
    table.rows[3].cells[0].text = "{%tr endfor %}"
    document.add_paragraph("{%p endfor %}")
    document.save(template_path)

    package, bids = _source_context()
    context = build_detailed_evaluation_context(package, bids)
    context["ds_bao_cao_dgct"] = context["detailed_evaluation_reports"]
    sealed, manifest = seal_docx_context("evaluation", context, [
        ("ds_bao_cao_dgct", "detailed_evaluation_reports", ""),
    ])
    rendered = generate_report_from_custom_template(template_path, sealed, manifest)
    output = Document(BytesIO(rendered.getvalue()))
    text = "\n".join(paragraph.text for paragraph in output.paragraphs)
    table_text = "\n".join(
        cell.text
        for rendered_table in output.tables
        for row in rendered_table.rows
        for cell in row.cells
    )
    assert "Nhà thầu: Công ty ABC - Đánh giá kỹ thuật" in text
    assert "Bảo đảm dự thầu" in table_text
