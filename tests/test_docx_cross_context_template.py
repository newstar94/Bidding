from io import BytesIO
from pathlib import Path

from docx import Document
import pytest

from backend.documents.custom_exporter import (
    TemplateRenderError,
    generate_report_from_custom_template,
)
from backend.documents.docx_context_policy import MANIFEST_VERSION
from backend.documents.docx_context_policy import project_docx_context, seal_docx_context
from backend.documents.docx_formula_service import apply_computed_mappings
from backend.documents.docx_mapping_service import apply_custom_mappings
from backend.documents.routes_docx import _scope_contracts_for_word_publication


def _all_document_text(document):
    return "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )


def test_combined_sample_template_documents_derived_variable_usage():
    template_path = (
        Path(__file__).resolve().parents[1]
        / "docs"
        / "BiddingFlow_Mau_Kiem_Thu_Xuat_Word.docx"
    )

    text = _all_document_text(Document(template_path))

    assert "6. Hướng dẫn sử dụng biến dẫn xuất" in text
    assert "{bangchu_gia_gt}" in text
    assert "{S_tg_dang_tai_kh}" in text
    assert "05/3/2026" in text
    assert "05/01/2026" in text
    assert "Mười hai triệu đồng" in text


def test_combined_sample_template_renders_evaluation_list_source_fields():
    template_path = (
        Path(__file__).resolve().parents[1]
        / "docs"
        / "BiddingFlow_Mau_Kiem_Thu_Xuat_Word.docx"
    )
    context = {
        "ds_to_chuyen_gia": [{
            "ho_ten": "Nguyễn Văn Chuyên Gia",
            "chuc_vu": "Tổ trưởng",
            "cong_viec": "Đánh giá hồ sơ",
            "so_chung_chi": "CG-01",
            "ngay_cap_chung_chi": "2026-01-02",
        }],
        "ds_to_tham_dinh": [{
            "ho_ten": "Trần Thị Thẩm Định",
            "chuc_vu": "Tổ trưởng",
            "cong_viec": "Thẩm định kết quả",
            "so_chung_chi": "TD-01",
            "ngay_cap_chung_chi": "2026-02-03",
        }],
        "ds_mo_thau": [{
            "ten_nha_thau": "Công ty Nhà thầu A",
            "loai_nha_thau": "Độc lập",
            "gia_du_thau": 123456789,
            "ty_le_giam_gia": 5,
            "gia_sau_giam_gia": 117283949,
            "hieu_luc_hsdt": "90 ngày",
            "gia_tri_dam_bao": 1000000,
            "danh_gia_hop_le": "Đạt",
            "danh_gia_nang_luc": "Đạt",
            "danh_gia_ky_thuat": "Đạt",
            "danh_gia_tai_chinh": "Đạt",
            "gia_xep_hang": 117283949,
            "danh_gia_ket_luan": "Xếp hạng 1",
        }],
        # These scalar aliases are present in a real evaluation context. They
        # must not be treated as fields of each list item.
        **{
            key: ""
            for key in (
                "ten_cg", "chuc_vu_cg", "cong_viec_cg",
                "so_chung_chi_cg", "ngay_cap_chung_chi_cg",
                "mt_ten_nt", "mt_loai_nt", "mt_gia_du_thau",
                "mt_ty_le_giam_gia", "mt_gia_sau_giam_gia",
                "mt_hieu_luc_hsdt", "mt_gia_tri_dam_bao",
                "mt_dg_hop_le", "mt_dg_nang_luc", "mt_dg_ky_thuat",
                "mt_dg_tai_chinh", "mt_gia_xep_hang", "mt_dg_ket_luan",
            )
        },
    }

    output = generate_report_from_custom_template(
        str(template_path),
        context,
        {
            "version": MANIFEST_VERSION,
            "document_type": "evaluation",
            "root_keys": sorted(context),
            "custom_root_keys": sorted(context),
            "image_fields": {},
            "media_organization_id": "org-a",
        },
    )

    text = _all_document_text(Document(BytesIO(output.getvalue())))
    assert "Nguyễn Văn Chuyên Gia" in text
    assert "Trần Thị Thẩm Định" in text
    assert "Công ty Nhà thầu A" in text


def test_combined_sample_template_renders_all_contracts_and_formats_money():
    template_path = (
        Path(__file__).resolve().parents[1]
        / "docs"
        / "BiddingFlow_Mau_Kiem_Thu_Xuat_Word.docx"
    )
    context = {
        "ds_hop_dong": [
            {
                "phan_loai": "Thẩm định",
                "ten_hop_dong": "Hợp đồng thẩm định",
                "so_hop_dong": "01/2026/HĐTĐ",
                "ngay_ky": "2026-08-01",
                "gia_tri": 9000000,
                "loai_hop_dong": "Trọn gói",
                "thoi_gian_thuc_hien": "30 ngày",
                "trang_thai_hop_dong": "Đang thực hiện",
            },
            {
                "phan_loai": "Tư vấn",
                "ten_hop_dong": "Hợp đồng tư vấn",
                "so_hop_dong": "01/2026/HĐTV",
                "ngay_ky": "2026-07-01",
                "gia_tri": 12000000,
                "loai_hop_dong": "Trọn gói",
                "thoi_gian_thuc_hien": "45 ngày",
                "trang_thai_hop_dong": "Đang thực hiện",
            },
        ],
        "ds_mo_thau": [{
            "ten_nha_thau": "Công ty Nhà thầu A",
            "loai_nha_thau": "Độc lập",
            "gia_du_thau": 100000000,
            "ty_le_giam_gia": 5,
            "gia_sau_giam_gia": 95000000,
            "hieu_luc_hsdt": "90 ngày",
            "gia_tri_dam_bao": 10000000,
            "danh_gia_hop_le": "Đạt",
            "danh_gia_nang_luc": "Đạt",
            "danh_gia_ky_thuat": "Đạt",
            "danh_gia_tai_chinh": "Đạt",
            "gia_xep_hang": "895000000",
            "danh_gia_ket_luan": "Xếp hạng 1",
        }],
        **{
            key: ""
            for key in (
                "hd_phan_loai", "ten_hd", "so_hd", "ngay_ky_hd",
                "hd_gia_tri", "loai_hd", "tg_thuc_hien_hd",
                "trang_thai_hd",
            )
        },
    }

    output = generate_report_from_custom_template(
        str(template_path),
        context,
        {
            "version": MANIFEST_VERSION,
            "document_type": "contract",
            "root_keys": sorted(context),
            "custom_root_keys": sorted(context),
            "image_fields": {},
            "media_organization_id": "org-a",
        },
    )

    text = _all_document_text(Document(BytesIO(output.getvalue())))
    assert "Hợp đồng thẩm định" in text
    assert "Hợp đồng tư vấn" in text
    assert "9.000.000" in text
    assert "12.000.000" in text
    assert "10.000.000" in text
    assert "895.000.000" in text
    assert "10000000" not in text
    assert "895000000" not in text


def test_combined_template_renders_contracts_from_evaluation_context():
    template_path = (
        Path(__file__).resolve().parents[1]
        / "docs"
        / "BiddingFlow_Mau_Kiem_Thu_Xuat_Word.docx"
    )
    context = project_docx_context(
        "evaluation",
        {
            "hop_dong_list": [
                {
                    "id": "contract-appraisal",
                    "phan_loai": "Thẩm định",
                    "ten_hop_dong": "Tư vấn thẩm định HSMT",
                    "so_hop_dong": "01/2026/HĐTĐ",
                    "gia_tri": 9000000,
                },
                {
                    "id": "contract-consulting",
                    "phan_loai": "Tư vấn",
                    "ten_hop_dong": "Tư vấn lập HSMT",
                    "so_hop_dong": "01/2026/HĐTV",
                    "gia_tri": 12000000,
                },
            ],
        },
    )
    apply_custom_mappings(
        context,
        [("ds_hop_dong", "hop_dong_list", "")],
    )

    output = generate_report_from_custom_template(
        str(template_path),
        context,
        {
            "version": MANIFEST_VERSION,
            "document_type": "evaluation",
            "root_keys": sorted(context),
            "custom_root_keys": ["ds_hop_dong"],
            "image_fields": {},
            "media_organization_id": "org-a",
        },
    )

    text = _all_document_text(Document(BytesIO(output.getvalue())))
    assert "Tư vấn thẩm định HSMT" in text
    assert "Tư vấn lập HSMT" in text
    assert "9.000.000" in text
    assert "12.000.000" in text


def test_combined_template_renders_only_consulting_contract_for_consultant_step():
    template_path = (
        Path(__file__).resolve().parents[1]
        / "docs"
        / "BiddingFlow_Mau_Kiem_Thu_Xuat_Word.docx"
    )
    context = project_docx_context(
        "evaluation",
        {
            "hop_dong_list": [
                {
                    "id": "contract-appraisal",
                    "phan_loai": "Thẩm định",
                    "ten_hop_dong": "Tư vấn thẩm định HSMT",
                },
                {
                    "id": "contract-consulting",
                    "phan_loai": "Tư vấn",
                    "ten_hop_dong": "Tư vấn lập HSMT",
                },
            ],
        },
    )
    _scope_contracts_for_word_publication(
        context,
        "consultant_evaluation_step_1",
    )
    apply_custom_mappings(
        context,
        [("ds_hop_dong", "hop_dong_list", "")],
    )

    output = generate_report_from_custom_template(
        str(template_path),
        context,
        {
            "version": MANIFEST_VERSION,
            "document_type": "evaluation",
            "root_keys": sorted(context),
            "custom_root_keys": ["ds_hop_dong"],
            "image_fields": {},
            "media_organization_id": "org-a",
        },
    )

    text = _all_document_text(Document(BytesIO(output.getvalue())))
    assert "Tư vấn lập HSMT" in text
    assert "Tư vấn thẩm định HSMT" not in text


def test_datetime_mapping_alias_keeps_hours_and_minutes(tmp_path):
    template_path = tmp_path / "datetime-alias.docx"
    template = Document()
    template.add_paragraph("Thời gian đăng tải: {tg_dang_tai_kh}")
    template.add_paragraph("Ngày đăng tải ngắn: {S_tg_dang_tai_kh}")
    template.save(template_path)
    context = {"tg_dang_tai_kh": "2026-03-05 14:55:00"}

    output = generate_report_from_custom_template(
        str(template_path),
        context,
        {
            "version": MANIFEST_VERSION,
            "document_type": "evaluation",
            "root_keys": ["tg_dang_tai_kh"],
            "custom_root_keys": ["tg_dang_tai_kh"],
            "datetime_root_keys": ["tg_dang_tai_kh"],
            "date_root_keys": ["tg_dang_tai_kh"],
            "money_root_keys": [],
            "image_fields": {},
            "media_organization_id": "org-a",
        },
    )

    text = _all_document_text(Document(BytesIO(output.getvalue())))
    assert "14:55 ngày 05/3/2026" in text
    assert "Ngày đăng tải ngắn: 05/3/2026" in text


@pytest.mark.parametrize(
    ("source", "expected"),
    [
        ("2026-01-05 14:55:00", "05/01/2026"),
        ("2026-02-05 14:55:00", "05/02/2026"),
        ("2026-03-05 14:55:00", "05/3/2026"),
        ("2026-04-05 14:55:00", "05/4/2026"),
        ("2026-05-05 14:55:00", "05/5/2026"),
        ("2026-06-05 14:55:00", "05/6/2026"),
        ("2026-07-05 14:55:00", "05/7/2026"),
        ("2026-08-05 14:55:00", "05/8/2026"),
        ("2026-09-05 14:55:00", "05/9/2026"),
        ("2026-10-05 14:55:00", "05/10/2026"),
        ("2026-11-05 14:55:00", "05/11/2026"),
        ("2026-12-05 14:55:00", "05/12/2026"),
    ],
)
def test_short_date_uses_leading_zero_only_for_january_and_february(
    tmp_path,
    source,
    expected,
):
    template_path = tmp_path / "short-date.docx"
    template = Document()
    template.add_paragraph("Ngày ngắn: {S_tg_dang_tai_kh}")
    template.save(template_path)
    context = {"tg_dang_tai_kh": source}

    output = generate_report_from_custom_template(
        str(template_path),
        context,
        {
            "version": MANIFEST_VERSION,
            "document_type": "evaluation",
            "root_keys": sorted(context),
            "custom_root_keys": sorted(context),
            "datetime_root_keys": sorted(context),
            "date_root_keys": sorted(context),
            "money_root_keys": [],
            "image_fields": {},
            "media_organization_id": "org-a",
        },
    )

    text = _all_document_text(Document(BytesIO(output.getvalue())))
    assert f"Ngày ngắn: {expected}" in text


@pytest.mark.parametrize(
    ("source", "expected"),
    [
        ("2026-01-05", "05/01/2026"),
        ("2026-02-05", "05/02/2026"),
        ("2026-03-05", "05/3/2026"),
        ("2026-04-05", "05/4/2026"),
        ("2026-05-05", "05/5/2026"),
        ("2026-06-05", "05/6/2026"),
        ("2026-07-05", "05/7/2026"),
        ("2026-08-05", "05/8/2026"),
        ("2026-09-05", "05/9/2026"),
        ("2026-10-05", "05/10/2026"),
        ("2026-11-05", "05/11/2026"),
        ("2026-12-05", "05/12/2026"),
    ],
)
@pytest.mark.parametrize(
    "formula",
    [
        "formatDate(source_date)",
        'formatDate(source_date, "dd/MM/yyyy")',
    ],
)
def test_computed_word_date_uses_the_same_approved_short_date_rule(
    source,
    expected,
    formula,
):
    context = {"source_date": source}

    apply_computed_mappings(
        context,
        [("computed_short_date", "__computed__", formula)],
    )

    assert context["computed_short_date"] == expected


def test_money_mapping_alias_supports_amount_in_words_for_numbers_and_numeric_strings(
    tmp_path,
):
    template_path = tmp_path / "money-in-words.docx"
    template = Document()
    template.add_paragraph("Giá gói thầu: {bangchu_gia_gt}")
    template.add_paragraph("Giá hợp đồng: {BangChu_hd_gia_tri}")
    template.save(template_path)
    context = {
        "gia_gt": 12000000,
        "hd_gia_tri": "9000000",
    }

    output = generate_report_from_custom_template(
        str(template_path),
        context,
        {
            "version": MANIFEST_VERSION,
            "document_type": "evaluation",
            "root_keys": sorted(context),
            "custom_root_keys": sorted(context),
            "datetime_root_keys": [],
            "date_root_keys": [],
            "money_root_keys": sorted(context),
            "image_fields": {},
            "media_organization_id": "org-a",
        },
    )

    text = _all_document_text(Document(BytesIO(output.getvalue())))
    assert "Giá gói thầu: Mười hai triệu đồng" in text
    assert "Giá hợp đồng: Chín triệu đồng" in text


def test_real_mapping_pipeline_keeps_money_numeric_until_amount_words_are_derived(
    tmp_path,
):
    template_path = tmp_path / "mapped-money-in-words.docx"
    template = Document()
    template.add_paragraph("Giá trị số: {gia_gt}")
    template.add_paragraph("Bằng chữ: {bangchu_gia_gt}")
    template.save(template_path)
    context = {"goi_thau": {"id": "package-1", "gia_goi_thau": 898000000}}
    mappings = [("gia_gt", "goi_thau", "gia_goi_thau")]

    apply_custom_mappings(context, mappings)
    context, manifest = seal_docx_context(
        "evaluation",
        context,
        mappings,
        organization_id="org-a",
    )
    output = generate_report_from_custom_template(
        str(template_path),
        context,
        manifest,
    )

    text = _all_document_text(Document(BytesIO(output.getvalue())))
    assert "Giá trị số: 898.000.000" in text
    assert "Bằng chữ: Tám trăm chín mươi tám triệu đồng" in text


def test_sample_template_renders_money_words_and_approved_short_date_from_real_mappings():
    template_path = (
        Path(__file__).resolve().parents[1]
        / "docs"
        / "BiddingFlow_Mau_Kiem_Thu_Xuat_Word.docx"
    )
    context = {
        "goi_thau": {"id": "package-1", "gia_goi_thau": 898000000},
        "ke_hoach": {"thoi_gian_dang_tai": "2026-03-05 14:55:00"},
    }
    mappings = [
        ("gia_gt", "goi_thau", "gia_goi_thau"),
        ("tg_dang_tai_kh", "ke_hoach_lcnt", "thoi_gian_dang_tai"),
    ]

    apply_custom_mappings(context, mappings)
    context, manifest = seal_docx_context(
        "evaluation",
        context,
        mappings,
        organization_id="org-a",
    )
    output = generate_report_from_custom_template(
        str(template_path),
        context,
        manifest,
    )

    text = _all_document_text(Document(BytesIO(output.getvalue())))
    assert "898.000.000" in text
    assert "Tám trăm chín mươi tám triệu đồng" in text
    assert "14:55 ngày 05/3/2026" in text
    assert "05/3/2026" in text


def test_template_skips_unavailable_list_when_assigned_to_plan_context(tmp_path):
    template_path = tmp_path / "cross-context.docx"
    template = Document()
    template.add_paragraph("Kế hoạch: {ma_kh}")
    template.add_paragraph("Hợp đồng ngoài ngữ cảnh: {ten_hd}")
    table = template.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Nhà thầu"
    table.cell(1, 0).text = "{#ds_mo_thau}{mt_ten_nt}{/ds_mo_thau}"
    template.save(template_path)

    output = generate_report_from_custom_template(
        str(template_path),
        {
            "ma_kh": "KH-01",
            "ke_hoach": {"ten_hd": "NESTED VALUE MUST NOT LEAK"},
        },
        {
            "version": MANIFEST_VERSION,
            "document_type": "plan",
            "root_keys": ["ke_hoach", "ma_kh"],
            "custom_root_keys": ["ma_kh"],
            "image_fields": {},
            "media_organization_id": "org-a",
        },
    )

    rendered = Document(BytesIO(output.getvalue()))
    assert "Kế hoạch: KH-01" in "\n".join(
        paragraph.text for paragraph in rendered.paragraphs
    )
    assert len(rendered.tables[0].rows) == 1
    assert "{ten_hd}" not in "\n".join(
        paragraph.text for paragraph in rendered.paragraphs
    )
    assert "NESTED VALUE MUST NOT LEAK" not in "\n".join(
        paragraph.text for paragraph in rendered.paragraphs
    )
    assert "ds_mo_thau" not in "\n".join(
        cell.text
        for table in rendered.tables
        for row in table.rows
        for cell in row.cells
    )


def test_cross_context_fallback_does_not_allow_jinja_expression_roots(tmp_path):
    template_path = tmp_path / "unsafe-expression.docx"
    template = Document()
    template.add_paragraph("{{ secret_outside_manifest }}")
    template.save(template_path)

    with pytest.raises(TemplateRenderError):
        generate_report_from_custom_template(
            str(template_path),
            {"ma_kh": "KH-01"},
            {
                "version": MANIFEST_VERSION,
                "document_type": "plan",
                "root_keys": ["ma_kh"],
                "custom_root_keys": ["ma_kh"],
                "image_fields": {},
                "media_organization_id": "org-a",
            },
        )
