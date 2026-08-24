from __future__ import annotations

from io import BytesIO
import json
import re
from zipfile import ZIP_DEFLATED, ZIP_STORED, ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree
from PIL import Image
import pytest

import backend.documents.document_worker as document_worker
import backend.documents.word_standardizer.automatic as automatic_policy
from backend.documents.document_worker import run_document_job
from backend.documents.document_ipc import WORKER_EVENT_PREFIX
from backend.documents.docx_context_policy import MANIFEST_VERSION
from backend.documents.word_standardizer import (
    process_docx,
    standardize_template_for_export,
)


_WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
_PLACEHOLDER = re.compile(r"(?:\{\{.*?\}\}|\{%.*?%\})")
_NUMBER = re.compile(r"(?<!\w)\d(?:[\d.,:/-]*\d)?(?!\w)")


def _seal_png() -> bytes:
    output = BytesIO()
    Image.new("RGB", (12, 12), color=(176, 20, 35)).save(
        output,
        format="PNG",
    )
    return output.getvalue()


def _administrative_template(*, conflict: bool = False) -> bytes:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x33, 0x44, 0x55)

    national_header = document.add_paragraph()
    header_run = national_header.add_run(
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    )
    header_run.font.name = "Arial"
    header_run.font.color.rgb = RGBColor(0x66, 0x22, 0x22)
    document.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    document.add_paragraph("Số: {{ document.number }}/TB-BF-2026")
    document.add_paragraph("THÔNG BÁO")
    if conflict:
        document.add_paragraph("HỢP ĐỒNG")
    document.add_paragraph("Về việc {{ document.subject }}")
    document.add_paragraph("Kính gửi: Đơn vị liên quan")
    document.add_paragraph(
        "Nội dung pháp lý phải giữ nguyên: giá trị 12.345,67 đồng; "
        "ngày 24/08/2026."
    )

    table = document.add_table(rows=1, cols=1)
    table_paragraph = table.cell(0, 0).paragraphs[0]
    table_paragraph.clear()
    table_run = table_paragraph.add_run(
        "Bảng số liệu 987.654 và {{ row.amount }}"
    )
    table_run.bold = True
    table_run.font.name = "Calibri"
    table_run.font.color.rgb = RGBColor(0x00, 0x55, 0xAA)

    signing = document.add_paragraph()
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), "17")
    bookmark_start.set(qn("w:name"), "signature_and_seal")
    signing._p.append(bookmark_start)
    signing.add_run("KT. GIÁM ĐỐC\nNguyễn Văn A")
    signing.add_run().add_picture(BytesIO(_seal_png()), width=Inches(0.22))
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "17")
    signing._p.append(bookmark_end)
    document.add_paragraph("Nơi nhận:")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _unknown_template() -> bytes:
    document = Document()
    document.add_paragraph("Biểu mẫu nghiệp vụ nội bộ 2026")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _strict_administrative_template() -> bytes:
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    for value in (
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "Độc lập - Tự do - Hạnh phúc",
        "NGHỊ ĐỊNH",
        "Nội dung pháp lý 214/2026/NĐ-CP được giữ nguyên.",
        "KT. BỘ TRƯỞNG",
        "Nguyễn Văn A",
    ):
        document.add_paragraph(value)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _dynamic_type_with_legal_citation() -> bytes:
    document = Document()
    for value in (
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "Độc lập - Tự do - Hạnh phúc",
        "{{ document.type }}",
        "Căn cứ Nghị định số 30/2020/NĐ-CP ngày 05/03/2020.",
    ):
        document.add_paragraph(value)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _with_zip_entry(content: bytes, name: str, value: bytes) -> bytes:
    output = BytesIO()
    with ZipFile(BytesIO(content)) as source, ZipFile(
        output,
        "w",
        ZIP_DEFLATED,
    ) as target:
        for info in source.infolist():
            target.writestr(info, source.read(info.filename))
        target.writestr(name, value)
    return output.getvalue()


def _with_alternate_word_namespace_prefix(content: bytes) -> bytes:
    output = BytesIO()
    with ZipFile(BytesIO(content)) as source, ZipFile(
        output,
        "w",
        ZIP_DEFLATED,
    ) as target:
        for info in source.infolist():
            value = source.read(info.filename)
            if info.filename == "word/document.xml":
                value = value.replace(b"xmlns:w=", b"xmlns:x=").replace(
                    b"w:", b"x:"
                )
            target.writestr(info, value)
    return output.getvalue()


def _word_xml_roots(content: bytes):
    with ZipFile(BytesIO(content)) as archive:
        for name in sorted(archive.namelist()):
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            try:
                yield name, etree.fromstring(archive.read(name))
            except etree.XMLSyntaxError:
                continue


def _text_records(content: bytes) -> tuple[tuple[str, str, str], ...]:
    records = []
    for part, root in _word_xml_roots(content):
        for node in root.xpath(".//w:t | .//w:instrText", namespaces=_WORD_NS):
            records.append(
                (part, etree.QName(node).localname, str(node.text or ""))
            )
    return tuple(records)


def _numeric_tokens(content: bytes) -> tuple[str, ...]:
    return tuple(
        match.group(0)
        for _part, _kind, value in _text_records(content)
        for match in _NUMBER.finditer(value)
    )


def _placeholder_tokens(content: bytes) -> tuple[str, ...]:
    return tuple(
        match.group(0)
        for _part, _kind, value in _text_records(content)
        for match in _PLACEHOLDER.finditer(value)
    )


def _canonical_tables(content: bytes) -> tuple[bytes, ...]:
    root = dict(_word_xml_roots(content))["word/document.xml"]
    return tuple(
        etree.tostring(table, method="c14n")
        for table in root.xpath(".//w:tbl", namespaces=_WORD_NS)
    )


def _canonical_signature_blocks(content: bytes) -> tuple[bytes, ...]:
    root = dict(_word_xml_roots(content))["word/document.xml"]
    blocks = []
    for paragraph in root.xpath(".//w:p", namespaces=_WORD_NS):
        text = "".join(
            paragraph.xpath(".//w:t/text()", namespaces=_WORD_NS)
        )
        if "KT. GIÁM ĐỐC" in text:
            blocks.append(etree.tostring(paragraph, method="c14n"))
    return tuple(blocks)


def _canonical_paragraph(content: bytes, expected_text: str) -> bytes:
    root = dict(_word_xml_roots(content))["word/document.xml"]
    for paragraph in root.xpath(".//w:p", namespaces=_WORD_NS):
        text = "".join(
            paragraph.xpath(".//w:t/text()", namespaces=_WORD_NS)
        )
        if text == expected_text:
            return etree.tostring(paragraph, method="c14n")
    raise AssertionError(f"Paragraph not found: {expected_text}")


def _media(content: bytes) -> dict[str, bytes]:
    with ZipFile(BytesIO(content)) as archive:
        return {
            name: archive.read(name)
            for name in sorted(archive.namelist())
            if name.startswith("word/media/")
        }


def _empty_manifest(document_type: str = "plan") -> dict:
    return {
        "version": MANIFEST_VERSION,
        "document_type": document_type,
        "root_keys": [],
        "custom_root_keys": [],
        "datetime_root_keys": [],
        "date_root_keys": [],
        "money_root_keys": [],
        "image_fields": {},
        "media_organization_id": "org-a",
    }


def _renderable_template() -> bytes:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x33, 0x44, 0x55)

    first = document.add_paragraph()
    run = first.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x66, 0x22, 0x22)
    document.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    document.add_paragraph("THÔNG BÁO")
    document.add_paragraph("Về việc kiểm tra tự động")
    document.add_paragraph("Kính gửi: Đơn vị liên quan")
    document.add_paragraph("Nội dung giữ nguyên 123.456.")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_batch_render_uses_one_worker_result_and_does_not_recompress_docx(
    tmp_path, monkeypatch,
):
    monkeypatch.setenv("WORD_EXPORT_STANDARDIZATION_MODE", "off")
    monkeypatch.setenv("BIDDING_WORD_EXPORT_CACHE_DIR", str(tmp_path / "cache"))
    first = Document()
    first.add_paragraph("Tài liệu thứ nhất")
    first_stream = BytesIO()
    first.save(first_stream)
    second = Document()
    second.add_paragraph("Tài liệu thứ hai")
    second_stream = BytesIO()
    second.save(second_stream)

    result = run_document_job(
        "render_docx_batch",
        {
            "templates": [
                {"template_content": first_stream.getvalue(), "filename": "A.docx"},
                {"template_content": second_stream.getvalue(), "filename": "B.docx"},
            ],
            "context": {},
            "context_manifest": _empty_manifest("evaluation"),
        },
    )

    with ZipFile(BytesIO(result)) as archive:
        assert [item.filename for item in archive.infolist()] == ["A.docx", "B.docx"]
        assert {item.compress_type for item in archive.infolist()} == {ZIP_STORED}
        assert Document(BytesIO(archive.read("A.docx"))).paragraphs[0].text == "Tài liệu thứ nhất"
        assert Document(BytesIO(archive.read("B.docx"))).paragraphs[0].text == "Tài liệu thứ hai"


def test_batch_cache_misses_are_standardized_and_rendered_in_one_worker(
    tmp_path, monkeypatch,
):
    """A cold batch must pay the sandbox startup cost exactly once."""

    monkeypatch.setenv("WORD_EXPORT_STANDARDIZATION_MODE", "off")
    monkeypatch.setenv("WORD_EXPORT_CACHE_ENABLED", "false")
    monkeypatch.setenv("DOCUMENT_WORKER_TEMP_DIR", str(tmp_path / "jobs"))
    original_popen = document_worker.subprocess.Popen
    process_count = 0

    def counting_popen(*args, **kwargs):
        nonlocal process_count
        process_count += 1
        return original_popen(*args, **kwargs)

    monkeypatch.setattr(document_worker.subprocess, "Popen", counting_popen)
    template = _renderable_template()

    result = run_document_job(
        "render_docx_batch",
        {
            "templates": [
                {"template_content": template, "filename": "A.docx"},
                {"template_content": template, "filename": "B.docx"},
            ],
            "context": {},
            "context_manifest": _empty_manifest("evaluation"),
        },
    )

    assert process_count == 1
    with ZipFile(BytesIO(result)) as archive:
        assert [item.filename for item in archive.infolist()] == ["A.docx", "B.docx"]


def _first_run_format(content: bytes) -> tuple[str | None, str | None]:
    document = Document(BytesIO(content))
    run = document.paragraphs[0].runs[0]
    color = str(run.font.color.rgb) if run.font.color.rgb is not None else None
    return run.font.name, color


def test_default_policy_applies_without_a_user_profile_or_action(monkeypatch):
    monkeypatch.delenv("WORD_EXPORT_STANDARDIZATION_MODE", raising=False)
    source = _administrative_template()

    result = standardize_template_for_export(source)

    assert result.content != source
    assert result.metadata["mode"] == "apply_safe"
    assert result.metadata["status"] == "APPLIED"
    assert result.metadata["effectiveProfile"] == "sector_template"
    assert result.metadata["preservation"] == "PASS"
    assert result.metadata["plannedTargetCount"] > 0


def test_high_confidence_administrative_content_selects_strict_automatically():
    source = _strict_administrative_template()

    result = standardize_template_for_export(source)

    assert result.metadata["status"] == "APPLIED"
    assert result.metadata["effectiveProfile"] == "n30_strict"
    assert "STRICT_TYPE_ALLOWLIST" in result.metadata["reasonCodes"]


def test_nghi_dinh_title_is_a_component_and_never_receives_body_indent():
    source = _strict_administrative_template()
    preview = process_docx(
        source,
        profile="n30_strict",
        mode="preview_fix",
    ).report

    title_field = next(
        field
        for field in preview["fields"]
        if field["semantic"] == "document.type"
    )
    body_targets = {
        target["paragraph"]
        for change in preview["plannedChanges"]
        if change["ruleId"].startswith("N30-BODY-")
        for target in change["targets"]
    }
    result = standardize_template_for_export(source)
    title = etree.fromstring(_canonical_paragraph(result.content, "NGHỊ ĐỊNH"))

    assert title_field["location"]["index"] == 2
    assert 2 not in body_targets
    assert title.find("w:pPr/w:ind", namespaces=_WORD_NS) is None
    assert title.find("w:pPr/w:jc", namespaces=_WORD_NS).get(qn("w:val")) == "center"


def test_conservative_export_context_never_promotes_to_strict():
    result = standardize_template_for_export(
        _strict_administrative_template(),
        document_type_hint="contract",
    )

    assert result.metadata["effectiveProfile"] == "sector_template"
    assert "CONSERVATIVE_SECTOR_RULES" in result.metadata["reasonCodes"]


def test_legal_citation_cannot_be_mistaken_for_a_document_type_title():
    source = _dynamic_type_with_legal_citation()

    result = standardize_template_for_export(source)

    assert result.content == source
    assert result.metadata["status"] == "BYPASSED_LOW_CONFIDENCE"
    assert result.metadata["detectedDocumentType"] == "unknown"
    assert result.metadata["effectiveProfile"] == "reference_only"


@pytest.mark.parametrize(
    ("mode", "expected_status"),
    (("off", "OFF"), ("shadow", "SHADOW")),
)
def test_non_mutating_modes_return_the_exact_source_bytes(mode, expected_status):
    source = _administrative_template()

    result = standardize_template_for_export(source, mode=mode)

    assert result.content == source
    assert result.metadata["mode"] == mode
    assert result.metadata["status"] == expected_status
    assert result.metadata["sourceSha256"] == result.metadata["outputSha256"]


def test_unknown_content_is_returned_unchanged_for_manual_reference():
    source = _unknown_template()

    result = standardize_template_for_export(source)

    assert result.content == source
    assert result.metadata["status"] == "BYPASSED_LOW_CONFIDENCE"
    assert result.metadata["effectiveProfile"] == "reference_only"
    assert result.metadata["detectedDocumentType"] == "unknown"


def test_conflicting_document_types_are_returned_unchanged():
    source = _administrative_template(conflict=True)

    result = standardize_template_for_export(source)

    assert result.content == source
    assert result.metadata["status"] == "BYPASSED_CONFLICT"
    assert result.metadata["effectiveProfile"] == "reference_only"
    assert result.metadata["documentTypeConflictCount"] > 0


def test_signed_opc_package_is_returned_byte_for_byte_unchanged():
    source = _with_zip_entry(
        _administrative_template(),
        "_xmlsignatures/sig1.xml",
        b"<Signature xmlns='urn:test'/>",
    )

    result = standardize_template_for_export(source)

    assert result.content == source
    assert result.metadata["status"] == "SKIPPED_SIGNED"
    assert result.metadata["effectiveProfile"] == "reference_only"
    assert result.metadata["reasonCodes"] == ["OPC_SIGNATURE"]


def test_applied_formatting_exactly_preserves_protected_content():
    source = _administrative_template()
    source_text = _text_records(source)
    source_numbers = _numeric_tokens(source)
    source_placeholders = _placeholder_tokens(source)
    source_tables = _canonical_tables(source)
    source_signatures = _canonical_signature_blocks(source)
    source_media = _media(source)

    result = standardize_template_for_export(source)

    assert result.metadata["status"] == "APPLIED"
    assert source_placeholders == (
        "{{ document.number }}",
        "{{ document.subject }}",
        "{{ row.amount }}",
    )
    assert source_tables and source_signatures and source_media
    assert _text_records(result.content) == source_text
    assert _numeric_tokens(result.content) == source_numbers
    assert _placeholder_tokens(result.content) == source_placeholders
    assert _canonical_tables(result.content) == source_tables
    assert _canonical_signature_blocks(result.content) == source_signatures
    assert _media(result.content) == source_media


def test_automatic_formatting_is_idempotent():
    first = standardize_template_for_export(_strict_administrative_template())
    second = standardize_template_for_export(first.content)

    assert first.metadata["status"] == "APPLIED"
    assert second.metadata["status"] == "NO_CHANGE"
    assert second.content == first.content


def test_standardizer_failure_falls_back_without_exposing_error_text(monkeypatch):
    source = _administrative_template()

    def fail_processing(*_args, **_kwargs):
        raise RuntimeError("secret record value must not escape")

    monkeypatch.setattr(automatic_policy, "process_docx", fail_processing)
    result = automatic_policy.standardize_template_for_export(source)

    assert result.content == source
    assert result.metadata["status"] == "BYPASSED_ERROR"
    assert result.metadata["errorType"] == "RuntimeError"
    assert "secret" not in str(result.metadata)


def test_oversized_paragraph_inventory_bypasses_before_repeated_analysis():
    document = Document()
    for index in range(automatic_policy.MAX_AUTOMATIC_PARAGRAPHS + 1):
        document.add_paragraph(f"Đoạn {index}")
    output = BytesIO()
    document.save(output)
    source = output.getvalue()

    result = standardize_template_for_export(source)

    assert result.content == source
    assert result.metadata["status"] == "BYPASSED_COMPLEXITY"
    assert result.metadata["paragraphCount"] > automatic_policy.MAX_AUTOMATIC_PARAGRAPHS
    assert result.metadata["reasonCodes"] == ["COMPLEXITY_BUDGET"]


def test_complexity_gate_counts_word_elements_by_namespace_not_prefix():
    document = Document()
    document.add_paragraph("Một")
    document.add_paragraph("Hai")
    output = BytesIO()
    document.save(output)
    alternate_prefix = _with_alternate_word_namespace_prefix(output.getvalue())

    complexity = automatic_policy._automatic_complexity(alternate_prefix)

    assert complexity["paragraphCount"] == 2
    assert complexity["runCount"] == 2
    assert complexity["exceeded"] is False


def test_tracked_formatting_change_paragraph_is_exactly_immutable():
    document = Document(BytesIO(_strict_administrative_template()))
    tracked = document.paragraphs[3]
    ppr_change = OxmlElement("w:pPrChange")
    ppr_change.set(qn("w:id"), "42")
    ppr_change.set(qn("w:author"), "Reviewer")
    previous = OxmlElement("w:pPr")
    ppr_change.append(previous)
    tracked._p.get_or_add_pPr().append(ppr_change)
    output = BytesIO()
    document.save(output)
    source = output.getvalue()
    expected = _canonical_paragraph(
        source,
        "Nội dung pháp lý 214/2026/NĐ-CP được giữ nguyên.",
    )

    result = standardize_template_for_export(source)

    assert result.metadata["status"] == "APPLIED"
    assert _canonical_paragraph(
        result.content,
        "Nội dung pháp lý 214/2026/NĐ-CP được giữ nguyên.",
    ) == expected


def test_plain_signature_title_and_signer_name_are_exactly_immutable():
    document = Document(BytesIO(_strict_administrative_template()))
    document.paragraphs[-2].text = "GIÁM ĐỐC"
    for paragraph in document.paragraphs[-2:]:
        paragraph.alignment = 1
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)
    output = BytesIO()
    document.save(output)
    source = output.getvalue()
    title = _canonical_paragraph(source, "GIÁM ĐỐC")
    signer = _canonical_paragraph(source, "Nguyễn Văn A")

    result = standardize_template_for_export(source)

    assert result.metadata["status"] == "APPLIED"
    assert _canonical_paragraph(result.content, "GIÁM ĐỐC") == title
    assert _canonical_paragraph(result.content, "Nguyễn Văn A") == signer


def test_isolated_worker_renders_the_automatically_standardized_template(
    monkeypatch,
):
    monkeypatch.delenv("WORD_EXPORT_STANDARDIZATION_MODE", raising=False)
    events = []
    monkeypatch.setattr(
        document_worker,
        "log_structured_event",
        lambda event, **kwargs: events.append((event, kwargs)),
    )
    source = _renderable_template()

    output = run_document_job(
        "render_docx",
        {
            "template_content": source,
            "context": {},
            "context_manifest": _empty_manifest(),
        },
        timeout_seconds=30,
    )

    assert _first_run_format(source) == ("Arial", "662222")
    assert _first_run_format(output) == ("Times New Roman", "000000")
    assert len(events) == 1
    event, event_options = events[0]
    assert event == "document.word_standardization"
    assert event_options["nonblocking"] is True
    assert event_options["fields"]["mode"] == "apply_safe"
    assert event_options["fields"]["status"] == "APPLIED"
    assert event_options["fields"]["preservation"] == "PASS"


def test_worker_event_sink_failure_cannot_fail_an_export(monkeypatch):
    def unavailable_sink(*_args, **_kwargs):
        raise OSError("log sink unavailable")

    monkeypatch.setattr(
        document_worker,
        "log_structured_event",
        unavailable_sink,
    )
    event = {
        "event": "document.word_standardization",
        "fields": {"mode": "shadow", "status": "SHADOW"},
    }

    document_worker._record_worker_events(
        (WORKER_EVENT_PREFIX + json.dumps(event)).encode("ascii")
    )


@pytest.mark.parametrize("mode", ("off", "shadow"))
def test_isolated_worker_rollback_modes_render_the_original_template(
    monkeypatch,
    mode,
):
    monkeypatch.setenv("WORD_EXPORT_STANDARDIZATION_MODE", mode)
    source = _renderable_template()

    output = run_document_job(
        "render_docx",
        {
            "template_content": source,
            "context": {},
            "context_manifest": _empty_manifest(),
        },
        timeout_seconds=30,
    )

    assert _first_run_format(output) == _first_run_format(source)
