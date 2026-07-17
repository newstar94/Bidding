from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from backend.documents.timeline_context_service import (
    TIMELINE_SECTIONS,
    _default_items,
    _number_section_items,
    _timeline_item_is_applicable,
)
from backend.documents.timeline_document_service import render_timeline_document


def _sample_context():
    items = _default_items()
    for item in items:
        item["display_date"] = "15/08/2026"
        item["is_planned_date"] = True
    return {
        "goi_thau": {"ma_goi_thau": "GT-01", "ten_goi_thau": "Gói kiểm thử Unicode"},
        "ke_hoach": {"ma_ke_hoach": "KH-01"},
        "to_chuc": {"ten_to_chuc": "Đơn vị kiểm thử"},
        "timeline_sections": [
            {
                "code": code,
                "title": title,
                "items": [item for item in items if item["ma_nhom"] == code],
            }
            for code, title, _definitions in TIMELINE_SECTIONS
        ],
        "planned_date_note": "Ngày màu đỏ là ngày dự kiến/chưa xác nhận.",
        "generated_date": "17/07/2026",
    }


def test_timeline_document_has_landscape_table_and_all_milestones():
    stream = render_timeline_document(
        "data/templates/words/mau_timeline_goi_thau.docx",
        _sample_context(),
    )
    document = Document(stream)
    assert document.sections[0].orientation == WD_ORIENT.LANDSCAPE
    assert len(document.tables) == 1
    assert len(document.tables[0].rows) == 1 + 5 + 48
    assert document.tables[0].rows[0]._tr.xpath("./w:trPr/w:tblHeader")
    text = "\n".join(cell.text for row in document.tables[0].rows for cell in row.cells)
    assert "Gói kiểm thử Unicode" not in text  # Metadata belongs above the checklist.
    assert "Chứng thư thẩm định giá, Báo giá" in text
    assert "BB hoàn thiện hợp đồng" in text
    assert "[[" not in "\n".join(paragraph.text for paragraph in document.paragraphs)
    planned_runs = [
        run
        for row in document.tables[0].rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text == "15/08/2026"
    ]
    assert planned_runs
    assert all(run.font.color.rgb is not None for run in planned_runs)
    table_width = document.tables[0]._tbl.tblPr.first_child_found_in("w:tblW")
    assert table_width.get(qn("w:type")) == "dxa"


def test_timeline_applicability_follows_plan_and_package_business_data():
    combined_plan = {"phe_duyet": "Dự toán và kế hoạch"}
    assert not _timeline_item_is_applicable({"ma_moc": "1.3"}, {}, combined_plan)
    assert _timeline_item_is_applicable({"ma_moc": "1.7"}, {}, combined_plan)

    competitive_package = {
        "hinh_thuc_lua_chon": "Chào hàng cạnh tranh",
        "phuong_thuc_lua_chon": "Một giai đoạn một túi hồ sơ",
        "yeu_cau_tham_dinh_hsmt": "Không",
    }
    assert not _timeline_item_is_applicable({"ma_moc": "3.1"}, competitive_package, {})
    assert not _timeline_item_is_applicable({"ma_moc": "4.2"}, competitive_package, {})
    assert not _timeline_item_is_applicable({"ma_moc": "5.10"}, competitive_package, {})
    assert not _timeline_item_is_applicable({"ma_moc": "5.5"}, competitive_package, {})
    assert _timeline_item_is_applicable({"ma_moc": "5.2"}, competitive_package, {})

    numbered = _number_section_items([
        {"ma_moc": "1.1"},
        {"ma_moc": "1.2"},
        {"ma_moc": "1.7"},
        {"ma_moc": "1.8"},
    ])
    assert [item["display_code"] for item in numbered] == ["1.1", "1.2", "1.3", "1.4"]
