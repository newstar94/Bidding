import io
import zipfile

from docx import Document
from lxml import etree

from backend.documents import custom_exporter
from backend.documents.docx_column_loop import expand_column_loops


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": WORD_NS}


def _docx_bytes(*, markers=True, split_markers=False, table_count=1):
    document = Document()
    for table_number in range(table_count):
        table = document.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "STT"
        table.cell(0, 1).text = f"Nội dung {table_number + 1}"
        table.cell(0, 2).text = "Nhà thầu"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "Tên"
        table.cell(1, 2).text = "Tên: {name}"
        table.cell(2, 0).text = "2"
        table.cell(2, 1).text = "Giá"
        table.cell(2, 2).text = "Giá: {price} / {global_note}"

        if markers:
            if split_markers:
                paragraph = table.cell(0, 2).paragraphs[0]
                paragraph.clear()
                paragraph.add_run("{#col bid")
                paragraph.add_run("ders}Nhà thầu")
                paragraph = table.cell(2, 2).paragraphs[0]
                paragraph.add_run("{/col bid")
                paragraph.add_run("ders}")
            else:
                table.cell(0, 2).text = "{#col bidders}Nhà thầu"
                table.cell(2, 2).paragraphs[0].add_run("{/col bidders}")
        if table_number + 1 < table_count:
            document.add_paragraph()

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _document_root(content):
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        return etree.fromstring(archive.read("word/document.xml"))


def _tables(content):
    return _document_root(content).xpath(".//w:tbl", namespaces=NS)


def _cells(row):
    return row.xpath("./w:tc", namespaces=NS)


def _cell_text(cell):
    return "".join(cell.xpath(".//w:t/text()", namespaces=NS))


def _expanded(content=None, context=None):
    source = content or _docx_bytes()
    data = context if context is not None else {
        "bidders": [
            {"name": "Công ty A", "price": "100"},
            {"name": "Công ty B", "price": "200"},
            {"name": "Liên danh C", "price": "300"},
        ],
        "global_note": "VNĐ",
    }
    return expand_column_loops(source, data)


def test_no_markers_returns_unchanged():
    source = _docx_bytes(markers=False)
    assert expand_column_loops(source, {}) is source


def test_expand_three_items_and_merge_header():
    table = _tables(_expanded())[0]
    rows = table.xpath("./w:tr", namespaces=NS)
    assert len(_cells(rows[0])) == 3
    assert len(_cells(rows[1])) == 5
    header = _cells(rows[0])[2]
    span = header.xpath("./w:tcPr/w:gridSpan/@w:val", namespaces=NS)
    assert span == ["3"]
    assert [_cell_text(cell) for cell in _cells(rows[1])[2:]] == [
        "Tên: Công ty A",
        "Tên: Công ty B",
        "Tên: Liên danh C",
    ]


def test_empty_list_keeps_one_blank_dynamic_column():
    result = _expanded(context={"bidders": []})
    table = _tables(result)[0]
    rows = table.xpath("./w:tr", namespaces=NS)
    assert len(_cells(rows[1])) == 3
    assert _cell_text(_cells(rows[1])[2]) == "Tên: "
    assert _cell_text(_cells(rows[2])[2]) == "Giá:  / "


def test_single_item_does_not_merge_header():
    result = _expanded(context={"bidders": [{"name": "A", "price": 10}]})
    table = _tables(result)[0]
    rows = table.xpath("./w:tr", namespaces=NS)
    assert not _cells(rows[0])[2].xpath("./w:tcPr/w:gridSpan", namespaces=NS)
    assert _cell_text(_cells(rows[1])[2]) == "Tên: A"


def test_static_text_is_preserved_for_every_column():
    rows = _tables(_expanded())[0].xpath("./w:tr", namespaces=NS)
    assert all(
        _cell_text(cell).startswith("Giá: ")
        for cell in _cells(rows[2])[2:]
    )


def test_markers_split_across_runs_are_removed():
    result = _expanded(content=_docx_bytes(split_markers=True))
    xml_text = etree.tostring(_document_root(result), encoding="unicode")
    assert "#col" not in xml_text
    assert "/col" not in xml_text
    assert "Công ty B" in xml_text


def test_table_grid_is_replaced_with_dynamic_columns():
    table = _tables(_expanded())[0]
    columns = table.xpath("./w:tblGrid/w:gridCol", namespaces=NS)
    assert len(columns) == 5
    widths = [int(column.get(f"{{{WORD_NS}}}w")) for column in columns]
    assert sum(widths[2:]) > 0
    assert max(widths[2:]) - min(widths[2:]) <= 1


def test_missing_list_behaves_like_empty_list():
    rows = _tables(_expanded(context={}))[0].xpath("./w:tr", namespaces=NS)
    assert len(_cells(rows[1])) == 3
    assert _cell_text(_cells(rows[1])[2]) == "Tên: "


def test_multiple_tables_are_expanded_independently():
    result = _expanded(content=_docx_bytes(table_count=2))
    tables = _tables(result)
    assert len(tables) == 2
    for table in tables:
        rows = table.xpath("./w:tr", namespaces=NS)
        assert len(_cells(rows[1])) == 5
        assert "Công ty A" in _cell_text(_cells(rows[1])[2])


def test_multiple_regions_in_one_table_expand_from_right_to_left():
    document = Document()
    table = document.add_table(rows=2, cols=4)
    table.cell(0, 0).text = "Cố định"
    table.cell(0, 1).text = "{#col left}Nhóm trái"
    table.cell(0, 2).text = "Phân cách"
    table.cell(0, 3).text = "{#col right}Nhóm phải"
    table.cell(1, 1).text = "{value}{/col left}"
    table.cell(1, 3).text = "{value}{/col right}"
    stream = io.BytesIO()
    document.save(stream)

    result = expand_column_loops(
        stream.getvalue(),
        {
            "left": [{"value": "L1"}, {"value": "L2"}],
            "right": [{"value": "R1"}, {"value": "R2"}, {"value": "R3"}],
        },
    )

    rows = _tables(result)[0].xpath("./w:tr", namespaces=NS)
    assert [_cell_text(cell) for cell in _cells(rows[1])] == [
        "",
        "L1",
        "L2",
        "",
        "R1",
        "R2",
        "R3",
    ]


def test_unknown_item_variable_is_preserved_for_docxtpl():
    result = _expanded(
        context={"bidders": [{"name": "A", "price": 10}]},
    )
    rows = _tables(result)[0].xpath("./w:tr", namespaces=NS)
    assert "{global_note}" in _cell_text(_cells(rows[2])[2])


def test_runtime_values_bypass_translated_template_cache(tmp_path):
    template_path = tmp_path / "columns.docx"
    template_path.write_bytes(_docx_bytes())
    custom_exporter._TRANSLATED_DOCXTPL_CACHE.clear()

    first = custom_exporter.generate_report_from_custom_template(
        template_path,
        {
            "bidders": [{"name": "Nhà thầu Một", "price": "111"}],
            "global_note": "VNĐ",
        },
    )
    second = custom_exporter.generate_report_from_custom_template(
        template_path,
        {
            "bidders": [{"name": "Nhà thầu Hai", "price": "222"}],
            "global_note": "VNĐ",
        },
    )

    assert "Nhà thầu Một" in "\n".join(
        cell.text for table in Document(first).tables for row in table.rows for cell in row.cells
    )
    assert "Nhà thầu Hai" in "\n".join(
        cell.text for table in Document(second).tables for row in table.rows for cell in row.cells
    )
    assert template_path not in custom_exporter._TRANSLATED_DOCXTPL_CACHE


def test_runtime_values_cannot_inject_template_syntax(tmp_path):
    template_path = tmp_path / "safe-columns.docx"
    template_path.write_bytes(_docx_bytes())
    supplied_name = "{global_note} {{ global_note }} {% if unsafe %}"

    rendered = custom_exporter.generate_report_from_custom_template(
        template_path,
        {
            "bidders": [{"name": supplied_name, "price": "100"}],
            "global_note": "VNĐ",
        },
    )
    text = "\n".join(
        cell.text
        for table in Document(rendered).tables
        for row in table.rows
        for cell in row.cells
    )

    assert supplied_name in text
