import asyncio
import io
import os
import subprocess
import threading
import time
import zipfile
from types import SimpleNamespace

import pytest
from docx import Document
from openpyxl import Workbook
from starlette.testclient import TestClient

from backend.app import app
from backend.documents.archive_validation import (
    UnsafeArchiveError,
    validate_ooxml_archive,
)
from backend.documents.custom_exporter import (
    TemplateRenderError,
    generate_report_from_custom_template,
)
from backend.documents.docx_context_policy import seal_docx_context
from backend.documents import document_worker
from backend.documents import routes_docx, routes_excel
from backend.documents.document_worker import (
    DocumentWorkerBusyError,
    DocumentWorkerError,
    DocumentWorkerInputError,
    DocumentWorkerTimeoutError,
    cleanup_stale_document_jobs,
    run_document_job,
)
from backend.documents.routes_docx import _validate_docx_upload
from backend.documents.routes_excel import _validate_excel_upload


def _docx_bytes(text="Hello {{ name }}"):
    output = io.BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(output)
    return output.getvalue()


def _xlsx_bytes():
    output = io.BytesIO()
    workbook = Workbook()
    workbook.active["A1"] = "safe"
    workbook.save(output)
    return output.getvalue()


def _xlsx_formula_bytes():
    output = io.BytesIO()
    workbook = Workbook()
    workbook.active["A1"] = "=WEBSERVICE(\"https://example.invalid\")"
    workbook.save(output)
    return output.getvalue()


def _with_external_link_part(content):
    source = io.BytesIO(content)
    output = io.BytesIO()
    with zipfile.ZipFile(source) as existing, zipfile.ZipFile(output, "w") as rewritten:
        for info in existing.infolist():
            rewritten.writestr(info, existing.read(info.filename))
        rewritten.writestr("xl/externalLinks/externalLink1.xml", "<externalLink/>")
    return output.getvalue()


def _minimal_content_types(main_part, content_type):
    return f"""<?xml version="1.0" encoding="UTF-8"?>
    <Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
      <Override PartName="/{main_part}" ContentType="{content_type}"/>
    </Types>""".encode()


def test_valid_docx_and_xlsx_uploads_pass_structural_validation():
    docx_data = _docx_bytes()
    generated_name = _validate_docx_upload("report.docx", docx_data)
    assert generated_name.startswith("report_")

    xlsx_data = _xlsx_bytes()
    upload = SimpleNamespace(
        filename="data.xlsx",
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    _validate_excel_upload(upload, xlsx_data)


def test_xlsx_import_rejects_formulas_and_external_link_parts():
    with pytest.raises(UnsafeArchiveError, match="công thức"):
        validate_ooxml_archive(_xlsx_formula_bytes(), "xlsx")
    with pytest.raises(UnsafeArchiveError, match="liên kết ngoài"):
        validate_ooxml_archive(_with_external_link_part(_xlsx_bytes()), "xlsx")


def test_upload_routes_require_authentication_before_parsing_files(monkeypatch):
    async def must_not_run(*_args, **_kwargs):
        raise AssertionError("Anonymous upload reached the document worker")

    monkeypatch.setattr(routes_docx, "run_document_job_async", must_not_run)
    monkeypatch.setattr(routes_excel, "run_document_job_async", must_not_run)
    with TestClient(app, base_url="https://testserver") as client:
        docx_response = client.post(
            "/api/templates/upload",
            files={"file": ("report.docx", _docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        excel_response = client.post(
            "/api/import-excel",
            data={"type": "goithau"},
            files={"file": ("data.xlsx", _xlsx_bytes(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        )

    assert docx_response.status_code == 403
    assert excel_response.status_code == 403


def test_authenticated_upload_routes_accept_valid_and_reject_invalid_files(monkeypatch):
    async def fake_document_job(operation, _payload, **_kwargs):
        return [] if operation == "parse_excel" else True

    monkeypatch.setattr(routes_docx, "run_document_job_async", fake_document_job)
    monkeypatch.setattr(routes_excel, "run_document_job_async", fake_document_job)
    monkeypatch.setattr(routes_docx, "_persist_user_template", lambda *_args: None)
    with TestClient(app, base_url="https://testserver") as client:
        login = client.post(
            "/api/auth/login",
            json={
                "username": os.environ.get("ADMIN_USERNAME", "admin"),
                "password": os.environ["ADMIN_PASSWORD"],
                "remember": False,
            },
        )
        assert login.status_code == 200
        headers = {"X-CSRF-Token": client.cookies.get("csrf_token")}

        valid_docx = client.post(
            "/api/templates/upload",
            headers=headers,
            files={"file": ("report.docx", _docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        valid_excel = client.post(
            "/api/import-excel",
            headers=headers,
            data={"type": "goithau"},
            files={"file": ("data.xlsx", _xlsx_bytes(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        )
        invalid_docx = client.post(
            "/api/templates/upload",
            headers=headers,
            files={"file": ("report.txt", b"not a document", "text/plain")},
        )
        invalid_excel = client.post(
            "/api/import-excel",
            headers=headers,
            data={"type": "goithau"},
            files={"file": ("data.exe", b"not a workbook", "application/octet-stream")},
        )

    assert valid_docx.status_code == 200
    assert valid_docx.json()["success"] is True
    assert valid_excel.status_code == 200
    assert valid_excel.json() == {"success": True, "rows": []}
    assert invalid_docx.status_code == 400
    assert invalid_excel.status_code == 400


def test_ooxml_rejects_path_traversal_entry():
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            _minimal_content_types(
                "word/document.xml",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
            ),
        )
        archive.writestr("word/document.xml", "<document/>")
        archive.writestr("../outside.xml", "<outside/>")

    with pytest.raises(UnsafeArchiveError, match="đường dẫn"):
        validate_ooxml_archive(output.getvalue(), "docx")


def test_ooxml_rejects_zip_bomb_compression_ratio():
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            _minimal_content_types(
                "word/document.xml",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
            ),
        )
        archive.writestr("word/document.xml", f"<document>{'A' * 500_000}</document>")

    with pytest.raises(UnsafeArchiveError, match="Tỷ lệ nén"):
        validate_ooxml_archive(output.getvalue(), "docx")
    with pytest.raises(DocumentWorkerInputError, match="Tỷ lệ nén"):
        run_document_job(
            "validate_docx",
            {"content": output.getvalue()},
            timeout_seconds=15,
        )


def test_docx_upload_rejects_unapproved_jinja_statement():
    data = _docx_bytes("{% include 'local-file' %}")

    with pytest.raises(ValueError, match="thẻ điều khiển"):
        _validate_docx_upload("unsafe.docx", data)


def test_sandbox_blocks_private_attribute_and_never_returns_raw_template(tmp_path):
    path = tmp_path / "unsafe.docx"
    path.write_bytes(_docx_bytes("{{ ''.__class__.__mro__ }}"))

    with pytest.raises(TemplateRenderError, match="không được hỗ trợ"):
        generate_report_from_custom_template(path, {})


def test_sandbox_still_renders_allowed_placeholders(tmp_path):
    path = tmp_path / "safe.docx"
    path.write_bytes(_docx_bytes("Hello {{ name | upper }}"))

    rendered = generate_report_from_custom_template(path, {"name": "Lan"})
    result = Document(rendered)

    assert "Hello LAN" in "\n".join(paragraph.text for paragraph in result.paragraphs)


def test_strict_undefined_rejects_missing_template_variable(tmp_path):
    path = tmp_path / "missing.docx"
    path.write_bytes(_docx_bytes("{{ missing_value }}"))

    with pytest.raises(TemplateRenderError, match="không thể kết xuất"):
        generate_report_from_custom_template(path, {})


@pytest.mark.parametrize(
    "expression",
    [
        "{{ value.upper() }}",
        "{{ value * 1000000 }}",
        "{{ values[dynamic_key] }}",
        "{% if ''.__class__ %}unsafe{% endif %}",
    ],
)
def test_template_grammar_rejects_calls_arithmetic_and_dynamic_access(expression):
    with pytest.raises(ValueError, match="không|chỉ hỗ trợ"):
        _validate_docx_upload("unsafe.docx", _docx_bytes(expression))


def test_document_worker_renders_and_exports_in_subprocess(tmp_path):
    template_path = tmp_path / "safe.docx"
    template_path.write_bytes(_docx_bytes("Hello {{ investor_name | upper }}"))
    context, context_manifest = seal_docx_context(
        "plan", {"investor_name": "Lan"}
    )

    rendered = run_document_job(
        "render_docx",
        {
            "template_path": str(template_path),
            "context": context,
            "context_manifest": context_manifest,
        },
    )
    rendered_document = Document(io.BytesIO(rendered))
    assert "Hello LAN" in "\n".join(
        paragraph.text for paragraph in rendered_document.paragraphs
    )

    excel_bytes = run_document_job(
        "export_excel",
        {"function": "create_excel_template", "args": ["goithau"]},
    )
    validate_ooxml_archive(excel_bytes, "xlsx")


def test_document_worker_rejects_non_allowlisted_operation():
    with pytest.raises(DocumentWorkerInputError, match="không được hỗ trợ"):
        run_document_job("run_command", {"command": "whoami"})


def test_document_worker_timeout_kills_process_and_cleans_job_dir(
    monkeypatch, tmp_path
):
    class TimeoutProcess:
        pid = os.getpid()
        returncode = None

        def __init__(self):
            self.killed = False
            self.calls = 0

        def communicate(self, timeout=None):
            self.calls += 1
            if self.calls == 1:
                raise subprocess.TimeoutExpired("worker", timeout)
            return b"", b""

        def kill(self):
            self.killed = True
            self.returncode = -9

    process = TimeoutProcess()
    monkeypatch.setenv("DOCUMENT_WORKER_TEMP_DIR", str(tmp_path / "worker"))
    monkeypatch.setattr(document_worker.subprocess, "Popen", lambda *_a, **_k: process)
    monkeypatch.setattr(document_worker, "_terminate_process", lambda child: child.kill())

    with pytest.raises(DocumentWorkerTimeoutError, match="thời gian"):
        run_document_job("validate_docx", {"content": _docx_bytes()}, timeout_seconds=1)

    assert process.killed
    assert not (tmp_path / "worker").exists()


def test_document_worker_crash_returns_safe_error_and_cleans_job_dir(
    monkeypatch, tmp_path
):
    class CrashedProcess:
        pid = os.getpid()
        returncode = 9

        def communicate(self, timeout=None):
            return b"", b"internal parser details"

    monkeypatch.setenv("DOCUMENT_WORKER_TEMP_DIR", str(tmp_path / "worker"))
    monkeypatch.setattr(
        document_worker.subprocess,
        "Popen",
        lambda *_a, **_k: CrashedProcess(),
    )

    with pytest.raises(DocumentWorkerError, match="dừng bất thường") as error:
        run_document_job("validate_docx", {"content": _docx_bytes()})

    assert "internal parser details" not in str(error.value)
    assert not (tmp_path / "worker").exists()


def test_document_worker_enforces_concurrency_quota(monkeypatch):
    monkeypatch.setenv("DOCUMENT_WORKER_MAX_CONCURRENCY", "1")
    monkeypatch.setenv("DOCUMENT_WORKER_QUEUE_TIMEOUT_SECONDS", "0")
    semaphore = document_worker._worker_semaphore()
    assert semaphore.acquire(blocking=False)
    try:
        with pytest.raises(DocumentWorkerBusyError, match="quá nhiều"):
            run_document_job("validate_docx", {"content": _docx_bytes()})
    finally:
        semaphore.release()


def test_async_document_worker_rejects_before_unbounded_executor_queue(monkeypatch):
    monkeypatch.setenv("DOCUMENT_WORKER_MAX_CONCURRENCY", "1")
    monkeypatch.setenv("DOCUMENT_WORKER_QUEUE_SIZE", "0")
    entered = threading.Event()
    release = threading.Event()

    def blocking_job(*_args, **_kwargs):
        entered.set()
        assert release.wait(timeout=5)
        return "done"

    monkeypatch.setattr(document_worker, "run_document_job", blocking_job)

    async def scenario():
        first = asyncio.create_task(
            document_worker.run_document_job_async("validate_docx", {"content": b"first"})
        )
        assert await asyncio.to_thread(entered.wait, 2)
        with pytest.raises(DocumentWorkerBusyError):
            await document_worker.run_document_job_async(
                "validate_docx", {"content": b"second"}
            )
        release.set()
        assert await first == "done"

    asyncio.run(scenario())


def test_cancelled_async_request_keeps_admission_until_worker_finishes(monkeypatch):
    monkeypatch.setenv("DOCUMENT_WORKER_MAX_CONCURRENCY", "1")
    monkeypatch.setenv("DOCUMENT_WORKER_QUEUE_SIZE", "0")
    entered = threading.Event()
    release = threading.Event()

    def blocking_job(*_args, **_kwargs):
        entered.set()
        assert release.wait(timeout=5)
        return "done"

    monkeypatch.setattr(document_worker, "run_document_job", blocking_job)

    async def scenario():
        first = asyncio.create_task(
            document_worker.run_document_job_async("validate_docx", {"content": b"first"})
        )
        assert await asyncio.to_thread(entered.wait, 2)
        first.cancel()
        with pytest.raises(asyncio.CancelledError):
            await first
        with pytest.raises(DocumentWorkerBusyError):
            await document_worker.run_document_job_async(
                "validate_docx", {"content": b"second"}
            )
        release.set()
        for _ in range(100):
            await asyncio.sleep(0.01)
            runtime = document_worker._async_runtime
            if runtime is not None and runtime.admission.acquire(blocking=False):
                runtime.admission.release()
                return
        raise AssertionError("worker admission was not released after completion")

    asyncio.run(scenario())


def test_document_worker_environment_omits_database_path_and_uses_private_cwd(
    monkeypatch, tmp_path
):
    monkeypatch.setenv("BIDDING_DB_PATH", str(tmp_path / "secret.db"))
    environment = document_worker._worker_environment(tmp_path)

    assert "BIDDING_DB_PATH" not in environment
    assert environment["DOCUMENT_WORKER_JOB_DIR"] == str(tmp_path)
    assert environment["PYTHONPATH"] == str(document_worker.PROJECT_ROOT)


def test_stale_document_job_cleanup(monkeypatch, tmp_path):
    root = tmp_path / "worker"
    stale = root / "job-stale"
    current = root / "job-current"
    stale.mkdir(parents=True)
    current.mkdir()
    old_time = time.time() - 120
    os.utime(stale, (old_time, old_time))
    monkeypatch.setenv("DOCUMENT_WORKER_TEMP_DIR", str(root))

    assert cleanup_stale_document_jobs(max_age_seconds=60) == 1
    assert not stale.exists()
    assert current.exists()
