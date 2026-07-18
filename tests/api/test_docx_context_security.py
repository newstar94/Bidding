import io
import sqlite3
from types import SimpleNamespace

import pytest
from docx import Document
from PIL import Image

from backend.documents import custom_exporter
from backend.documents.custom_exporter import (
    TemplateRenderError,
    _collect_image_tasks,
    convert_images_in_context,
    generate_report_from_custom_template,
)
from backend.documents.docx_context_policy import (
    filter_mapping_rows,
    project_docx_context,
    seal_docx_context,
    sensitive_capability_groups_present,
    validate_docx_context_manifest,
    validate_mapping_definition,
)
from backend.documents.docx_service import enrich_bids_with_contractor_fields
from backend.documents.document_worker_entry import _run_operation
from backend.shared.access_policy import DocumentExportCapabilities


ALL_SENSITIVE_CAPABILITIES = DocumentExportCapabilities.allow_all()


def _template(path, text):
    document = Document()
    document.add_paragraph(text)
    document.save(path)


def test_plan_dto_strips_auth_secrets_unknown_roots_and_nested_payloads():
    context = project_docx_context(
        "plan",
        {
            "user": {
                "ten_dang_nhap": "lan",
                "ho_ten": "Lan",
                "email": "lan@example.test",
                "mat_khau": "password-hash",
                "ma_xac_minh": "verification-secret",
            },
            "ke_hoach": {
                "ma_ke_hoach": "KH-01",
                "organization_id": "org-secret",
                "cv_da_thuc_hien": [
                    {
                        "ten_cong_viec": "Công việc an toàn",
                        "mat_khau": "nested-secret",
                    }
                ],
            },
            "auth_session": {"token_hash": "session-secret"},
            "nha_thau": [{"so_tai_khoan": "not-in-plan"}],
        },
    )

    assert context["user"] == {
        "ten_dang_nhap": "lan",
        "ho_ten": "Lan",
        "email": "lan@example.test",
    }
    assert context["ke_hoach"]["cv_da_thuc_hien"] == [
        {"ten_cong_viec": "Công việc an toàn"}
    ]
    assert "auth_session" not in context
    assert "nha_thau" not in context
    assert "organization_id" not in context["ke_hoach"]


def test_report_dto_keeps_only_managed_related_sensitive_fields():
    context = project_docx_context(
        "evaluation",
        {
            "nha_thau": [
                {
                    "id": "bid-1",
                    "nha_thau_id": "contractor-1",
                    "so_tai_khoan": "123456789",
                    "noi_mo_tai_khoan": "Ngân hàng A",
                    "anh_dau": "images/nha_thau/stamp.png",
                    "mat_khau": "must-not-pass",
                }
            ],
            "to_chuyen_gia": [
                {
                    "id": "expert-1",
                    "ho_ten": "Chuyên gia A",
                    "so_cccd": "001234567890",
                    "anh_chung_chi": "images/chuyen_gia/cert.png",
                    "anh_chu_ky": "images/chuyen_gia/../nha_thau/stamp.png",
                    "token_hash": "must-not-pass",
                }
            ],
        },
        ALL_SENSITIVE_CAPABILITIES,
    )

    contractor = context["nha_thau"][0]
    expert = context["to_chuyen_gia"][0]
    assert contractor["so_tai_khoan"] == "123456789"
    assert contractor["noi_mo_tai_khoan"] == "Ngân hàng A"
    assert contractor["anh_dau"] == "images/nha_thau/stamp.png"
    assert "mat_khau" not in contractor
    assert expert["so_cccd"] == "001234567890"
    assert expert["anh_chung_chi"] == "images/chuyen_gia/cert.png"
    assert expert["anh_chu_ky"] == ""
    assert "token_hash" not in expert


def test_mapping_policy_rejects_auth_secrets_and_keeps_supported_sensitive_fields():
    with pytest.raises(ValueError, match="DTO"):
        validate_mapping_definition("leak", "tai_khoan", "mat_khau")
    with pytest.raises(ValueError, match="không được phép"):
        validate_mapping_definition("leak", "auth_sessions", "token_hash")

    validate_mapping_definition("stk_nt", "nha_thau", "so_tai_khoan")
    validate_mapping_definition("cccd_cg", "chuyen_gia", "so_cccd")
    validate_mapping_definition("chu_ky_cg", "chuyen_gia", "anh_chu_ky")

    rows = [
        ("leak", "tai_khoan", "mat_khau"),
        ("stk_nt", "nha_thau", "so_tai_khoan"),
    ]
    assert filter_mapping_rows(
        rows, "evaluation", ALL_SENSITIVE_CAPABILITIES
    ) == [
        ("stk_nt", "nha_thau", "so_tai_khoan")
    ]


def test_custom_mapping_values_are_reprojected_instead_of_cloned():
    context, manifest = seal_docx_context(
        "evaluation",
        {
            "danh_sach_nt": [
                {
                    "ten_nha_thau": "Nhà thầu A",
                    "so_tai_khoan": "123456789",
                    "mat_khau": "must-not-pass",
                }
            ],
            "cccd_cg": {"so_cccd": "must-not-pass-as-object"},
        },
        [
            ("danh_sach_nt", "nha_thau", ""),
            ("cccd_cg", "chuyen_gia", "so_cccd"),
        ],
        ALL_SENSITIVE_CAPABILITIES,
    )

    assert context["danh_sach_nt"] == [
        {
            "ten_nha_thau": "Nhà thầu A",
            "so_tai_khoan": "123456789",
        }
    ]
    assert context["cccd_cg"] is None
    assert manifest["custom_root_keys"] == ["cccd_cg", "danh_sach_nt"]


def test_sensitive_docx_fields_default_deny_and_follow_each_capability():
    source = {
        "nha_thau": [{
            "ten_nha_thau": "Nhà thầu A",
            "so_tai_khoan": "123456789",
            "anh_dau": "images/nha_thau/stamp.png",
        }],
        "to_chuyen_gia": [{
            "ho_ten": "Chuyên gia A",
            "so_cccd": "001234567890",
            "anh_chu_ky": "images/chuyen_gia/signature.png",
        }],
    }

    denied = project_docx_context("evaluation", source)
    assert "so_tai_khoan" not in denied["nha_thau"][0]
    assert "anh_dau" not in denied["nha_thau"][0]
    assert "so_cccd" not in denied["to_chuyen_gia"][0]
    assert "anh_chu_ky" not in denied["to_chuyen_gia"][0]

    financial_only = project_docx_context(
        "evaluation",
        source,
        DocumentExportCapabilities(financial=True),
    )
    assert financial_only["nha_thau"][0]["so_tai_khoan"] == "123456789"
    assert "anh_dau" not in financial_only["nha_thau"][0]
    assert "so_cccd" not in financial_only["to_chuyen_gia"][0]

    identity_and_signature = project_docx_context(
        "evaluation",
        source,
        DocumentExportCapabilities(identity=True, signature=True),
    )
    assert "so_tai_khoan" not in identity_and_signature["nha_thau"][0]
    assert identity_and_signature["nha_thau"][0]["anh_dau"].endswith("stamp.png")
    assert identity_and_signature["to_chuyen_gia"][0]["so_cccd"] == "001234567890"
    assert identity_and_signature["to_chuyen_gia"][0]["anh_chu_ky"].endswith(
        "signature.png"
    )


def test_manifest_never_grants_image_fields_without_signature_capability():
    context, manifest = seal_docx_context(
        "evaluation",
        {"nha_thau": [{"anh_dau": "images/nha_thau/stamp.png"}]},
        capabilities=DocumentExportCapabilities(financial=True, identity=True),
    )

    assert "anh_dau" not in context["nha_thau"][0]
    assert manifest["image_fields"] == {}


def test_empty_image_allowlist_does_not_fall_back_to_sensitive_defaults(
    monkeypatch,
):
    context = {"anh_dau": "images/nha_thau/stamp.png"}
    monkeypatch.setattr(
        custom_exporter,
        "_collect_image_tasks",
        lambda _data, _root, tasks=None, allowed_image_fields=None: (
            [] if allowed_image_fields == {} else [("unexpected",)]
        ),
    )

    convert_images_in_context(SimpleNamespace(sections=[]), context, {})

    assert context["anh_dau"] == "images/nha_thau/stamp.png"


def test_sensitive_audit_summary_contains_only_capability_names():
    context = project_docx_context(
        "evaluation",
        {
            "nha_thau": [{
                "so_tai_khoan": "123456789",
                "anh_dau": "images/nha_thau/stamp.png",
            }],
            "to_chuyen_gia": [{"so_cccd": "001234567890"}],
        },
        ALL_SENSITIVE_CAPABILITIES,
    )

    summary = sensitive_capability_groups_present(context)

    assert summary == {"financial", "identity", "signature"}
    serialized_summary = repr(sorted(summary))
    assert "123456789" not in serialized_summary
    assert "001234567890" not in serialized_summary
    assert "stamp.png" not in serialized_summary


def test_context_manifest_is_exact_and_template_cannot_use_unknown_root(tmp_path):
    context, manifest = seal_docx_context(
        "plan",
        {
            "investor_name": "Chủ đầu tư A",
            "password_hash": "must-not-pass",
        },
    )
    assert context == {"investor_name": "Chủ đầu tư A"}
    assert manifest["root_keys"] == ["investor_name"]

    with pytest.raises(ValueError, match="không khớp"):
        validate_docx_context_manifest(
            {**context, "password_hash": "must-not-pass"}, manifest
        )

    template_path = tmp_path / "unknown-root.docx"
    _template(template_path, "{{ password_hash }}")
    with pytest.raises(TemplateRenderError, match="không thể kết xuất"):
        generate_report_from_custom_template(
            template_path,
            context,
            context_manifest=manifest,
        )


def test_document_worker_rejects_docx_context_without_manifest():
    with pytest.raises(ValueError, match="thiếu manifest"):
        _run_operation(
            "render_docx",
            {"template_path": "unused.docx", "context": {}},
        )


def test_image_embedding_requires_allowed_field_folder_and_contained_file(
    monkeypatch, tmp_path
):
    image_root = tmp_path / "images"
    contractor_dir = image_root / "nha_thau"
    expert_dir = image_root / "chuyen_gia"
    contractor_dir.mkdir(parents=True)
    expert_dir.mkdir(parents=True)
    (contractor_dir / "stamp.png").write_bytes(b"managed-stamp")
    (expert_dir / "signature.png").write_bytes(b"managed-signature")
    monkeypatch.setattr(custom_exporter, "IMAGE_DIR", image_root)

    context = {
        "anh_dau": "images/nha_thau/stamp.png",
        "avatar": "images/nha_thau/stamp.png",
        "anh_chung_chi": "images/chuyen_gia/../nha_thau/stamp.png",
        "nested": {"anh_chu_ky": "images/chuyen_gia/signature.png"},
    }
    tasks = _collect_image_tasks(context, tmp_path)

    assert [(task[1], task[2]) for task in tasks] == [
        ("anh_dau", str(contractor_dir / "stamp.png")),
        ("anh_chu_ky", str(expert_dir / "signature.png")),
    ]


def test_worker_payload_renders_granted_bank_identity_and_image_without_auth_secrets(
    monkeypatch,
    tmp_path,
):
    image_root = tmp_path / "images"
    contractor_dir = image_root / "nha_thau"
    contractor_dir.mkdir(parents=True)
    stamp_path = contractor_dir / "stamp.png"
    Image.new("RGBA", (16, 16), (0, 90, 180, 255)).save(stamp_path)
    monkeypatch.setattr(custom_exporter, "IMAGE_DIR", image_root)

    template_path = tmp_path / "sensitive-fields.docx"
    template = Document()
    template.add_paragraph(
        "{% for nt in nha_thau %}Bank={{ nt.so_tai_khoan }} {{ nt.anh_dau }}{% endfor %}"
    )
    template.add_paragraph(
        "{% for cg in to_chuyen_gia %}CCCD={{ cg.so_cccd }}{% endfor %}"
    )
    template.save(template_path)
    context, manifest = seal_docx_context(
        "evaluation",
        {
            "nha_thau": [{
                "so_tai_khoan": "123456789",
                "anh_dau": "images/nha_thau/stamp.png",
                "mat_khau": "must-not-pass",
            }],
            "to_chuyen_gia": [{
                "so_cccd": "001234567890",
                "ma_xac_minh": "must-not-pass",
            }],
            "auth_session": {"token_hash": "must-not-pass"},
        },
        capabilities=ALL_SENSITIVE_CAPABILITIES,
    )

    serialized_payload = repr(context)
    assert "123456789" in serialized_payload
    assert "001234567890" in serialized_payload
    assert "images/nha_thau/stamp.png" in serialized_payload
    assert "mat_khau" not in serialized_payload
    assert "ma_xac_minh" not in serialized_payload
    assert "token_hash" not in serialized_payload

    rendered = _run_operation(
        "render_docx",
        {
            "template_path": str(template_path),
            "context": context,
            "context_manifest": manifest,
        },
    )
    rendered_document = Document(io.BytesIO(rendered))
    rendered_text = "\n".join(
        paragraph.text for paragraph in rendered_document.paragraphs
    )

    assert "Bank=123456789" in rendered_text
    assert "CCCD=001234567890" in rendered_text
    assert len(rendered_document.inline_shapes) == 1
    assert "must-not-pass" not in rendered_text


def test_contractor_enrichment_is_scoped_to_active_workspace():
    connection = sqlite3.connect(":memory:")
    connection.row_factory = sqlite3.Row
    cursor = connection.cursor()
    cursor.execute(
        """
        CREATE TABLE nha_thau (
            id TEXT PRIMARY KEY,
            organization_id TEXT NOT NULL,
            ten_nha_thau TEXT,
            ten_viet_tat TEXT,
            ma_nha_thau TEXT,
            ma_so_thue TEXT,
            nguoi_dai_dien TEXT,
            chuc_vu_dai_dien TEXT,
            danh_xung TEXT,
            dia_chi TEXT,
            dia_chi_goc TEXT,
            so_dien_thoai TEXT,
            email TEXT,
            so_tai_khoan TEXT,
            noi_mo_tai_khoan TEXT,
            ma_ngan_hang TEXT
        )
        """
    )
    cursor.execute(
        """INSERT INTO nha_thau (
               id, organization_id, ten_nha_thau, so_tai_khoan
           ) VALUES ('cross-org', 'org-other', 'Nhà thầu khác', '999999')"""
    )
    bids = [{"nha_thau_id": "cross-org", "ten_nha_thau": "Snapshot"}]

    enrich_bids_with_contractor_fields(cursor, bids, "org-active")

    assert bids == [{"nha_thau_id": "cross-org", "ten_nha_thau": "Snapshot"}]
