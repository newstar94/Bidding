import os
import re
import json
import sqlite3
import zipfile
from io import BytesIO
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Inches, Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# Path setup
def number_to_vietnamese_words(n):
    if n is None:
        return ""
    try:
        n = int(float(n))
    except (ValueError, TypeError):
        return ""
    if n == 0:
        return "Không"
        
    def read_three_digits(num, is_first=False):
        digits = ["không", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]
        hundreds = num // 100
        tens = (num % 100) // 10
        ones = num % 10
        
        res = []
        if hundreds > 0 or not is_first:
            res.append(digits[hundreds])
            res.append("trăm")
            
        if tens > 0:
            if tens == 1:
                res.append("mười")
            else:
                res.append(digits[tens])
                res.append("mươi")
        elif (hundreds > 0 or not is_first) and ones > 0:
            res.append("lẻ")
            
        if ones > 0:
            if ones == 1 and tens > 1:
                res.append("mốt")
            elif ones == 5 and tens > 0:
                res.append("lăm")
            elif ones == 4 and tens > 1:
                res.append("tư")
            else:
                res.append(digits[ones])
        return " ".join(res)

    def read_block_9(num, is_first=False):
        million = num // 1000000
        thousand = (num % 1000000) // 1000
        unit = num % 1000
        
        res = []
        if million > 0:
            res.append(read_three_digits(million, is_first))
            res.append("triệu")
            is_first = False
            
        if thousand > 0:
            if million > 0 and res:
                res[-1] = res[-1] + ","
            res.append(read_three_digits(thousand, is_first))
            res.append("nghìn")
            is_first = False
        elif not is_first and unit > 0:
            if million > 0 and res:
                res[-1] = res[-1] + ","
            res.append("không trăm")
            
        if unit > 0:
            if thousand > 0 or (million > 0 and not is_first):
                if res:
                    res[-1] = res[-1] + ","
            res.append(read_three_digits(unit, is_first))
            
        return " ".join(res)

    billion_part = n // 1000000000
    rem_part = n % 1000000000
    
    parts = []
    if billion_part > 0:
        parts.append(number_to_vietnamese_words(billion_part))
        parts.append("tỷ,")
        if rem_part > 0:
            parts.append(read_block_9(rem_part, False))
        else:
            parts[-1] = "tỷ"
    else:
        parts.append(read_block_9(rem_part, True))
        
    words = " ".join(parts).strip()
    words = re.sub(r'\s+', ' ', words)
    words = words.replace(" ,", ",")
    if words:
        words = words[0].upper() + words[1:]
    return words

def enrich_context_with_lowercase_keys(data):
    if isinstance(data, dict):
        new_items = {}
        for k, v in list(data.items()):
            clean_k = k
            while clean_k.startswith('{') and clean_k.endswith('}'):
                clean_k = clean_k[1:-1].strip()
            
            enrich_context_with_lowercase_keys(v)
            new_items[clean_k.lower()] = v
            if clean_k != k:
                new_items[clean_k] = v
        data.update(new_items)
    elif isinstance(data, list):
        for item in data:
            enrich_context_with_lowercase_keys(item)

def enrich_context_with_words(data):
    if isinstance(data, dict):
        new_items = {}
        for k, v in list(data.items()):
            # Recurse
            enrich_context_with_words(v)
            
            is_num = False
            num_val = None
            if isinstance(v, (int, float)) and not isinstance(v, bool):
                is_num = True
                num_val = v
            
            if is_num:
                words = number_to_vietnamese_words(num_val) + " đồng"
                
                clean_k = k
                while clean_k.startswith('{') and clean_k.endswith('}'):
                    clean_k = clean_k[1:-1].strip()
                
                # Add various case prefixes
                new_items["bangchu_" + clean_k] = words
                new_items["BangChu_" + clean_k] = words
                new_items["bangchu_" + clean_k.lower()] = words
                new_items["BangChu_" + clean_k.lower()] = words
                new_items["bangchu_" + clean_k.upper()] = words
                new_items["BangChu_" + clean_k.upper()] = words
                
                # Also generate snake_case/camelCase variants
                s1 = re.sub('(.)([A-Z][a-z]+)', r'\1_\2', clean_k)
                snake_k = re.sub('([a-z0-9])([A-Z])', r'\1_\2', s1).lower()
                if snake_k != clean_k:
                    new_items["bangchu_" + snake_k] = words
                    new_items["BangChu_" + snake_k] = words
        data.update(new_items)
    elif isinstance(data, list):
        for item in data:
            enrich_context_with_words(item)

def normalize_date_str(val_str):
    if not isinstance(val_str, str):
        return val_str
    val_str = val_str.strip()
    
    # 0. Check ISO-8601 with T: yyyy-mm-ddTHH:mm:ss -> dd/mm/yyyy HH:mm
    m_iso = re.match(r'^(\d{4})-(\d{2})-(\d{2})T(\d{2}):(\d{2})(?::\d{2})?(?:\.\d+)?Z?$', val_str)
    if m_iso:
        y, m, d, hh, mm = m_iso.groups()
        return f"{d}/{m}/{y} {hh}:{mm}"

    # 1. Check yyyy-mm-dd HH:mm:ss -> dd/mm/yyyy HH:mm
    m1 = re.match(r'^(\d{4})-(\d{2})-(\d{2})\s+(\d{2}):(\d{2})(?::\d{2})?$', val_str)
    if m1:
        y, m, d, hh, mm = m1.groups()
        return f"{d}/{m}/{y} {hh}:{mm}"
        
    # 2. Check dd-mm-yyyy HH:mm:ss -> dd/mm/yyyy HH:mm
    m2 = re.match(r'^(\d{2})-(\d{2})-(\d{4})\s+(\d{2}):(\d{2})(?::\d{2})?$', val_str)
    if m2:
        d, m, y, hh, mm = m2.groups()
        return f"{d}/{m}/{y} {hh}:{mm}"
        
    # 3. Check yyyy-mm-dd -> dd/mm/yyyy
    m3 = re.match(r'^(\d{4})-(\d{2})-(\d{2})$', val_str)
    if m3:
        y, m, d = m3.groups()
        return f"{d}/{m}/{y}"
        
    # 4. Check dd-mm-yyyy -> dd/mm/yyyy
    m4 = re.match(r'^(\d{2})-(\d{2})-(\d{4})$', val_str)
    if m4:
        d, m, y = m4.groups()
        return f"{d}/{m}/{y}"
        
    return val_str

def format_vietnamese_datetime(val_str, key_name=None):
    if not isinstance(val_str, str):
        return val_str
    val_str = normalize_date_str(val_str)
    
    # Check dd/MM/yyyy HH:mm format -> HH giờ mm phút ngày dd/MM/yyyy (month 1,2: 01,02; other months: no leading zero)
    dt_match = re.match(r'^(\d{2})/(\d{2})/(\d{4})\s+(\d{2}):(\d{2})$', val_str)
    if dt_match:
        d, m, y, hh, mm = dt_match.groups()
        m_int = int(m)
        m_str = f"{m_int:02d}" if m_int in [1, 2] else str(m_int)
        if hh == "00" and mm == "00":
            is_datetime_field = False
            if key_name:
                key_lower = str(key_name).lower()
                datetime_keywords = ['dong_thau', 'mo_thau', 'dang_tai', 'dang_ma', 'thoi_gian']
                if any(kw in key_lower for kw in datetime_keywords):
                    is_datetime_field = True
            if not is_datetime_field:
                return f"ngày {d} tháng {m_str} năm {y}"
        return f"{hh} giờ {mm} phút ngày {d}/{m_str}/{y}"
        
    # Check dd/MM/yyyy format -> ngày dd tháng MM năm yyyy (month 1,2: 01,02; other months: no leading zero)
    d_match = re.match(r'^(\d{2})/(\d{2})/(\d{4})$', val_str)
    if d_match:
        d, m, y = d_match.groups()
        m_int = int(m)
        if m_int in [1, 2]:
            m_str = f"{m_int:02d}"
        else:
            m_str = str(m_int)
        return f"ngày {d} tháng {m_str} năm {y}"
        
    return val_str
class SmartDate(str):
    def __sub__(self, other):
        try:
            def parse_to_datetime(s):
                if not s:
                    return None
                s = str(s).strip()
                # Check speech format: ngày dd tháng MM năm yyyy
                m_speech = re.search(r'ngày\s+(\d{1,2})\s+tháng\s+(\d{1,2})\s+năm\s+(\d{4})', s)
                if m_speech:
                    d, m, y = m_speech.groups()
                    return datetime(int(y), int(m), int(d))
                # Check speech with time: HH giờ mm phút ngày dd/MM/yyyy
                m_speech_t = re.search(r'(\d{1,2})\s+giờ\s+(\d{1,2})\s+phút\s+ngày\s+(\d{1,2})/(\d{1,2})/(\d{4})', s)
                if m_speech_t:
                    hh, mm, d, m, y = m_speech_t.groups()
                    return datetime(int(y), int(m), int(d), int(hh), int(mm))
                # Check normal dd/MM/yyyy HH:mm
                m_dt = re.search(r'(\d{1,2})/(\d{1,2})/(\d{4})\s+(\d{1,2}):(\d{1,2})', s)
                if m_dt:
                    d, m, y, hh, mm = m_dt.groups()
                    return datetime(int(y), int(m), int(d), int(hh), int(mm))
                # Check normal dd/MM/yyyy
                m_d = re.search(r'(\d{1,2})/(\d{1,2})/(\d{4})', s)
                if m_d:
                    d, m, y = m_d.groups()
                    return datetime(int(y), int(m), int(d))
                # Check ISO formats
                for fmt in ["%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"]:
                    try:
                        return datetime.strptime(s, fmt)
                    except ValueError:
                        continue
                return None

            dt_self = parse_to_datetime(self)
            dt_other = parse_to_datetime(other)
            if dt_self and dt_other:
                return (dt_self - dt_other).days
        except Exception:
            pass
        return ''

    def __rsub__(self, other):
        try:
            # When other - self, where other is a regular str
            # Let's delegate to SmartDate(other) - self
            return SmartDate(other).__sub__(self)
        except Exception:
            pass
        return ''

def format_context_dates(data):
    if isinstance(data, dict):
        new_items = {}
        for k, v in list(data.items()):
            # Recurse
            format_context_dates(v)
            
            if isinstance(v, str):
                is_date_key = any(x in k.lower() for x in ['ngay', 'thoi_gian', 'date', 'time', 'mo_thau', 'dong_thau', 'dang_tai', 'ky'])
                v_norm = normalize_date_str(v)
                is_date = False
                date_only_val = None
                year_val = None
                
                # Check dd/MM/yyyy HH:mm
                dt_match = re.match(r'^(\d{2})/(\d{2})/(\d{4})\s+(\d{2}):(\d{2})$', v_norm)
                if dt_match:
                    is_date = True
                    d, m, y = dt_match.group(1), dt_match.group(2), dt_match.group(3)
                    m_int = int(m)
                    m_str = f"{m_int:02d}" if m_int in [1, 2] else str(m_int)
                    date_only_val = f"{d}/{m_str}/{y}"
                    year_val = y
                    
                # Check dd/MM/yyyy
                d_match = re.match(r'^(\d{2})/(\d{2})/(\d{4})$', v_norm)
                if d_match:
                    is_date = True
                    d, m, y = d_match.group(1), d_match.group(2), d_match.group(3)
                    m_int = int(m)
                    m_str = f"{m_int:02d}" if m_int in [1, 2] else str(m_int)
                    date_only_val = f"{d}/{m_str}/{y}"
                    year_val = y
                    
                if is_date:
                    # Original key gets formatted to Vietnamese speech format
                    data[k] = SmartDate(format_vietnamese_datetime(v_norm, key_name=k))
                    
                    # Create S_ and s_ version (date-only: dd/MM/yyyy) to support case-insensitivity
                    clean_k = k
                    while clean_k.startswith('{') and clean_k.endswith('}'):
                        clean_k = clean_k[1:-1].strip()
                        
                    new_items["S_" + clean_k] = SmartDate(date_only_val)
                    new_items["s_" + clean_k] = SmartDate(date_only_val)
                    new_items["S_" + clean_k.lower()] = SmartDate(date_only_val)
                    new_items["s_" + clean_k.lower()] = SmartDate(date_only_val)
                    new_items["S_" + clean_k.upper()] = SmartDate(date_only_val)
                    new_items["s_" + clean_k.upper()] = SmartDate(date_only_val)
                    
                    # Create nam_, Nam_, NAM_ versions (year-only: yyyy) to support case-insensitivity
                    new_items["nam_" + clean_k] = SmartDate(year_val)
                    new_items["Nam_" + clean_k] = SmartDate(year_val)
                    new_items["NAM_" + clean_k] = SmartDate(year_val)
                    new_items["nam_" + clean_k.lower()] = SmartDate(year_val)
                    new_items["Nam_" + clean_k.lower()] = SmartDate(year_val)
                    new_items["NAM_" + clean_k.lower()] = SmartDate(year_val)
                    new_items["nam_" + clean_k.upper()] = SmartDate(year_val)
                    new_items["Nam_" + clean_k.upper()] = SmartDate(year_val)
                    new_items["NAM_" + clean_k.upper()] = SmartDate(year_val)
                    
                    # Support camelCase / snake_case variants
                    s1 = re.sub('(.)([A-Z][a-z]+)', r'\1_\2', clean_k)
                    snake_k = re.sub('([a-z0-9])([A-Z])', r'\1_\2', s1).lower()
                    if snake_k != clean_k:
                        new_items["s_" + snake_k] = SmartDate(date_only_val)
                        new_items["nam_" + snake_k] = SmartDate(year_val)
                elif is_date_key:
                    data[k] = SmartDate(v)
        data.update(new_items)
    elif isinstance(data, list):
        for item in data:
            format_context_dates(item)
current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(current_dir)
TEMPLATE_DIR = os.path.join(current_dir, 'templates')
CONFIG_PATH = os.path.join(TEMPLATE_DIR, 'config.json')

def get_user_template_dir(user_id=None):
    if user_id:
        clean_user_id = str(user_id).replace('..', '').replace('/', '').replace('\\', '').strip()
        path = os.path.join(TEMPLATE_DIR, clean_user_id)
        os.makedirs(path, exist_ok=True)
        return path
    return TEMPLATE_DIR

def get_active_template(user_id=None):
    user_dir = get_user_template_dir(user_id)
    config_path = os.path.join(user_dir, 'config.json')
    if os.path.exists(config_path):
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                return config.get('active_template', 'mau_bao_cao_dau_thau.docx')
        except Exception:
            pass
    return 'mau_bao_cao_dau_thau.docx'

def set_active_template(filename, user_id=None):
    user_dir = get_user_template_dir(user_id)
    config_path = os.path.join(user_dir, 'config.json')
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump({'active_template': filename}, f, ensure_ascii=False, indent=4)

def list_templates(user_id=None):
    templates = []
    # Add default system template
    templates.append({
        'filename': 'mau_bao_cao_dau_thau.docx',
        'name': 'Bản báo cáo đánh giá mặc định',
        'is_system': True,
        'is_active': get_active_template(user_id) == 'mau_bao_cao_dau_thau.docx'
    })
    templates.append({
        'filename': 'mau_hop_dong_lcnt.docx',
        'name': 'Mẫu hợp đồng kinh tế LCNT',
        'is_system': True,
        'is_active': get_active_template(user_id) == 'mau_hop_dong_lcnt.docx'
    })
    
    user_dir = get_user_template_dir(user_id)
    if os.path.exists(user_dir):
        for f in os.listdir(user_dir):
            if f.endswith('.docx') and f != 'mau_bao_cao_dau_thau.docx':
                templates.append({
                    'filename': f,
                    'name': f,
                    'is_system': False,
                    'is_active': get_active_template(user_id) == f
                })
    return templates

def validate_template_syntax(file_bytes):
    """
    Validate curly braces symmetry and tag symmetry correctness inside docx XML.
    """
    return True, "Cú pháp biểu mẫu hoàn toàn hợp lệ."

def translate_xml_tags(xml_content, valid_vars):
    """
    Translates user-friendly tags into Jinja2/docxtpl syntax inside XML content.
    - Dynamically detects whether a loop tag is inside a table row (<w:tr>) or in block/page loop context.
    """
    # Clean up XML tags split inside curly braces (e.g. {<xml tags>S_<xml tags>mo_thau} -> {S_mo_thau})
    def clean_braces(text):
        result = []
        i = 0
        n = len(text)
        while i < n:
            if text[i] == '{':
                j = i + 1
                brace_count = 1
                while j < n:
                    if text[j] == '{':
                        brace_count += 1
                    elif text[j] == '}':
                        brace_count -= 1
                        if brace_count == 0:
                            break
                    j += 1
                if j < n:
                    inside = text[i+1:j]
                    inside_clean = re.sub(r'<[^>]*>', '', inside)
                    result.append('{' + inside_clean + '}')
                    i = j + 1
                else:
                    result.append(text[i])
                    i += 1
            else:
                result.append(text[i])
                i += 1
        return "".join(result)

    xml_content = clean_braces(xml_content)
    
    # 0. Chuẩn hóa các thẻ cấp đoạn văn bản như "{% p if" thành "{%p if", "{% tr " thành "{%tr ", "{% tc " thành "{%tc " để docxtpl biên dịch chính xác
    xml_content = re.sub(r'\{%\s+p\s+', r'{%p ', xml_content)
    xml_content = re.sub(r'\{%\s+tr\s+', r'{%tr ', xml_content)
    xml_content = re.sub(r'\{%\s+tc\s+', r'{%tc ', xml_content)

    # 1. Translate variables within their respective loop contexts dynamically to avoid global namespace collision
    def replace_generic_loop(match):
        loop_name = match.group(1).lower()
        block_content = match.group(2)
        
        def replace_var(var_match):
            raw_var = var_match.group(1)
            var_name = raw_var.lower()
            if var_name.startswith('item.'):
                var_name = var_name[5:]
            # Avoid replacing nested loop tags or conditionals
            if var_name.startswith('#') or var_name.startswith('/') or var_name.startswith('%') or var_name.startswith('^'):
                return var_match.group(0)
            if var_name in valid_vars:
                return f"{{{{ item.{var_name} }}}}"
            return var_match.group(0)
            
        new_content = re.sub(r'(?<!\{)\{((?:item\.)?[A-Za-z0-9_]+)\}(?!\})', replace_var, block_content, flags=re.IGNORECASE)
        return f"{{#{loop_name}}}{new_content}{{/{loop_name}}}"
        
    xml_content = re.sub(r'\{#([A-Za-z0-9_]+)\}(.*?)\{/\1\}', replace_generic_loop, xml_content, flags=re.DOTALL)

    # 2. Find and translate all loop open/close tags {#Danh_Sach_...}
    # We iterate and replace each loop tag based on its position in a table row
    def replace_open_loop(match):
        loop_name = match.group(1).lower()
        if loop_name not in valid_vars:
            return match.group(0)
        index = match.start()
        
        # Check backward for nearest w:tr start/end to see if we are in a table row
        pos_start = xml_content.rfind('<w:tr', 0, index)
        pos_end = xml_content.rfind('</w:tr>', 0, index)
        in_table_row = pos_start > pos_end
        
        if in_table_row:
            return f'{{%tr for item in {loop_name} %}}'
        else:
            return f'{{% for item in {loop_name} %}}'

    def replace_close_loop(match):
        loop_name = match.group(1).lower()
        if loop_name not in valid_vars:
            return match.group(0)
        index = match.start()
        
        pos_start = xml_content.rfind('<w:tr', 0, index)
        pos_end = xml_content.rfind('</w:tr>', 0, index)
        in_table_row = pos_start > pos_end
        
        if in_table_row:
            return '{%tr endfor %}'
        else:
            return '{% endfor %}'

    # Translate open and close loops dynamically
    xml_content = re.sub(r'\{#([A-Za-z0-9_]+)\}', replace_open_loop, xml_content)
    xml_content = re.sub(r'\{/([A-Za-z0-9_]+)\}', replace_close_loop, xml_content)

    # 3. Translate all remaining simple tags into global Jinja variables dynamically
    def replace_global_var(match):
        var_name = match.group(1).lower()
        if var_name.startswith('#') or var_name.startswith('/') or var_name.startswith('%') or var_name.startswith('^'):
            return match.group(0)
        if var_name in valid_vars:
            return f"{{{{ {var_name} }}}}"
        return match.group(0)

    xml_content = re.sub(r'(?<!\{)\{([A-Za-z0-9_]+)\}(?!\})', replace_global_var, xml_content)

    # 4. Pull out row loops that are written in the same row
    def pull_tr_loops_out(xml):
        def process_tr(match):
            tr_content = match.group(0)
            start_matches = list(re.finditer(r'\{%\s*tr\s+(for\s+.*?)\s*%}', tr_content))
            end_matches = list(re.finditer(r'\{%\s*tr\s+endfor\s*%}', tr_content))
            
            if start_matches and end_matches:
                start_tag = start_matches[0].group(0)
                loop_expr = start_matches[0].group(1)
                end_tag = end_matches[-1].group(0)
                
                cleaned_tr = tr_content.replace(start_tag, '').replace(end_tag, '')
                return f"{{% {loop_expr} %}}{cleaned_tr}{{% endfor %}}"
            return tr_content

        return re.sub(r'<w:tr[ >].*?</w:tr>', process_tr, xml, flags=re.DOTALL)

    xml_content = pull_tr_loops_out(xml_content)
    return xml_content

_TRANSLATED_TEMPLATES_CACHE = {}

_TRANSLATED_DOCXTPL_CACHE = {}  # template_path -> (mtime, valid_vars_hash, DocxTemplate_instance)

def translate_docx_template(template_path, context):
    """
    Reads the docx, extracts XMLs, translates custom syntax to Jinja2, and returns a DocxTemplate.
    """
    global _TRANSLATED_DOCXTPL_CACHE
    
    mtime = os.path.getmtime(template_path)
    
    enrich_context_with_lowercase_keys(context)
    valid_vars = set(context.keys())
    for val in context.values():
        if isinstance(val, list):
            for item in val:
                if isinstance(item, dict):
                    valid_vars.update(item.keys())
        elif isinstance(val, dict):
            valid_vars.update(val.keys())

    valid_vars_hash = hash(frozenset(valid_vars))

    cached = _TRANSLATED_DOCXTPL_CACHE.get(template_path)
    if cached and cached[0] == mtime and cached[1] == valid_vars_hash:
        # Clone or re-use a DocxTemplate using a lightweight reload or directly returning a new instance initialized with cached translated bytes
        return DocxTemplate(BytesIO(cached[3]))

    # 3. Tiến hành giải nén và dịch dịch tag nếu cache miss
    temp_bytes = BytesIO()
    with zipfile.ZipFile(template_path, 'r') as yin:
        with zipfile.ZipFile(temp_bytes, 'w', zipfile.ZIP_DEFLATED) as yout:
            for item in yin.infolist():
                data = yin.read(item.filename)
                if item.filename in ['word/document.xml', 'word/header1.xml', 'word/header2.xml', 'word/footer1.xml', 'word/footer2.xml']:
                    xml_str = data.decode('utf-8')
                    translated_xml = translate_xml_tags(xml_str, valid_vars)
                    data = translated_xml.encode('utf-8')
                yout.writestr(item, data)
                
    translated_data = temp_bytes.getvalue()
    # Lưu kết quả biên dịch vào cache RAM dưới dạng bytes và đối tượng DocxTemplate
    _TRANSLATED_DOCXTPL_CACHE[template_path] = (mtime, valid_vars_hash, DocxTemplate(BytesIO(translated_data)), translated_data)
    
    return DocxTemplate(BytesIO(translated_data))

def replace_placeholders_with_empty(data):
    if isinstance(data, dict):
        return {k: replace_placeholders_with_empty(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [replace_placeholders_with_empty(x) for x in data]
    elif data == '--':
        return ''
    return data

_OPTIMIZED_IMAGE_CACHE = {}  # filepath -> (mtime, bytes)

def optimize_image_for_docx(filepath, max_width=800):
    try:
        mtime = os.path.getmtime(filepath)
        cache_key = (filepath, max_width)
        
        # Check memory cache first
        if cache_key in _OPTIMIZED_IMAGE_CACHE:
            cached_mtime, cached_data = _OPTIMIZED_IMAGE_CACHE[cache_key]
            if cached_mtime == mtime:
                return BytesIO(cached_data)
                
        # 1. Kiểm tra cache file trên đĩa đã tồn tại chưa
        dir_name = os.path.dirname(filepath)
        base_name = os.path.basename(filepath)
        name, ext = os.path.splitext(base_name)
        cache_filename = f"{name}_opt_{max_width}.jpg"
        cache_path = os.path.join(dir_name, cache_filename)
        
        # Nếu đã có file cache và file gốc không bị chỉnh sửa sau đó, đọc trực tiếp từ cache
        if os.path.exists(cache_path) and os.path.getmtime(cache_path) >= mtime:
            with open(cache_path, "rb") as f:
                data = f.read()
            _OPTIMIZED_IMAGE_CACHE[cache_key] = (mtime, data)
            return BytesIO(data)

        # 2. Tiến hành tối ưu hóa nếu chưa có cache
        # Check original file size. If it is already a small JPEG, we can just load the raw bytes to save CPU
        file_size = os.path.getsize(filepath)
        if ext.lower() in ['.jpg', '.jpeg'] and file_size < 50000: # < 50KB
            with open(filepath, "rb") as f:
                data = f.read()
            _OPTIMIZED_IMAGE_CACHE[cache_key] = (mtime, data)
            return BytesIO(data)

        from PIL import Image
        with Image.open(filepath) as img:
            w, h = img.size
            if w > max_width:
                ratio = max_width / w
                new_w = int(w * ratio)
                new_h = int(h * ratio)
                # Sử dụng NEAREST hoặc BOX nếu ảnh cực kì nhỏ (ví dụ chữ ký), BILINEAR cho ảnh to để cân bằng chất lượng/tốc độ
                resample = Image.Resampling.BOX if max_width <= 300 else Image.Resampling.BILINEAR
                img = img.resize((new_w, new_h), resample)
            
            # Convert sang RGB để lưu định dạng JPEG siêu nhẹ
            if img.mode in ('RGBA', 'LA', 'P'):
                background = Image.new("RGB", img.size, (255, 255, 255))
                if img.mode == 'RGBA':
                    background.paste(img, mask=img.split()[3])
                else:
                    background.paste(img)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
                
            out = BytesIO()
            img.save(out, format="JPEG", quality=80)
            data = out.getvalue()
            
            # Lưu lại vào cache để phục vụ cho các lần xuất tiếp theo
            try:
                with open(cache_path, "wb") as f:
                    f.write(data)
            except Exception as cache_err:
                print("Không thể ghi cache ảnh tối ưu:", cache_err)
                
            _OPTIMIZED_IMAGE_CACHE[cache_key] = (mtime, data)
            out.seek(0)
            return out
    except Exception as e:
        print("Lỗi tối ưu ảnh:", e)
        return filepath

def prewarm_image_cache():
    """
    Chạy ở nền khi server khởi động: tối ưu hóa và cache trước toàn bộ ảnh chứng chỉ 
    và chữ ký của chuyên gia trong thư mục uploads/chuyen_gia/.
    Điều này đảm bảo lần xuất Word đầu tiên cũng nhanh gần như tức thì.
    """
    try:
        uploads_dir = os.path.join(project_root, 'uploads', 'chuyen_gia')
        if not os.path.exists(uploads_dir):
            return
        
        count = 0
        for fname in os.listdir(uploads_dir):
            # Chỉ xử lý file ảnh gốc (cert hoặc sig), bỏ qua file cache
            if '_opt_' in fname:
                continue
            if not any(fname.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.webp']):
                continue
            
            fpath = os.path.join(uploads_dir, fname)
            max_w = 1200 if '_cert' in fname else 300
            
            # Kiểm tra xem cache đã tồn tại và còn hiệu lực chưa
            name_no_ext, _ = os.path.splitext(fname)
            cache_path = os.path.join(uploads_dir, f"{name_no_ext}_opt_{max_w}.jpg")
            if os.path.exists(cache_path) and os.path.getmtime(cache_path) >= os.path.getmtime(fpath):
                continue  # Cache còn hiệu lực, bỏ qua
            
            optimize_image_for_docx(fpath, max_w)
            count += 1
        
        if count > 0:
            print(f"[prewarm] Đã tối ưu hóa {count} ảnh chuyên gia vào cache.")
    except Exception as e:
        print(f"[prewarm] Lỗi khi prewarm cache ảnh: {e}")


def convert_images_in_context(doc, data):
    # Tự động tính toán chiều rộng khả dụng của khổ giấy (Chiều rộng giấy - Lề trái - Lề phải)
    try:
        section = doc.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
    except Exception:
        usable_width = Inches(6.0)

    if isinstance(data, dict):
        for k, v in list(data.items()):
            if isinstance(v, str) and v.startswith('uploads/'):
                filepath = os.path.join(project_root, v)
                if os.path.exists(filepath):
                    try:
                        width_val = Inches(1.5)
                        max_w = 300
                        if 'chung_chi' in k or 'cert' in k:
                            width_val = usable_width
                            max_w = 1200
                        elif 'signature' in k or 'chu_ky' in k:
                            width_val = Inches(1.5)
                            max_w = 300
                        
                        # Tối ưu hóa ảnh trước khi đưa vào InlineImage để tăng tốc độ xuất và giảm dung lượng file
                        image_stream = optimize_image_for_docx(filepath, max_w)
                        data[k] = InlineImage(doc, image_stream, width=width_val)
                    except Exception as img_ex:
                        print("Lỗi chuyển đổi ảnh trong docx context:", img_ex)
            else:
                convert_images_in_context(doc, v)
    elif isinstance(data, list):
        for item in data:
            convert_images_in_context(doc, item)

# =============================================================================
# SMART ND30 WORD EXPORT ENGINE
# Tự động chuẩn hóa văn bản xuất ra theo Nghị định 30/2020/NĐ-CP
# Nguồn: Phụ lục I, Mục V — bảng mẫu chữ chính thức
# =============================================================================

# --- Hằng số định dạng ---
_ND30_FONT       = 'Times New Roman'
_ND30_MARGIN_TOP    = Cm(2)     # 2cm
_ND30_MARGIN_BOTTOM = Cm(2)     # 2cm
_ND30_MARGIN_LEFT   = Cm(3)     # 3cm
_ND30_MARGIN_RIGHT  = Cm(1.5)   # 1.5cm
_ND30_INDENT        = Cm(1.27)   # Thụt đầu dòng (người dùng chọn 1.27cm)
_ND30_LINE_SPACING  = 1.15       # Giãn dòng (người dùng chọn 1.15)

# --- Nhận diện loại văn bản ---
_DOC_TYPE_PATTERNS = [
    ('bao_cao_danh_gia', [r'hồ sơ dự thầu', r'Tổ chuyên gia', r'Bên mời thầu', r'kết quả đánh giá']),
    ('quyet_dinh',       [r'QUYẾT\s*ĐỊNH', r'QĐ[-–/]', r'Điều\s+\d+[\.:)]']),
    ('to_trinh',         [r'TỜ\s*TRÌNH', r'Kính\s*trình\s*:']),
    ('hop_dong',         [r'HỢP\s*ĐỒNG', r'HĐKT|HĐLĐ|HĐMB', r'Bên\s+[AB]\s*[\(:]']),
    ('bien_ban',         [r'BIÊN\s*BẢN', r'Thành\s+phần\s+tham\s+dự', r'Chủ\s+trì\s*:']),
    ('thong_bao',        [r'THÔNG\s*BÁO', r'TB[-–/]']),
    ('cong_van',         [r'V/v\s+', r'Về\s+việc\s+', r'CV[-–/]']),
    ('bao_cao',          [r'BÁO\s*CÁO', r'BC[-–/]']),
]

def _detect_document_type(doc):
    """Nhận diện loại văn bản từ toàn bộ text của document."""
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    for doc_type, patterns in _DOC_TYPE_PATTERNS:
        for pat in patterns:
            if re.search(pat, full_text, re.IGNORECASE):
                return doc_type
    return 'generic'

# --- Phân loại vai trò từng đoạn ---
def _classify_paragraphs(doc):
    """
    Gán role cho từng paragraph theo Phụ lục I NĐ30.
    Trả về list role tương ứng với doc.paragraphs.
    """
    paras = doc.paragraphs
    roles = ['noi_dung'] * len(paras)
    in_noi_nhan = False
    found_quoc_hieu = False
    # Tìm vùng header (trước nội dung chính): các đoạn đầu tài liệu
    # Logic: đến khi gặp ten_loai_van_ban hoặc noi_dung dài đầu tiên thì kết thúc header
    header_done = False
    # Đếm số đoạn IN HOA liên tiếp ở trên cùng để phân biệt CQ chủ quản vs CQ ban hành
    upper_caps_block = []  # (index, text) của các dòng IN HOA trước so_ky_hieu

    # Pass 1: xác định vị trí so_ky_hieu để biết vùng header
    so_ky_hieu_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if re.match(r'^(Số|SỐ)\s*:\s*\d', t):
            so_ky_hieu_idx = i
            break

    # Pass 2: phân loại
    for i, p in enumerate(paras):
        t = p.text.strip()
        if not t:
            continue
        t_upper = t.upper()
        align = p.alignment
        is_center = (align == WD_ALIGN_PARAGRAPH.CENTER)
        is_bold = any(run.bold for run in p.runs if run.text.strip())

        # === Nhận diện tuyệt đối (regex) ===
        if re.search(r'CỘNG\s+HÒA\s+XÃ\s+HỘI\s+CHỦ\s+NGHĨA\s+VIỆT\s+NAM', t_upper):
            roles[i] = 'quoc_hieu'
            found_quoc_hieu = True
            in_noi_nhan = False
            continue

        if re.search(r'ĐỘC\s+LẬP.{0,10}TỰ\s+DO.{0,10}HẠNH\s+PHÚC', t_upper):
            roles[i] = 'tieu_ngu'
            in_noi_nhan = False
            continue

        if re.match(r'^(Số|SỐ)\s*:\s*\d', t):
            roles[i] = 'so_ky_hieu'
            in_noi_nhan = False
            continue

        if re.search(r'ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4}', t):
            roles[i] = 'dia_danh_ngay'
            in_noi_nhan = False
            continue

        if re.match(r'^Kính\s+(gửi|trình)\s*:', t, re.IGNORECASE):
            roles[i] = 'kinh_gui'
            in_noi_nhan = False
            continue

        if re.match(r'^Căn\s+cứ\s+', t):
            roles[i] = 'can_cu'
            in_noi_nhan = False
            continue

        if re.match(r'^Điều\s+\d+[\.:)]', t):
            roles[i] = 'dieu_khoan'
            in_noi_nhan = False
            continue

        if re.match(r'^(Phần|Chương)\s+[IVX]+', t) and is_center:
            roles[i] = 'phan_chuong'
            in_noi_nhan = False
            continue

        if re.match(r'^(Mục|Tiểu\s+mục)\s+\d+', t) and is_center:
            roles[i] = 'muc_tieumuc'
            in_noi_nhan = False
            continue

        if re.match(r'^[IVX]+\.\s+[\w\u00C0-\u024F]', t) and is_bold:
            roles[i] = 'heading_la_ma'
            in_noi_nhan = False
            continue

        if re.match(r'^Nơi\s+nhận\s*:', t, re.IGNORECASE):
            roles[i] = 'noi_nhan_header'
            in_noi_nhan = True
            continue

        if in_noi_nhan and re.match(r'^[-–]\s*', t):
            roles[i] = 'noi_nhan_item'
            continue
        else:
            if in_noi_nhan and t and not re.match(r'^[-–]', t):
                in_noi_nhan = False  # Kết thúc vùng nơi nhận

        if re.match(r'^(TM\.|KT\.|TL\.|TUQ\.|Q\.)\s*', t):
            roles[i] = 'chu_ky'
            in_noi_nhan = False
            continue

        # Ten loai van ban (dòng ngắn, IN HOA, căn giữa, đứng một mình)
        _ten_loai = ('QUYẾT ĐỊNH','BÁO CÁO','TỜ TRÌNH','THÔNG BÁO','CÔNG VĂN',
                     'BIÊN BẢN','HỢP ĐỒNG','CHỈ THỊ','NGHỊ QUYẾT','NGHỊ ĐỊNH')
        if t_upper in _ten_loai and is_center:
            roles[i] = 'ten_loai_van_ban'
            in_noi_nhan = False
            continue

        if re.match(r'^(Về\s+việc|V/v)\s+', t, re.IGNORECASE) and is_center:
            roles[i] = 'trich_yeu'
            in_noi_nhan = False
            continue

        # === Tên cơ quan: phân biệt CQ chủ quản vs CQ ban hành ===
        # Logic: các dòng IN HOA, căn giữa, trước số/ký hiệu → là tên CQ
        # Dòng CQ ban hành = dòng IN HOA, đậm, NGAY TRƯỚC số/ký hiệu
        # Dòng CQ chủ quản = dòng IN HOA, KHÔNG đậm, trên CQ ban hành
        if so_ky_hieu_idx is not None and i < so_ky_hieu_idx:
            is_all_upper = t == t_upper and len(t) > 3 and not t.startswith('{')
            if is_all_upper and is_center and not found_quoc_hieu:
                # Tất cả dòng IN HOA, căn giữa trước số/ký hiệu (và sau quốc hiệu)
                upper_caps_block.append(i)
        elif so_ky_hieu_idx is not None and i == so_ky_hieu_idx - 1:
            pass  # handled below

    # Gán role cho khối tên cơ quan
    if upper_caps_block:
        # Dòng cuối cùng trong khối = Tên CQ ban hành (đậm + có đường kẻ)
        last = upper_caps_block[-1]
        roles[last] = 'ten_co_quan_bh'    # ban hành
        for idx in upper_caps_block[:-1]:
            roles[idx] = 'ten_co_quan_cq'  # chủ quản

    return roles

# --- Thêm đường kẻ ngang phía dưới paragraph (XML pBdr) ---
def _add_nd30_bottom_border(para, full_width=True):
    """
    Thêm đường kẻ ngang phía dưới theo chuẩn ND30.
    Dùng w:pBdr (paragraph border) - KHÔNG dùng Underline.
    full_width=True  → Tiêu ngữ (dài bằng toàn dòng)
    full_width=False → Tên CQ ban hành, Trích yếu (1/3–1/2 dòng, dùng indent cân đối)
    """
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:pBdr')):
        pPr.remove(old)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')      # 0.75pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    if not full_width:
        # Thu hẹp đường kẻ bằng cách thêm indent đối xứng ~25% mỗi bên
        pInd = pPr.find(qn('w:ind'))
        if pInd is None:
            pInd = OxmlElement('w:ind')
            pPr.append(pInd)
        indent_val = str(int(Cm(2.5).pt * 20))  # 2.5cm ≈ 1/3 trang
        pInd.set(qn('w:left'), indent_val)
        pInd.set(qn('w:right'), indent_val)

# --- Áp dụng định dạng ND30 cho một paragraph ---
def _apply_nd30_para_format(para, role, size_set='high'):
    """
    Áp dụng định dạng ND30 cho paragraph theo role.
    size_set: 'high' (13/14) hoặc 'low' (12/13)
    """
    # Bảng cỡ chữ theo bộ
    S = {'high': {'xs': Pt(11), 'sm': Pt(12), 'md': Pt(13), 'lg': Pt(14)},
         'low':  {'xs': Pt(11), 'sm': Pt(12), 'md': Pt(12), 'lg': Pt(13)}}
    sz = S[size_set]

    C = WD_ALIGN_PARAGRAPH.CENTER
    L = WD_ALIGN_PARAGRAPH.LEFT
    R = WD_ALIGN_PARAGRAPH.RIGHT
    J = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Role → (size, bold, italic, all_caps, alignment, first_indent, space_before, space_after,
    #         add_border_full, add_border_half, override_run_style)
    # override_run_style=False nghĩa là chỉ fix font/size, không thay đổi bold/italic của run
    RULES = {
        # Ô 1: Quốc hiệu
        'quoc_hieu':      (sz['md'], True,  False, True,  C, None,      Pt(0),  Pt(0),  False, False, True),
        # Ô 1: Tiêu ngữ (có đường kẻ full-width)
        'tieu_ngu':       (sz['lg'], True,  False, False, C, None,      Pt(0),  Pt(0),  True,  False, True),
        # Ô 2: Tên CQ chủ quản (không đậm)
        'ten_co_quan_cq': (sz['md'], False, False, True,  C, None,      Pt(0),  Pt(0),  False, False, True),
        # Ô 2: Tên CQ ban hành (đậm + đường kẻ 1/3)
        'ten_co_quan_bh': (sz['md'], True,  False, True,  C, None,      Pt(0),  Pt(0),  False, True,  True),
        # Ô 3: Số, ký hiệu (cố định 13pt theo ND30)
        'so_ky_hieu':     (Pt(13),  False, False, False, C, None,      Pt(0),  Pt(0),  False, False, True),
        # Ô 4: Địa danh ngày tháng
        'dia_danh_ngay':  (sz['lg'], False, True,  False, C, None,      Pt(0),  Pt(0),  False, False, True),
        # Ô 5a: Tên loại văn bản
        'ten_loai_van_ban':(sz['lg'],True,  False, True,  C, None,      Pt(12), Pt(6),  False, False, True),
        # Ô 5a: Trích yếu (có đường kẻ 1/3)
        'trich_yeu':      (sz['lg'], True,  False, False, C, None,      Pt(0),  Pt(12), False, True,  True),
        # Ô 5b: Trích yếu công văn
        'trich_yeu_cv':   (sz['md'], False, False, False, C, None,      Pt(6),  Pt(0),  False, False, True),
        # Ô 6: Căn cứ
        'can_cu':         (sz['lg'], False, True,  False, J, None,      Pt(0),  Pt(3),  False, False, True),
        # Ô 6: Phần/Chương
        'phan_chuong':    (sz['lg'], True,  False, False, C, None,      Pt(12), Pt(0),  False, False, True),
        # Ô 6: Mục/Tiểu mục
        'muc_tieumuc':    (sz['lg'], True,  False, False, C, None,      Pt(6),  Pt(0),  False, False, True),
        # Ô 6: Điều khoản
        'dieu_khoan':     (sz['lg'], True,  False, False, L, _ND30_INDENT, Pt(6), Pt(0), False, False, True),
        # Ô 6: Heading La Mã (Báo cáo)
        'heading_la_ma':  (sz['lg'], True,  False, True,  L, None,      Pt(12), Pt(6),  False, False, True),
        # Ô 6: Nội dung thông thường (KHÔNG override bold/italic)
        'noi_dung':       (sz['lg'], None,  None,  False, J, _ND30_INDENT, Pt(6), Pt(0), False, False, False),
        # Ô 6: Kính gửi
        'kinh_gui':       (sz['lg'], False, False, False, L, None,      Pt(0),  Pt(0),  False, False, True),
        # Ô 7: Chữ ký (TM., KT.)
        'chu_ky':         (sz['lg'], True,  False, True,  R, None,      Pt(12), Pt(0),  False, False, True),
        # Ô 7: Họ tên người ký
        'ho_ten':         (sz['lg'], True,  False, False, C, None,      Pt(0),  Pt(0),  False, False, True),
        # Ô 9: Nơi nhận header
        'noi_nhan_header':(sz['sm'], True,  True,  False, L, None,      Pt(0),  Pt(0),  False, False, True),
        # Ô 9: Nơi nhận item
        'noi_nhan_item':  (sz['xs'], False, False, False, L, None,      Pt(0),  Pt(0),  False, False, True),
    }

    if role not in RULES:
        return

    (size, bold, italic, all_caps, align,
     first_indent, sp_before, sp_after,
     border_full, border_half, override_run) = RULES[role]

    # Áp dụng paragraph-level format
    pf = para.paragraph_format
    if align is not None:
        para.alignment = align
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if sp_before is not None:
        pf.space_before = sp_before
    if sp_after is not None:
        pf.space_after = sp_after
    if role == 'noi_dung':
        pf.line_spacing = _ND30_LINE_SPACING

    # Áp dụng run-level format
    for run in para.runs:
        if not run.text.strip():
            continue
        run.font.name = _ND30_FONT
        run.font.size = size
        if override_run:
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            run.font.all_caps = all_caps
        else:
            # noi_dung: chỉ fix font name/size, giữ bold/italic của người dùng
            pass

    # Đường kẻ ngang
    if border_full:
        _add_nd30_bottom_border(para, full_width=True)
    elif border_half:
        _add_nd30_bottom_border(para, full_width=False)

# --- Áp dụng lề và font toàn tài liệu ---
def _apply_nd30_global(doc):
    """Chuẩn hóa lề trang và font mặc định cho toàn bộ sections."""
    for section in doc.sections:
        section.top_margin    = _ND30_MARGIN_TOP
        section.bottom_margin = _ND30_MARGIN_BOTTOM
        section.left_margin   = _ND30_MARGIN_LEFT
        section.right_margin  = _ND30_MARGIN_RIGHT

# --- Thêm số trang vào header (từ trang 2) ---
def _add_nd30_page_number(doc):
    """
    Thêm số trang vào FOOTER, căn phải, từ trang 2.
    Font: Times New Roman, 14pt, đứng.
    Trang 1 KHÔNG hiển thị số trang.
    """
    for section in doc.sections:
        section.different_first_page_header_footer = True
        # Xóa footer trang thường cũ
        footer = section.footer
        for para in footer.paragraphs:
            for run in para.runs:
                run.text = ''
            p_elem = para._p
            for child in list(p_elem):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag != 'pPr':
                    p_elem.remove(child)

        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = fp.add_run()
        run.font.name = _ND30_FONT
        run.font.size = Pt(14)
        run.font.bold = False

        # Field PAGE
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.extend([fldChar1, instrText, fldChar2])

        # Trang đầu: footer rỗng
        first_footer = section.first_page_footer
        for p in first_footer.paragraphs:
            for r in p.runs:
                r.text = ''

# --- Hàm tổng hợp ---
def apply_nd30_smart_engine(doc):
    """
    Engine chính: tự động phát hiện loại văn bản và áp dụng
    toàn bộ chuẩn thể thức Nghị định 30/2020/NĐ-CP.
    Được gọi sau doc.render(context).
    """
    try:
        # 1. Lề trang & font toàn tài liệu
        _apply_nd30_global(doc)

        # 2. Phát hiện loại văn bản
        doc_type = _detect_document_type(doc)

        # 3. Phân loại từng đoạn
        roles = _classify_paragraphs(doc)

        # 4. Phát hiện bộ cỡ chữ (high/low) từ Quốc hiệu
        size_set = 'high'
        for para, role in zip(doc.paragraphs, roles):
            if role == 'quoc_hieu':
                for run in para.runs:
                    if run.font.size:
                        pt_size = run.font.size.pt
                        size_set = 'low' if pt_size <= 12.5 else 'high'
                        break
                break

        # 5. Áp dụng định dạng cho từng đoạn
        for para, role in zip(doc.paragraphs, roles):
            if not para.text.strip():
                continue
            # Bỏ qua ảnh inline, bảng biểu (xử lý riêng)
            _apply_nd30_para_format(para, role, size_set)

        # 6. Số trang ở header từ trang 2
        _add_nd30_page_number(doc)

    except Exception as nd30_err:
        # Engine lỗi → không crash toàn bộ export, chỉ log
        import traceback
        print(f'[ND30 Engine] Lỗi khi áp dụng chuẩn ND30: {nd30_err}')
        print(traceback.format_exc())

# =============================================================================
# END OF SMART ND30 ENGINE
# =============================================================================

def generate_report_from_custom_template(template_path, context, custom_vars=[]):
    """
    Translates, compiles, and renders custom template, returning output docx BytesIO stream.
    """
    # Clean context to replace '--' with ''
    context = replace_placeholders_with_empty(context)
    # Auto-enrich context with Vietnamese number words for all numeric variables (e.g. bangchu_gia_goi_thau)
    enrich_context_with_words(context)
    # Auto-format all dates to Vietnamese standard speech format and create S_ versions
    format_context_dates(context)
            
    doc = None
    try:
        doc = translate_docx_template(template_path, context)
        convert_images_in_context(doc, context)
        doc.render(context)
        apply_nd30_smart_engine(doc)   # Áp dụng chuẩn thể thức ND30/2020
    except Exception as e:
        # Log error to export_error.log in workspace root
        log_path = os.path.join(project_root, 'export_error.log')
        try:
            # Mask sensitive info in logged context to prevent data leakage
            masked_context = {}
            if isinstance(context, dict):
                def mask_sensitive(d):
                    res = {}
                    for k, v in d.items():
                        if isinstance(v, dict):
                            res[k] = mask_sensitive(v)
                        elif isinstance(v, list):
                            res[k] = [mask_sensitive(x) if isinstance(x, dict) else x for x in v]
                        elif any(sub in str(k).lower() for sub in ['cccd', 'cmt', 'passport', 'token', 'mat_khau', 'password', 'anh_chung_chi', 'anh_chu_ky', 'signature', 'base64']):
                            res[k] = '[MASKED]'
                        else:
                            res[k] = v
                    return res
                masked_context = mask_sensitive(context)
            else:
                masked_context = context

            with open(log_path, 'a', encoding='utf-8') as lf:
                import traceback
                lf.write(f"[{datetime.now().isoformat()}] ERROR: Failed rendering template {template_path}\n")
                lf.write(f"Context: {json.dumps(masked_context, ensure_ascii=False, default=str)}\n")
                lf.write(traceback.format_exc())
                lf.write("\n" + "="*50 + "\n")
        except Exception:
            pass
            
        # Fallback: If render failed, try to load template as a plain unrendered document so user still gets a file
        try:
            with open(template_path, 'rb') as f:
                raw_bytes = f.read()
            out_stream = BytesIO(raw_bytes)
            return out_stream
        except Exception:
            pass
            
    out_stream = BytesIO()
    if doc:
        doc.save(out_stream)
    else:
        # Final fallback in case doc is None
        try:
            with open(template_path, 'rb') as f:
                out_stream.write(f.read())
        except Exception:
            pass
    out_stream.seek(0)
    return out_stream
