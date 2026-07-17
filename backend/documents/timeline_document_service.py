"""Create and render the system-owned package timeline Word template."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# compact_reference_guide with the named ``timeline_a4_landscape`` geometry
# and the BiddingFlow teal table-header override requested by the checklist.
FONT_NAME = "Calibri"
HEADER_FILL = "17627D"
GROUP_FILL = "D9EEF5"
ALT_FILL = "F4FAFC"
GRID_COLOR = "31566B"
PLANNED_RED = RGBColor(199, 53, 67)
TABLE_WIDTHS_DXA = (620, 3900, 2900, 2500, 1700, 3818)
TABLE_WIDTH_DXA = sum(TABLE_WIDTHS_DXA)


def _set_run_font(run, *, size=9, bold=False, color=None):
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_margins(cell, *, top=70, start=80, bottom=70, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_table_geometry(table):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    tbl_width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_width.set(qn("w:type"), "dxa")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "0")
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in TABLE_WIDTHS_DXA:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        unique_cells = []
        seen_cells = set()
        for cell in row.cells:
            marker = id(cell._tc)
            if marker not in seen_cells:
                seen_cells.add(marker)
                unique_cells.append(cell)
        row_widths = (TABLE_WIDTHS_DXA[0], sum(TABLE_WIDTHS_DXA[1:])) if len(unique_cells) == 2 else TABLE_WIDTHS_DXA
        for index, cell in enumerate(unique_cells):
            width = row_widths[min(index, len(row_widths) - 1)]
            cell.width = Cm(width / 1440 * 2.54)
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:color"), GRID_COLOR)


def _add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, separate, fallback, end))
    _set_run_font(run, size=8)


def _format_cell(cell, text="", *, bold=False, size=8.5, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(str(text or ""))
    _set_run_font(run, size=size, bold=bold, color=color)
    return run


def _replace_tokens(document, replacements):
    containers = list(document.paragraphs)
    for section in document.sections:
        containers.extend(section.header.paragraphs)
        containers.extend(section.footer.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)
    for paragraph in containers:
        original = "".join(run.text for run in paragraph.runs)
        updated = original
        for token, value in replacements.items():
            updated = updated.replace(token, str(value or ""))
        if updated != original:
            for run in paragraph.runs:
                run.text = ""
            target = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            target.text = updated
            _set_run_font(target, size=8.5)


def create_timeline_template(output_path):
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.15)
    section.right_margin = Cm(1.15)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(header.add_run("[[ORG_NAME]]  |  Gói: [[PACKAGE_CODE]]"), size=8, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(footer.add_run("Trang "), size=8)
    _add_field(footer, "PAGE")
    _set_run_font(footer.add_run(" / "), size=8)
    _add_field(footer, "NUMPAGES")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    _set_run_font(title.add_run("CHECK LIST"), size=14, bold=True)
    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(7)
    _set_run_font(metadata.add_run("[[PACKAGE_CODE]] - [[PACKAGE_NAME]]  |  Kế hoạch: [[PLAN_CODE]]"), size=8.5, bold=True)

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ("STT", "Công việc", "Đơn vị ban hành", "Số văn bản", "Ngày ký", "Ghi chú")
    for index, label in enumerate(headers):
        _format_cell(table.rows[0].cells[index], label, bold=True, size=8.5, color=RGBColor(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_fill(table.rows[0].cells[index], HEADER_FILL)
    _set_repeat_table_header(table.rows[0])
    _prevent_row_split(table.rows[0])
    _set_table_geometry(table)
    _set_table_borders(table)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    _set_run_font(note.add_run("[[PLANNED_NOTE]]"), size=8, color=PLANNED_RED)
    _set_run_font(note.add_run("  |  Ngày xuất: [[GENERATED_DATE]]"), size=8)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def render_timeline_document(template_path, context):
    document = Document(template_path)
    package = context.get("goi_thau") or {}
    plan = context.get("ke_hoach") or {}
    organization = context.get("to_chuc") or {}
    replacements = {
        "[[ORG_NAME]]": organization.get("ten_to_chuc") or organization.get("name") or "",
        "[[PACKAGE_CODE]]": package.get("ma_goi_thau") or "",
        "[[PACKAGE_NAME]]": package.get("ten_goi_thau") or "",
        "[[PLAN_CODE]]": plan.get("ma_ke_hoach") or "",
        "[[PLANNED_NOTE]]": context.get("planned_date_note") or "",
        "[[GENERATED_DATE]]": context.get("generated_date") or "",
    }
    _replace_tokens(document, replacements)
    table = document.tables[0]
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    alternating = False
    for section in context.get("timeline_sections") or []:
        group_row = table.add_row()
        group_row.cells[1].merge(group_row.cells[5])
        _format_cell(group_row.cells[0], section.get("code"), bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _format_cell(group_row.cells[1], section.get("title"), bold=True, size=9)
        for cell in group_row.cells:
            _set_cell_fill(cell, GROUP_FILL)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
        _prevent_row_split(group_row)

        for item in section.get("items") or []:
            row = table.add_row()
            values = (
                item.get("display_code") or item.get("ma_moc"), item.get("cong_viec"), item.get("don_vi_ban_hanh"),
                item.get("so_van_ban"), item.get("display_date"), item.get("ghi_chu"),
            )
            for index, value in enumerate(values):
                align = WD_ALIGN_PARAGRAPH.CENTER if index in {0, 4} else WD_ALIGN_PARAGRAPH.LEFT
                color = PLANNED_RED if index == 4 and item.get("is_planned_date") else None
                _format_cell(row.cells[index], value, size=8.2, color=color, align=align)
                if alternating:
                    _set_cell_fill(row.cells[index], ALT_FILL)
            _prevent_row_split(row)
            alternating = not alternating
    _set_table_geometry(table)
    _set_table_borders(table)
    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream
