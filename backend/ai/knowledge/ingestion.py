"""Prepare and activate expert-approved local knowledge documents."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from hashlib import sha256
from pathlib import Path
import re
from typing import Any
import uuid

from docx import Document


DOCUMENT_TYPES = frozenset(
    {
        "LEGAL_DOCUMENT",
        "INTERNAL_POLICY",
        "PROCESS_GUIDE",
        "BIDDINGFLOW_HELP",
        "TEMPLATE_GUIDE",
        "APPROVED_QA",
    }
)
CONFIDENTIALITY_LEVELS = frozenset({"public", "internal", "confidential"})
ALLOWED_SUFFIXES = frozenset({".md", ".txt", ".docx"})
MAX_DOCUMENT_BYTES = 20 * 1024 * 1024
MAX_CHUNK_CHARS = 1800
_HEADING_RE = re.compile(r"^\s{0,3}#{1,6}\s+(.+?)\s*#*\s*$")
_SUSPICIOUS_INSTRUCTIONS = (
    re.compile(r"ignore\s+(all\s+)?previous\s+instructions?", re.IGNORECASE),
    re.compile(r"reveal\s+(the\s+)?system\s+prompt", re.IGNORECASE),
    re.compile(r"bỏ\s+qua\s+(mọi\s+)?(chỉ dẫn|hướng dẫn)\s+trước", re.IGNORECASE),
)


class KnowledgeIngestionError(ValueError):
    """Document failed validation and must not enter the active registry."""


@dataclass(frozen=True)
class PreparedChunk:
    section: str
    content: str
    page_number: int | None = None


@dataclass(frozen=True)
class PreparedDocument:
    metadata: dict[str, Any]
    content_hash: str
    chunks: tuple[PreparedChunk, ...]


def _required_text(metadata: dict[str, Any], field: str) -> str:
    value = str(metadata.get(field) or "").strip()
    if not value:
        raise KnowledgeIngestionError(f"Thiếu metadata bắt buộc: {field}.")
    return value


def _iso_date(metadata: dict[str, Any], field: str, *, optional: bool = False) -> str | None:
    raw = metadata.get(field)
    if raw in (None, "") and optional:
        return None
    value = _required_text(metadata, field)
    try:
        return date.fromisoformat(value[:10]).isoformat()
    except ValueError as exc:
        raise KnowledgeIngestionError(f"Metadata {field} phải có dạng YYYY-MM-DD.") from exc


def _validate_metadata(metadata: dict[str, Any]) -> dict[str, Any]:
    if not isinstance(metadata, dict):
        raise KnowledgeIngestionError("Metadata tài liệu phải là JSON object.")
    document_type = _required_text(metadata, "document_type").upper()
    if document_type not in DOCUMENT_TYPES:
        raise KnowledgeIngestionError("document_type không được hỗ trợ.")
    status = _required_text(metadata, "status").casefold()
    if status != "approved":
        raise KnowledgeIngestionError("Chỉ tài liệu có status=approved mới được kích hoạt.")
    confidentiality = _required_text(metadata, "confidentiality").casefold()
    if confidentiality not in CONFIDENTIALITY_LEVELS:
        raise KnowledgeIngestionError("confidentiality không hợp lệ.")
    organization_id = str(metadata.get("organization_id") or "").strip() or None
    if organization_id is None and confidentiality == "confidential":
        raise KnowledgeIngestionError("Tài liệu confidential phải thuộc một organization.")
    effective_from = _iso_date(metadata, "effective_from")
    effective_to = _iso_date(metadata, "effective_to", optional=True)
    if effective_to and effective_to < effective_from:
        raise KnowledgeIngestionError("effective_to không được trước effective_from.")
    source_url = str(metadata.get("source_url") or "").strip()
    if source_url and not (
        (source_url.startswith("/") and not source_url.startswith("//"))
        or source_url.startswith("https://")
        or source_url.startswith("http://")
    ):
        raise KnowledgeIngestionError("source_url phải là URL http(s) hoặc route nội bộ.")
    return {
        "title": _required_text(metadata, "title")[:300],
        "document_number": _required_text(metadata, "document_number")[:160],
        "issuing_authority": _required_text(metadata, "issuing_authority")[:300],
        "document_type": document_type,
        "issued_date": _iso_date(metadata, "issued_date"),
        "effective_from": effective_from,
        "effective_to": effective_to,
        "version": _required_text(metadata, "version")[:80],
        "status": status,
        "organization_id": organization_id,
        "confidentiality": confidentiality,
        "source_url": source_url[:2000],
    }


def _read_source(path: Path) -> str:
    suffix = path.suffix.casefold()
    if suffix not in ALLOWED_SUFFIXES:
        raise KnowledgeIngestionError("Chỉ hỗ trợ tài liệu .md, .txt và .docx.")
    try:
        size = path.stat().st_size
    except OSError as exc:
        raise KnowledgeIngestionError("Không thể đọc tệp tài liệu.") from exc
    if size <= 0 or size > MAX_DOCUMENT_BYTES:
        raise KnowledgeIngestionError("Kích thước tài liệu không hợp lệ hoặc vượt quá 20 MB.")
    if suffix == ".docx":
        try:
            document = Document(path)
        except Exception as exc:
            raise KnowledgeIngestionError("Tệp DOCX không hợp lệ.") from exc
        lines: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = str(paragraph.style.name if paragraph.style else "")
            if style.casefold().startswith("heading"):
                level_match = re.search(r"(\d+)", style)
                level = max(1, min(6, int(level_match.group(1)) if level_match else 2))
                lines.append(f"{'#' * level} {text}")
            else:
                lines.append(text)
        return "\n\n".join(lines).strip()
    try:
        return path.read_text(encoding="utf-8-sig").replace("\r\n", "\n").replace("\r", "\n").strip()
    except (OSError, UnicodeError) as exc:
        raise KnowledgeIngestionError("Tệp văn bản phải dùng UTF-8 hợp lệ.") from exc


def _split_large(section: str, text: str) -> list[PreparedChunk]:
    text = text.strip()
    if not text:
        return []
    chunks = []
    remaining = text
    while len(remaining) > MAX_CHUNK_CHARS:
        cut = remaining.rfind("\n", 0, MAX_CHUNK_CHARS)
        if cut < MAX_CHUNK_CHARS // 2:
            cut = remaining.rfind(". ", 0, MAX_CHUNK_CHARS)
            if cut >= MAX_CHUNK_CHARS // 2:
                cut += 1
        if cut < MAX_CHUNK_CHARS // 2:
            cut = MAX_CHUNK_CHARS
        chunks.append(PreparedChunk(section=section, content=remaining[:cut].strip()))
        remaining = remaining[cut:].strip()
    if remaining:
        chunks.append(PreparedChunk(section=section, content=remaining))
    return chunks


def _chunk_markdown(content: str, fallback_section: str) -> tuple[PreparedChunk, ...]:
    chunks: list[PreparedChunk] = []
    section = fallback_section
    paragraphs: list[str] = []

    def flush() -> None:
        if paragraphs:
            chunks.extend(_split_large(section, "\n\n".join(paragraphs)))
            paragraphs.clear()

    for block in re.split(r"\n\s*\n", content):
        clean = block.strip()
        if not clean:
            continue
        heading = _HEADING_RE.match(clean)
        if heading:
            flush()
            section = heading.group(1).strip()[:300]
            continue
        paragraphs.append(clean)
    flush()
    return tuple(chunks)


def prepare_document(path: str | Path, metadata: dict[str, Any]) -> PreparedDocument:
    """Validate and deterministically prepare a local document without writing it."""

    source = Path(path)
    clean_metadata = _validate_metadata(metadata)
    content = _read_source(source)
    if not content:
        raise KnowledgeIngestionError("Tài liệu không có nội dung văn bản.")
    for pattern in _SUSPICIOUS_INSTRUCTIONS:
        if pattern.search(content):
            raise KnowledgeIngestionError("Tài liệu chứa instruction đáng ngờ và cần chuyên gia kiểm tra.")
    chunks = _chunk_markdown(content, clean_metadata["title"])
    if not chunks:
        raise KnowledgeIngestionError("Không thể tạo đoạn tri thức từ tài liệu.")
    return PreparedDocument(
        metadata=clean_metadata,
        content_hash=sha256(content.encode("utf-8")).hexdigest(),
        chunks=chunks,
    )


def ingest_approved_document(
    cursor,
    path: str | Path,
    metadata: dict[str, Any],
    *,
    approved_by: str,
    document_id: str | None = None,
) -> dict[str, Any]:
    """Activate one approved version and retire the previous active version."""

    approver = str(approved_by or "").strip()
    if not approver:
        raise KnowledgeIngestionError("approved_by là bắt buộc.")
    prepared = prepare_document(path, metadata)
    values = prepared.metadata
    duplicate = cursor.execute(
        """SELECT id FROM ai_knowledge_documents
           WHERE content_hash = ?
             AND COALESCE(organization_id, '') = COALESCE(?, '')
           LIMIT 1""",
        (prepared.content_hash, values["organization_id"]),
    ).fetchone()
    if duplicate:
        raise KnowledgeIngestionError("Tài liệu có content_hash đã tồn tại trong registry.")
    resolved_id = str(document_id or uuid.uuid4())
    cursor.execute(
        """UPDATE ai_knowledge_documents
           SET status = 'retired', updated_at = CURRENT_TIMESTAMP
           WHERE status = 'active'
             AND document_type = ? AND document_number = ?
             AND COALESCE(organization_id, '') = COALESCE(?, '')""",
        (values["document_type"], values["document_number"], values["organization_id"]),
    )
    cursor.execute(
        """INSERT INTO ai_knowledge_documents
           (id, organization_id, title, document_number, issuing_authority,
            document_type, issued_date, effective_from, effective_to, version,
            status, confidentiality, approved_by, approved_at, source_url,
            content_hash, source_filename)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'active', ?, ?,
                   CURRENT_TIMESTAMP, ?, ?, ?)""",
        (
            resolved_id,
            values["organization_id"],
            values["title"],
            values["document_number"],
            values["issuing_authority"],
            values["document_type"],
            values["issued_date"],
            values["effective_from"],
            values["effective_to"],
            values["version"],
            values["confidentiality"],
            approver,
            values["source_url"],
            prepared.content_hash,
            Path(path).name[:255],
        ),
    )
    cursor.executemany(
        """INSERT INTO ai_knowledge_chunks
           (id, document_id, chunk_index, section, page_number, content, char_count)
           VALUES (?, ?, ?, ?, ?, ?, ?)""",
        [
            (
                str(uuid.uuid4()),
                resolved_id,
                index,
                chunk.section,
                chunk.page_number,
                chunk.content,
                len(chunk.content),
            )
            for index, chunk in enumerate(prepared.chunks)
        ],
    )
    return {
        "documentId": resolved_id,
        "contentHash": prepared.content_hash,
        "chunkCount": len(prepared.chunks),
        "status": "active",
    }
